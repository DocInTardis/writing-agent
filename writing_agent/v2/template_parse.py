"""Template Parse module.

This module belongs to `writing_agent.v2` in the writing-agent codebase.
"""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path
from dataclasses import dataclass
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


@dataclass(frozen=True)
class TemplateInfo:
    name: str
    required_h2: list[str]
    outline: list[tuple[int, str]]


def parse_template_file(path: Path, name: str) -> TemplateInfo:
    path = prepare_template_file(path)
    suffix = path.suffix.lower()
    if suffix in {".docx"}:
        try:
            return TemplateInfo(name=name, required_h2=_parse_docx_h2(path), outline=_parse_docx_outline_strict(path))
        except Exception:
            converted = _convert_doc_to_docx(path, force=True)
            if converted and converted.exists():
                return TemplateInfo(
                    name=name,
                    required_h2=_parse_docx_h2(converted),
                    outline=_parse_docx_outline_strict(converted),
                )
            raise
    if suffix in {".html", ".htm"}:
        html = path.read_text(encoding="utf-8", errors="replace")
        return TemplateInfo(name=name, required_h2=_parse_html_h2(html), outline=_parse_html_outline(html))
    # default: treat as text
    text = path.read_text(encoding="utf-8", errors="replace")
    return TemplateInfo(name=name, required_h2=_parse_text_h2(text), outline=_parse_text_outline(text))


def prepare_template_file(path: Path) -> Path:
    suffix = path.suffix.lower()
    if suffix == ".doc":
        converted = _convert_doc_to_docx(path)
        return converted or path
    if suffix == ".docx" and (not _looks_like_docx(path) or not _is_docx_readable(path)):
        converted = _convert_doc_to_docx(path, force=True)
        return converted or path
    return path


def _convert_doc_to_docx(path: Path, *, force: bool = False) -> Path | None:
    src = Path(path)
    if not src.exists():
        return None
    if src.suffix.lower() == ".docx":
        out_path = src.with_name(src.stem + "_converted.docx")
    else:
        out_path = src.with_suffix(".docx")
    if out_path.exists() and not force:
        return out_path

    soffice = shutil.which("soffice")
    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", str(src.parent), str(src)],
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )
            if out_path.exists():
                return out_path
            if src.suffix.lower() == ".docx" and not out_path.exists():
                alt = src.with_suffix(".doc")
                try:
                    alt.write_bytes(src.read_bytes())
                    subprocess.run(
                        [soffice, "--headless", "--convert-to", "docx", "--outdir", str(src.parent), str(alt)],
                        capture_output=True,
                        text=True,
                        timeout=60,
                        check=False,
                    )
                    alt_out = alt.with_suffix(".docx")
                    if alt_out.exists():
                        return alt_out
                finally:
                    if alt.exists():
                        alt.unlink(missing_ok=True)
        except Exception:
            pass

    try:
        in_path = str(src).replace("'", "''")
        out = str(out_path).replace("'", "''")
        ps_script = (
            "$in = '{0}';"
            "$out = '{1}';"
            "$word = New-Object -ComObject Word.Application;"
            "$word.Visible = $false;"
            "$doc = $word.Documents.Open($in, $false, $true);"
            "$doc.SaveAs([ref]$out, 16);"
            "$doc.Close();"
            "$word.Quit();"
        ).format(in_path, out)
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps_script],
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
        )
        if out_path.exists():
            return out_path
    except Exception:
        return None
    return None


def _looks_like_docx(path: Path) -> bool:
    try:
        with path.open("rb") as f:
            sig = f.read(4)
        return sig.startswith(b"PK")
    except Exception:
        return False


def _is_docx_readable(path: Path) -> bool:
    try:
        Document(str(path))
        return True
    except Exception:
        return False


def _parse_docx_h2(path: Path) -> list[str]:
    outline = _parse_docx_outline_strict(path)
    h1_list: list[str] = []
    h2_list: list[str] = []
    for lvl, txt in outline:
        if lvl == 1:
            _push_unique(h1_list, txt)
        elif lvl == 2:
            _push_unique(h2_list, txt)
    if len(h1_list) >= 2:
        return h1_list
    if h2_list:
        return h2_list
    return h1_list


def _parse_docx_outline(path: Path) -> list[tuple[int, str]]:
    doc = Document(str(path))
    out: list[tuple[int, str]] = []
    chapter_re = re.compile(r"^\s*第[一二三四五六七八九十百0-9]+章\s*.*$")
    sub_re = re.compile(
        r"^\s*(?:"
        r"\d+(?:\.\d+)+\s+.+"
        r"|第[一二三四五六七八九十百0-9]+节\s*.*"
        r"|[一二三四五六七八九十]+、\s*.+"
        r"|（[一二三四五六七八九十]+）\s*.+"
        r")$",
    )
    for p in _iter_docx_paragraphs(doc):
        txt = (p.text or "").strip()
        if not txt:
            continue
        style = ""
        try:
            style = str(p.style.name or "")
        except Exception:
            style = ""
        s = style.lower()
        if _is_toc_style(s) or _looks_like_toc_entry(txt):
            continue
        is_h1 = ("heading 1" in s) or ("标题 1" in style) or ("标题1" in style)
        is_h2 = ("heading 2" in s) or ("标题 2" in style) or ("标题2" in style)
        if is_h1 or chapter_re.match(txt):
            out.append((1, txt))
            continue
        if is_h2 or sub_re.match(txt):
            out.append((2, txt))
    return out


def _parse_docx_outline_strict(path: Path) -> list[tuple[int, str]]:
    doc = Document(str(path))
    style_hits: list[tuple[int, str]] = []
    for p in _iter_docx_paragraphs(doc):
        txt = (p.text or "").strip()
        if not txt:
            continue
        style = ""
        try:
            style = str(p.style.name or "")
        except Exception:
            style = ""
        s = style.lower()
        if _is_toc_style(s) or _looks_like_toc_entry(txt):
            continue
        if ("heading 1" in s) or ("标题 1" in style) or ("标题1" in style):
            style_hits.append((1, txt))
            continue
        if (
            ("heading 2" in s)
            or ("heading 3" in s)
            or ("标题 2" in style)
            or ("标题2" in style)
            or ("标题 3" in style)
            or ("标题3" in style)
        ):
            style_hits.append((2, txt))
    if style_hits:
        return style_hits
    return _parse_docx_outline(path)


def _iter_docx_paragraphs(doc: Document):
    body = doc.element.body  # type: ignore[attr-defined]
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def _is_toc_style(style_lower: str) -> bool:
    return ("toc" in style_lower) or ("目录" in style_lower)


def _looks_like_toc_entry(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if re.search(r"[\\.·…]{2,}\s*\d+$", t):
        return True
    return False


def _parse_html_h2(html: str) -> list[str]:
    headings = re.findall(r"<h2\b[^>]*>(.*?)</h2>", html or "", flags=re.IGNORECASE | re.DOTALL)
    out: list[str] = []
    for h in headings:
        txt = re.sub(r"<[^>]+>", "", h).strip()
        if txt:
            _push_unique(out, txt)
    return out


def _parse_html_outline(html: str) -> list[tuple[int, str]]:
    out: list[tuple[int, str]] = []
    for m in re.finditer(r"<h([12])\b[^>]*>(.*?)</h\\1>", html or "", flags=re.IGNORECASE | re.DOTALL):
        level = int(m.group(1))
        txt = re.sub(r"<[^>]+>", "", m.group(2) or "").strip()
        if txt:
            out.append((level, txt))
    return out


def _parse_text_h2(text: str) -> list[str]:
    out: list[str] = []
    for line in (text or "").splitlines():
        m = re.match(r"^\s*##\s+(.+?)\s*$", line)
        if m:
            _push_unique(out, m.group(1).strip())
    return out


def _parse_text_outline(text: str) -> list[tuple[int, str]]:
    out: list[tuple[int, str]] = []
    for line in (text or "").splitlines():
        m1 = re.match(r"^\s*#\s+(.+?)\s*$", line)
        if m1:
            out.append((1, m1.group(1).strip()))
            continue
        m2 = re.match(r"^\s*##\s+(.+?)\s*$", line)
        if m2:
            out.append((2, m2.group(1).strip()))
    return out


def _push_unique(out: list[str], s: str) -> None:
    if s and s not in out:
        out.append(s)
