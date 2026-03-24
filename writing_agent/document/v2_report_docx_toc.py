from __future__ import annotations

import re
import unicodedata
from collections.abc import Callable
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm, Pt

from writing_agent.document.v2_report_docx_helpers import (
    _add_field_simple,
    _add_inline_runs,
    _add_internal_hyperlink,
    _apply_toc_heading_style,
    _sanitize_heading_text,
)
from writing_agent.v2.doc_format import DocBlock


def collect_toc_entries(blocks: list[DocBlock], *, levels: int) -> list[tuple[int, str]]:
    max_level = max(1, min(4, int(levels or 3)))
    out: list[tuple[int, str]] = []
    seen: set[tuple[int, str]] = set()
    for block in blocks:
        if block.type != "heading":
            continue
        level = max(1, min(6, int(block.level or 1)))
        if level > max_level:
            continue
        title = _sanitize_heading_text((block.text or "").strip())
        if not title:
            continue
        key = (level, title)
        if key in seen:
            continue
        seen.add(key)
        out.append((level, title))
        if len(out) >= 120:
            break
    return out


def build_toc_anchor_entries(entries: list[tuple[int, str]], *, entry_factory: Callable[..., Any]) -> list[Any]:
    def _bookmark_name(title: str, idx: int) -> str:
        slug = re.sub(r"[^0-9A-Za-z_]+", "_", str(title or "").strip())
        slug = re.sub(r"_+", "_", slug).strip("_")
        if not slug:
            slug = f"h{idx:03d}"
        if not re.match(r"^[A-Za-z_]", slug):
            slug = f"h_{slug}"
        return f"toc_{idx:03d}_{slug[:24]}"

    out: list[Any] = []
    for idx, (level, title) in enumerate(entries, start=1):
        out.append(entry_factory(level=int(level or 1), title=str(title or ""), anchor=_bookmark_name(title, idx)))
    return out


def display_width(text: str) -> int:
    total = 0
    for ch in str(text or ""):
        total += 2 if unicodedata.east_asian_width(ch) in {"W", "F"} else 1
    return total


def estimate_toc_pages(entries: list[tuple[int, str]], blocks: list[DocBlock], *, start_page: int = 1) -> list[int]:
    if not entries:
        return []

    def _norm(value: str) -> str:
        return re.sub(r"\s+", "", str(value or "").strip().lower())

    def _block_cost(block: DocBlock) -> int:
        block_type = str(block.type or "")
        if block_type == "heading":
            return 70
        if block_type == "paragraph":
            return max(40, len(str(block.text or "")))
        if block_type == "list":
            items = getattr(block, "items", None) or []
            try:
                return max(80, sum(len(str(item or "")) for item in list(items)))
            except Exception:
                return 120
        if block_type == "table":
            return 700
        if block_type == "figure":
            return 600
        return 120

    key_to_indices: dict[tuple[int, str], list[int]] = {}
    for idx, (lvl, title) in enumerate(entries):
        key = (int(lvl or 1), _norm(title))
        key_to_indices.setdefault(key, []).append(idx)
    pages = [max(1, int(start_page)) for _ in entries]
    page = max(1, int(start_page))
    budget = 0
    page_capacity = 1800
    for block in blocks:
        if block.type == "heading":
            level = max(1, min(6, int(block.level or 1)))
            title = _sanitize_heading_text((block.text or "").strip())
            key = (level, _norm(title))
            queue = key_to_indices.get(key) or []
            if queue:
                idx = queue.pop(0)
                pages[idx] = page
        budget += _block_cost(block)
        while budget >= page_capacity:
            page += 1
            budget -= page_capacity
    return pages


def render_toc_preview_text(entries: list[tuple[int, str]], blocks: list[DocBlock]) -> str:
    if not entries:
        return ""
    pages = estimate_toc_pages(entries, blocks, start_page=1)
    target_width = 72
    lines: list[str] = []
    for idx, (level, title) in enumerate(entries):
        depth = max(1, min(4, int(level or 1)))
        indent = "  " * max(0, depth - 1)
        text = f"{indent}{title}".rstrip()
        page_num = max(1, int(pages[idx] if idx < len(pages) else 1))
        fill = max(6, target_width - display_width(text) - len(str(page_num)))
        lines.append(f"{text}{'.' * fill}{page_num}")
    return "\n".join(lines).strip()


def add_toc(
    doc: Document,
    *,
    levels: int,
    entries: list[Any] | None,
    blocks: list[DocBlock] | None,
    resolve_toc_style_name: Callable[[Document, int], str | None],
    dynamic_toc_field_enabled: bool,
    toc_hyperlink_enabled: bool,
    toc_field_lock_enabled: bool,
    toc_clickable_links_enabled: bool,
) -> None:
    heading = doc.add_paragraph()
    try:
        heading.style = "TOC Heading"
    except Exception:
        pass
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_toc_heading_style(heading, "目录")

    toc_entries = list(entries or [])
    plain_entries = [
        (int(_entry_value(item, "level") or 1), str(_entry_value(item, "title") or ""))
        for item in toc_entries
    ]
    if toc_clickable_links_enabled:
        if not toc_entries:
            paragraph = doc.add_paragraph("无目录项")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            return
        pages = estimate_toc_pages(plain_entries, list(blocks or []), start_page=1)
        for idx, entry in enumerate(toc_entries):
            level = max(1, min(4, int(_entry_value(entry, "level") or 1)))
            title = str(_entry_value(entry, "title") or "")
            anchor = str(_entry_value(entry, "anchor") or "").strip()
            page_num = max(1, int(pages[idx] if idx < len(pages) else 1))
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                toc_style_name = resolve_toc_style_name(doc, level)
                if toc_style_name:
                    paragraph.style = toc_style_name
            except Exception:
                pass
            try:
                paragraph.paragraph_format.left_indent = Cm(max(0.0, float(level - 1) * 0.75))
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            try:
                section = doc.sections[-1]
                right_tab = section.page_width - section.left_margin - section.right_margin
                tab_stops = paragraph.paragraph_format.tab_stops
                try:
                    tab_stops.clear_all()
                except Exception:
                    pass
                tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            except Exception:
                pass
            if anchor:
                _add_internal_hyperlink(paragraph, title, anchor)
            else:
                _add_inline_runs(paragraph, title)
            paragraph.add_run("\t")
            if anchor:
                _add_internal_hyperlink(paragraph, str(page_num), anchor)
            else:
                paragraph.add_run(str(page_num))
        return

    if dynamic_toc_field_enabled:
        paragraph = doc.add_paragraph()
        switches = [f'\\\\o "1-{levels}"']
        if toc_hyperlink_enabled:
            switches.append("\\\\h")
        switches.extend(["\\\\z", "\\\\u"])
        instruction = "TOC " + " ".join(switches)
        preview = render_toc_preview_text(plain_entries, list(blocks or []))
        _add_field_simple(paragraph, instruction, preview, lock=toc_field_lock_enabled)
        return

    if not toc_entries:
        paragraph = doc.add_paragraph("无目录项")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return

    for entry in toc_entries:
        level = int(_entry_value(entry, "level") or 1)
        title = str(_entry_value(entry, "title") or "")
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            paragraph.paragraph_format.left_indent = Cm(max(0.0, float(level - 1) * 0.75))
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        except Exception:
            pass
        _add_inline_runs(paragraph, title)


def _entry_value(entry: Any, name: str) -> Any:
    if isinstance(entry, dict):
        return entry.get(name)
    return getattr(entry, name, None)
