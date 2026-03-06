"""V2 Report Docx module.

This module belongs to `writing_agent.document` in the writing-agent codebase.
"""

from __future__ import annotations

import os
import re
from io import BytesIO
from dataclasses import dataclass
from typing import Any
import unicodedata

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from writing_agent.models import FormattingRequirements
from writing_agent.v2.doc_format import DocBlock, ParsedDoc, parse_report_text
from writing_agent.v2.figure_render import render_figure_svg
from writing_agent.document.v2_report_docx_helpers import *  # noqa: F401,F403

try:
    import cairosvg
except Exception:  # pragma: no cover - optional dependency
    cairosvg = None


@dataclass(frozen=True)
class ExportPrefs:
    include_cover: bool = True
    include_toc: bool = True
    toc_levels: int = 3
    include_header: bool = True
    page_numbers: bool = True
    header_text: str = ""
    footer_text: str = ""
    page_margins_cm: float = 2.5
    page_margin_top_cm: float | None = None
    page_margin_bottom_cm: float | None = None
    page_margin_left_cm: float | None = None
    page_margin_right_cm: float | None = None
    page_size: str = "A4"


@dataclass(frozen=True)
class TocAnchorEntry:
    level: int
    title: str
    anchor: str


FIRST_LINE_INDENT_CM = 0.85


def _center_chapter_headings_enabled() -> bool:
    raw = str(os.environ.get("WRITING_AGENT_CENTER_CHAPTER_HEADINGS", "1")).strip().lower()
    return raw not in {"0", "false", "no", "off"}


def _update_fields_on_open_enabled() -> bool:
    # Default to disabled to avoid Word open-time field update prompts.
    raw = str(os.environ.get("WRITING_AGENT_DOCX_UPDATE_FIELDS_ON_OPEN", "0")).strip().lower()
    return raw in {"1", "true", "yes", "on"}


def _dynamic_toc_field_enabled() -> bool:
    # Keep dynamic TOC enabled by default (user requirement).
    raw = str(os.environ.get("WRITING_AGENT_DOCX_DYNAMIC_TOC_FIELD", "1")).strip().lower()
    return raw in {"1", "true", "yes", "on"}


def _toc_field_lock_enabled() -> bool:
    # Lock TOC field by default to reduce open-time update prompts.
    raw = str(os.environ.get("WRITING_AGENT_DOCX_TOC_FIELD_LOCK", "1")).strip().lower()
    return raw in {"1", "true", "yes", "on"}


def _toc_hyperlink_enabled() -> bool:
    # Hyperlink switch can trigger extra update behavior in some Word setups.
    raw = str(os.environ.get("WRITING_AGENT_DOCX_TOC_HYPERLINK", "0")).strip().lower()
    return raw in {"1", "true", "yes", "on"}


def _page_field_lock_enabled() -> bool:
    raw = str(os.environ.get("WRITING_AGENT_DOCX_PAGE_FIELD_LOCK", "1")).strip().lower()
    return raw in {"1", "true", "yes", "on"}


def _toc_clickable_links_enabled() -> bool:
    # Default to clickable manual TOC (Ctrl+Click jump), independent from field update.
    raw = str(os.environ.get("WRITING_AGENT_DOCX_TOC_CLICKABLE_LINKS", "1")).strip().lower()
    return raw in {"1", "true", "yes", "on"}


def _toc_footer_postprocess_enabled() -> bool:
    # Disabled by default: XML re-serialization of footer parts can introduce
    # namespace edge-cases that trigger Word repair prompts on some versions.
    raw = str(os.environ.get("WRITING_AGENT_DOCX_TOC_FOOTER_POSTPROCESS", "0")).strip().lower()
    return raw in {"1", "true", "yes", "on"}


class V2ReportDocxExporter:
    """\n    Builds a graduation-design style .docx:\n    - Cover page (no page number)\n    - TOC field (Word updates on open)\n    - Main content with heading numbering\n    - Footer page numbers\n    """

    def _resolve_toc_style_name(self, doc: Document, level: int) -> str | None:
        lvl = max(1, min(9, int(level or 1)))
        for candidate in (f"TOC {lvl}", f"TOC{lvl}"):
            try:
                _ = doc.styles[candidate]
                return candidate
            except Exception:
                continue
        return None

    def build_from_text(
        self,
        text: str,
        formatting: FormattingRequirements,
        prefs: ExportPrefs,
        *,
        template_path: str | None = None,
    ) -> bytes:
        parsed = parse_report_text(text)
        return self.build_from_parsed(parsed, formatting, prefs, template_path=template_path)

    def build_from_parsed(
        self,
        parsed: ParsedDoc,
        formatting: FormattingRequirements,
        prefs: ExportPrefs,
        *,
        template_path: str | None = None,
    ) -> bytes:
        heading_styles: dict[int, str] | None = None
        if template_path:
            doc = Document(template_path)
            heading_styles = _detect_template_heading_styles(doc)
            _truncate_template_body(doc)
        else:
            doc = Document()
        self._apply_page_setup(doc, prefs)
        self._apply_styles(doc, formatting)
        if _update_fields_on_open_enabled():
            _enable_update_fields(doc)
        else:
            _disable_update_fields(doc)

        raw_title = _strip_inline_markers((parsed.title or "").strip())
        if raw_title in {"\u672a\u547d\u540d\u6587\u6863", "\u672a\u547d\u540d\u62a5\u544a"}:
            raw_title = ""
        title = raw_title or "\u62a5\u544a"
        blocks = list(parsed.blocks or [])
        blocks = _remove_disallowed_sections(blocks)
        # 优化: 移除虚拟图表插入,让模型自然生成
        # blocks = _ensure_min_figures(blocks)
        # blocks = _ensure_min_tables(blocks)
        blocks = _ensure_reference_section(blocks)
        # Normalize heading structure before TOC/content emission.
        if blocks and blocks[0].type == "heading" and int(blocks[0].level or 0) == 1:
            blocks = blocks[1:]
        blocks = _promote_headings_if_no_h1(blocks)
        toc_entries = self._collect_toc_entries(blocks, levels=max(1, min(4, int(prefs.toc_levels))))
        toc_anchors = self._build_toc_anchor_entries(toc_entries)

        header_text = (prefs.header_text or "").strip() or title
        footer_text = (prefs.footer_text or "").strip()

        cover_section = doc.sections[0]
        toc_section = None
        main_section = cover_section

        if prefs.include_cover:
            self._add_cover(doc, title)
            if prefs.include_toc:
                toc_section = self._new_section(doc, start_page_numbering=True, numbering_format="upperRoman")
                self._add_toc(
                    doc,
                    levels=max(1, min(4, int(prefs.toc_levels))),
                    entries=toc_anchors,
                    blocks=blocks,
                )
                main_section = self._new_section(doc, start_page_numbering=True, numbering_format="decimal")
            else:
                main_section = self._new_section(doc, start_page_numbering=True, numbering_format="decimal")
        else:
            if prefs.include_toc:
                toc_section = cover_section
                _set_section_page_numbering(toc_section, start_at=1, numbering_format="upperRoman")
                self._add_toc(
                    doc,
                    levels=max(1, min(4, int(prefs.toc_levels))),
                    entries=toc_anchors,
                    blocks=blocks,
                )
                main_section = self._new_section(doc, start_page_numbering=True, numbering_format="decimal")
            else:
                main_section = cover_section
                _set_section_page_numbering(main_section, start_at=1, numbering_format="decimal")

        if prefs.page_numbers:
            _clear_header_footer(cover_section)
            _remove_section_page_numbering(cover_section)
            if toc_section is not None:
                _clear_header_footer(toc_section)
                self._set_footer_page_numbers(toc_section, "", page_format="ROMAN")
            _clear_header_footer(main_section)
            self._set_footer_page_numbers(main_section, footer_text, page_format=None)
        if prefs.include_header:
            self._set_header(main_section, header_text)

        self._emit_content(
            doc,
            blocks,
            formatting,
            number_headings=False,
            apply_para_format=True,
            heading_styles=heading_styles,
            toc_anchor_entries=toc_anchors,
        )
        _ensure_reference_citations(doc)
        _strip_cover_section_numbering(doc)
        buf = _save_doc(doc)
        if prefs.page_numbers and prefs.include_toc and _toc_footer_postprocess_enabled():
            buf = _postprocess_toc_footer_numbers(buf)
        return buf

    def _apply_page_setup(self, doc: Document, prefs: ExportPrefs) -> None:
        sec = doc.sections[0]
        size = str(prefs.page_size or "A4").upper()
        sizes = {
            "A4": (21.0, 29.7),
            "A5": (14.8, 21.0),
            "LETTER": (21.59, 27.94),
        }
        width_cm, height_cm = sizes.get(size, sizes["A4"])
        sec.page_width = Cm(width_cm)
        sec.page_height = Cm(height_cm)
        m = float(prefs.page_margins_cm or 2.5)
        top = _pick_margin(prefs.page_margin_top_cm, m + 0.5)
        bottom = _pick_margin(prefs.page_margin_bottom_cm, m)
        left = _pick_margin(prefs.page_margin_left_cm, m + 0.5)
        right = _pick_margin(prefs.page_margin_right_cm, max(1.5, m - 0.5))
        sec.top_margin = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin = Cm(left)
        sec.right_margin = Cm(right)

    def _apply_styles(self, doc: Document, formatting: FormattingRequirements) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = formatting.font_name or "宋体"
        normal.font.size = Pt(formatting.font_size_pt or 10.5)
        line_spacing = formatting.line_spacing or 1.5
        if line_spacing >= 4:
            normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            normal.paragraph_format.line_spacing = Pt(line_spacing)
        else:
            normal.paragraph_format.line_spacing = line_spacing
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
        try:
            r_pr = normal._element.get_or_add_rPr()  # type: ignore[attr-defined]
            r_fonts = r_pr.get_or_add_rFonts()
            r_fonts.set(qn("w:eastAsia"), formatting.font_name_east_asia or "宋体")
        except Exception:
            pass

        # Graduation-design-friendly heading defaults (can be overridden by user template later).
        try:
            h1 = doc.styles["Heading 1"]
            h1.font.name = formatting.heading1_font_name or formatting.font_name or "黑体"
            h1.font.size = Pt(formatting.heading1_size_pt or 22)
            h1.font.bold = True
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_pr = h1._element.get_or_add_rPr()  # type: ignore[attr-defined]
            r_pr.get_or_add_rFonts().set(
                qn("w:eastAsia"),
                formatting.heading1_font_name_east_asia or "黑体",
            )
            h1.paragraph_format.space_before = Pt(17)
            h1.paragraph_format.space_after = Pt(16.5)
            h1.paragraph_format.line_spacing = 2.4
        except Exception:
            pass
        try:
            h2 = doc.styles["Heading 2"]
            h2.font.name = formatting.heading2_font_name or formatting.font_name or "黑体"
            h2.font.size = Pt(formatting.heading2_size_pt or 16)
            h2.font.bold = True
            h2.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if _center_chapter_headings_enabled() else WD_ALIGN_PARAGRAPH.LEFT
            )
            r_pr = h2._element.get_or_add_rPr()  # type: ignore[attr-defined]
            r_pr.get_or_add_rFonts().set(
                qn("w:eastAsia"),
                formatting.heading2_font_name_east_asia or "黑体",
            )
            h2.paragraph_format.space_before = Pt(13)
            h2.paragraph_format.space_after = Pt(13)
            h2.paragraph_format.line_spacing = 1.73
        except Exception:
            pass
        try:
            h3 = doc.styles["Heading 3"]
            h3.font.name = formatting.heading3_font_name or formatting.font_name or "黑体"
            h3.font.size = Pt(formatting.heading3_size_pt or 16)
            h3.font.bold = True
            h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r_pr = h3._element.get_or_add_rPr()  # type: ignore[attr-defined]
            r_pr.get_or_add_rFonts().set(
                qn("w:eastAsia"),
                formatting.heading3_font_name_east_asia or "黑体",
            )
            h3.paragraph_format.space_before = Pt(13)
            h3.paragraph_format.space_after = Pt(13)
            h3.paragraph_format.line_spacing = 1.73
        except Exception:
            pass
        for i in range(1, 4):
            try:
                toc_style_name = self._resolve_toc_style_name(doc, i)
                if not toc_style_name:
                    continue
                toc_style = doc.styles[toc_style_name]
                toc_style.font.name = formatting.font_name or "瀹嬩綋"
                toc_style.font.size = Pt(formatting.font_size_pt or 12)
                toc_style.font.bold = False
                toc_para = toc_style.paragraph_format
                toc_para.space_before = Pt(0)
                toc_para.space_after = Pt(0)
                toc_para.line_spacing = 1.5
                r_pr = toc_style._element.get_or_add_rPr()  # type: ignore[attr-defined]
                r_pr.get_or_add_rFonts().set(qn("w:eastAsia"), formatting.font_name_east_asia or "瀹嬩綋")
            except Exception:
                pass

    def _add_cover(self, doc: Document, title: str) -> None:
        p = doc.add_paragraph()
        try:
            p.style = "Title"
        except Exception:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        try:
            run.font.name = "黑体"
            r_pr = run._element.get_or_add_rPr()  # type: ignore[attr-defined]
            r_pr.get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        except Exception:
            pass
        doc.add_paragraph("")
        doc.add_paragraph("")
        # The next section will start a new page; avoid double breaks here.

    def _collect_toc_entries(self, blocks: list[DocBlock], *, levels: int) -> list[tuple[int, str]]:
        max_level = max(1, min(4, int(levels or 3)))
        out: list[tuple[int, str]] = []
        seen: set[tuple[int, str]] = set()
        for block in blocks:
            if block.type != "heading":
                continue
            level = max(1, min(6, int(block.level or 1)))
            if level > max_level:
                continue
            title = _sanitize_heading_text((block.text or "").strip())
            if not title:
                continue
            key = (level, title)
            if key in seen:
                continue
            seen.add(key)
            out.append((level, title))
            if len(out) >= 120:
                break
        return out

    def _build_toc_anchor_entries(self, entries: list[tuple[int, str]]) -> list[TocAnchorEntry]:
        def _bookmark_name(title: str, idx: int) -> str:
            slug = re.sub(r"[^0-9A-Za-z_]+", "_", str(title or "").strip())
            slug = re.sub(r"_+", "_", slug).strip("_")
            if not slug:
                slug = f"h{idx:03d}"
            if not re.match(r"^[A-Za-z_]", slug):
                slug = f"h_{slug}"
            # Word bookmark names are safest when kept short and ASCII.
            return f"toc_{idx:03d}_{slug[:24]}"

        out: list[TocAnchorEntry] = []
        for idx, (level, title) in enumerate(entries, start=1):
            out.append(TocAnchorEntry(level=int(level or 1), title=str(title or ""), anchor=_bookmark_name(title, idx)))
        return out

    def _display_width(self, text: str) -> int:
        total = 0
        for ch in str(text or ""):
            if unicodedata.east_asian_width(ch) in {"W", "F"}:
                total += 2
            else:
                total += 1
        return total

    def _estimate_toc_pages(self, entries: list[tuple[int, str]], blocks: list[DocBlock], *, start_page: int = 1) -> list[int]:
        if not entries:
            return []
        def _norm(value: str) -> str:
            return re.sub(r"\s+", "", str(value or "").strip().lower())

        def _block_cost(block: DocBlock) -> int:
            t = str(block.type or "")
            if t == "heading":
                return 70
            if t == "paragraph":
                return max(40, len(str(block.text or "")))
            if t == "list":
                items = getattr(block, "items", None) or []
                try:
                    return max(80, sum(len(str(x or "")) for x in list(items)))
                except Exception:
                    return 120
            if t == "table":
                return 700
            if t == "figure":
                return 600
            return 120

        key_to_indices: dict[tuple[int, str], list[int]] = {}
        for idx, (lvl, title) in enumerate(entries):
            key = (int(lvl or 1), _norm(title))
            key_to_indices.setdefault(key, []).append(idx)
        pages = [max(1, int(start_page)) for _ in entries]
        page = max(1, int(start_page))
        budget = 0
        page_capacity = 1800
        for block in blocks:
            if block.type == "heading":
                level = max(1, min(6, int(block.level or 1)))
                title = _sanitize_heading_text((block.text or "").strip())
                key = (level, _norm(title))
                queue = key_to_indices.get(key) or []
                if queue:
                    idx = queue.pop(0)
                    pages[idx] = page
            budget += _block_cost(block)
            while budget >= page_capacity:
                page += 1
                budget -= page_capacity
        return pages

    def _render_toc_preview_text(self, entries: list[tuple[int, str]], blocks: list[DocBlock]) -> str:
        if not entries:
            return ""
        pages = self._estimate_toc_pages(entries, blocks, start_page=1)
        target_width = 72
        lines: list[str] = []
        for idx, (level, title) in enumerate(entries):
            depth = max(1, min(4, int(level or 1)))
            indent = "  " * max(0, depth - 1)
            text = f"{indent}{title}".rstrip()
            page_num = max(1, int(pages[idx] if idx < len(pages) else 1))
            fill = max(6, target_width - self._display_width(text) - len(str(page_num)))
            lines.append(f"{text}{'.' * fill}{page_num}")
        return "\n".join(lines).strip()

    def _add_toc(
        self,
        doc: Document,
        *,
        levels: int,
        entries: list[TocAnchorEntry] | None = None,
        blocks: list[DocBlock] | None = None,
    ) -> None:
        h = doc.add_paragraph()
        try:
            h.style = "TOC Heading"
        except Exception:
            pass
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_toc_heading_style(h, "\u76ee \u5f55")
        toc_entries = list(entries or [])
        plain_entries = [(int(x.level or 1), str(x.title or "")) for x in toc_entries]
        if _toc_clickable_links_enabled():
            if not toc_entries:
                p = doc.add_paragraph("\u65e0\u76ee\u5f55\u9879")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                return
            pages = self._estimate_toc_pages(plain_entries, list(blocks or []), start_page=1)
            for idx, entry in enumerate(toc_entries):
                level = max(1, min(4, int(entry.level or 1)))
                title = str(entry.title or "")
                anchor = str(entry.anchor or "").strip()
                page_num = max(1, int(pages[idx] if idx < len(pages) else 1))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                try:
                    toc_style_name = self._resolve_toc_style_name(doc, level)
                    if toc_style_name:
                        p.style = toc_style_name
                except Exception:
                    pass
                try:
                    p.paragraph_format.left_indent = Cm(max(0.0, float(level - 1) * 0.75))
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                except Exception:
                    pass
                try:
                    sec = doc.sections[-1]
                    right_tab = sec.page_width - sec.left_margin - sec.right_margin
                    tab_stops = p.paragraph_format.tab_stops
                    try:
                        tab_stops.clear_all()
                    except Exception:
                        pass
                    tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                except Exception:
                    pass
                if anchor:
                    _add_internal_hyperlink(p, title, anchor)
                else:
                    _add_inline_runs(p, title)
                p.add_run("\t")
                if anchor:
                    _add_internal_hyperlink(p, str(page_num), anchor)
                else:
                    p.add_run(str(page_num))
            return

        if _dynamic_toc_field_enabled():
            p = doc.add_paragraph()
            switches = [f'\\\\o \"1-{levels}\"']
            if _toc_hyperlink_enabled():
                switches.append("\\\\h")
            switches.extend(["\\\\z", "\\\\u"])
            instr = "TOC " + " ".join(switches)
            preview = self._render_toc_preview_text(plain_entries, list(blocks or []))
            _add_field_simple(p, instr, preview, lock=_toc_field_lock_enabled())
            return

        if not toc_entries:
            p = doc.add_paragraph("\u65e0\u76ee\u5f55\u9879")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            return

        for entry in toc_entries:
            level = int(entry.level or 1)
            title = str(entry.title or "")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p.paragraph_format.left_indent = Cm(max(0.0, float(level - 1) * 0.75))
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            _add_inline_runs(p, title)

    def _new_section(
        self,
        doc: Document,
        *,
        start_page_numbering: bool,
        numbering_format: str | None = None,
    ):
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        try:
            sec.header.is_linked_to_previous = False
            sec.footer.is_linked_to_previous = False
        except Exception:
            pass
        if start_page_numbering:
            _set_section_page_numbering(sec, start_at=1, numbering_format=numbering_format)
        return sec

    def _set_header(self, sec, title: str) -> None:
        hdr = sec.header
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        p.text = title
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.runs[0] if p.runs else p.add_run("")
            run.font.name = "宋体"
            run.font.size = Pt(9)
            r_pr = run._element.get_or_add_rPr()  # type: ignore[attr-defined]
            r_pr.get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass
        _add_bottom_border(p)

    def _set_footer_page_numbers(self, sec, footer_text: str = "", *, page_format: str | None = None) -> None:
        def _write_footer(footer) -> None:
            if footer is None:
                return
            for p in list(footer.paragraphs):
                try:
                    p._element.getparent().remove(p._element)  # type: ignore[attr-defined]
                except Exception:
                    pass
            p = footer.add_paragraph()
            _force_paragraph_center(p)
            try:
                p.paragraph_format.left_indent = Cm(0)
                p.paragraph_format.first_line_indent = Cm(0)
            except Exception:
                pass
            if footer_text:
                p.add_run(footer_text)
            if page_format:
                _add_field_simple(p, f"PAGE \\\\* {page_format}", "", lock=_page_field_lock_enabled())
            else:
                _add_field_simple(p, "PAGE", "", lock=_page_field_lock_enabled())

        _write_footer(getattr(sec, "footer", None))
        try:
            sec.different_first_page_header_footer = False
        except Exception:
            pass
        _strip_non_default_header_footer_refs(sec)

    def _emit_content(
        self,
        doc: Document,
        blocks: list[DocBlock],
        formatting: FormattingRequirements,
        *,
        number_headings: bool,
        apply_para_format: bool,
        heading_styles: dict[int, str] | None = None,
        toc_anchor_entries: list[TocAnchorEntry] | None = None,
    ) -> None:
        h1 = 0
        h2 = 0
        h3 = 0
        table_no = 0
        fig_no = 0
        current_section = ""
        ref_index = 0
        seen_h1 = False
        has_reference_heading = any(
            b.type == "heading"
            and _is_reference_title(
                _normalize_section_title(_sanitize_heading_text((b.text or "").strip()))
            )
            for b in blocks
        )
        inserted_reference_heading = False
        ref_buffer: list[tuple[int | None, str]] = []
        bookmark_id = 1

        def _norm_key(level: int, title: str) -> tuple[int, str]:
            return (
                max(1, min(6, int(level or 1))),
                re.sub(r"\s+", "", str(title or "").strip().lower()),
            )

        anchor_queue: dict[tuple[int, str], list[str]] = {}
        for item in list(toc_anchor_entries or []):
            key = _norm_key(int(item.level or 1), str(item.title or ""))
            anchor_queue.setdefault(key, []).append(str(item.anchor or ""))

        def _attach_anchor(paragraph: Paragraph | None, level: int, heading_title: str) -> None:
            nonlocal bookmark_id
            if paragraph is None:
                return
            key = _norm_key(level, heading_title)
            queue = anchor_queue.get(key) or []
            if not queue:
                return
            anchor = str(queue.pop(0) or "").strip()
            if not anchor:
                return
            _add_bookmark(paragraph, anchor, bookmark_id)
            bookmark_id += 1

        def flush_reference_buffer() -> None:
            nonlocal ref_buffer, ref_index
            if not ref_buffer:
                return
            indexed = [(idx, text) for idx, text in ref_buffer if idx is not None]
            free = [text for idx, text in ref_buffer if idx is None]
            indexed.sort(key=lambda x: x[0])
            ordered = [text for _, text in indexed] + free
            merged: list[str] = []
            for item in ordered:
                if merged and (len(item) <= 8 or re.fullmatch(r"[\d\W]+", item)):
                    merged[-1] = (merged[-1].rstrip("，,") + " " + item).strip()
                else:
                    merged.append(item)
            for item in merged:
                ref_index += 1
                t = f"[{ref_index}] {item}"
                p = doc.add_paragraph()
                if apply_para_format:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    try:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                    except Exception:
                        pass
                try:
                    p.paragraph_format.left_indent = Cm(FIRST_LINE_INDENT_CM)
                    p.paragraph_format.first_line_indent = Cm(-FIRST_LINE_INDENT_CM)
                except Exception:
                    pass
                _add_inline_runs(p, t)
            ref_buffer = []

        def handle_paragraph_text(raw_text: str) -> None:
            nonlocal current_section, inserted_reference_heading, has_reference_heading, ref_index, ref_buffer
            for chunk in _split_paragraph_chunks(raw_text):
                t = _sanitize_paragraph_text(chunk)
                if not t:
                    continue
                bullet = re.match(r"^\s*[\u2022\u00B7]\s+(.+)$", t)
                num_bullet = re.match(r"^\s*\d+[.\uFF0E\u3001\)]\s*(.+)$", t)
                is_ref_line = bool(re.match(r"^\s*\[\d+\]\s+", t))
                if is_ref_line and not _is_reference_title(current_section):
                    if not inserted_reference_heading:
                        ref_title = "\u53c2\u8003\u6587\u732e"
                        if number_headings:
                            _add_heading(doc, ref_title, level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
                        else:
                            ref_level = 2
                            style_name = heading_styles.get(ref_level) if heading_styles else None
                            _add_heading(doc, ref_title, level=ref_level, style_name=style_name)
                        current_section = ref_title
                        ref_index = 0
                        inserted_reference_heading = True
                        has_reference_heading = True
                if _is_reference_title(current_section) and _is_reference_noise(t):
                    continue
                if _is_reference_title(current_section):
                    m = re.match(r"^\s*\[(\d+)\]\s*(.+)$", t)
                    if m:
                        ref_buffer.append((int(m.group(1)), m.group(2).strip()))
                    else:
                        ref_buffer.append((None, t.strip()))
                    continue
                if num_bullet:
                    item = num_bullet.group(1).strip()
                    try:
                        p = doc.add_paragraph(style="List Number")
                        _add_inline_runs(p, item)
                    except KeyError:
                        p = doc.add_paragraph()
                        _add_inline_runs(p, f"{item}")
                    if apply_para_format:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        try:
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                        except Exception:
                            pass
                    continue
                if bullet:
                    item = bullet.group(1).strip()
                    try:
                        p = doc.add_paragraph(style="List Bullet")
                        _add_inline_runs(p, item)
                    except KeyError:
                        p = doc.add_paragraph()
                        _add_inline_runs(p, f"\u2022 {item}")
                    if apply_para_format:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        try:
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                        except Exception:
                            pass
                    continue
                p = doc.add_paragraph()
                if apply_para_format:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    try:
                        p.paragraph_format.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                    except Exception:
                        pass
                _add_inline_runs(p, t)

        for b in blocks:
            if b.type == "heading":
                lvl = int(b.level or 1)
                raw_text = _sanitize_heading_text((b.text or "").strip())
                text, tail = _split_heading_tail(raw_text)
                if not text:
                    continue
                if _is_reference_title(current_section) and ref_buffer:
                    flush_reference_buffer()
                if number_headings:
                    clean = _strip_chapter_prefix(text)
                    if _is_reference_title(_normalize_section_title(clean)):
                        if seen_h1:
                            doc.add_page_break()
                        seen_h1 = True
                        heading_paragraph = _add_heading(doc, clean, level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
                        _attach_anchor(heading_paragraph, lvl, text)
                        current_section = _normalize_section_title(clean)
                        ref_index = 0
                        continue
                    if lvl <= 1:
                        if seen_h1:
                            doc.add_page_break()
                        seen_h1 = True
                        h1 += 1
                        h2 = 0
                        h3 = 0
                        t = f"\u7b2c{h1}\u7ae0 {clean}"
                        heading_paragraph = _add_heading(doc, t, level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
                        _attach_anchor(heading_paragraph, lvl, text)
                    elif lvl == 2:
                        if h1 == 0:
                            h1 = 1
                        h2 += 1
                        h3 = 0
                        t = f"{h1}.{h2} {clean}"
                        heading_paragraph = _add_heading(
                            doc,
                            t,
                            level=2,
                            align=WD_ALIGN_PARAGRAPH.CENTER if _center_chapter_headings_enabled() else WD_ALIGN_PARAGRAPH.LEFT,
                        )
                        _attach_anchor(heading_paragraph, lvl, text)
                    else:
                        if h1 == 0:
                            h1 = 1
                        if h2 == 0:
                            h2 = 1
                        h3 += 1
                        t = f"{h1}.{h2}.{h3} {clean}"
                        heading_paragraph = _add_heading(doc, t, level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
                        _attach_anchor(heading_paragraph, lvl, text)
                else:
                    level = max(1, min(3, lvl))
                    if level == 1:
                        if seen_h1:
                            doc.add_page_break()
                        seen_h1 = True
                    if heading_styles:
                        style_name = heading_styles.get(level)
                        heading_paragraph = _add_heading(doc, text, level=level, style_name=style_name)
                    else:
                        heading_paragraph = _add_heading(doc, text, level=level)
                    _attach_anchor(heading_paragraph, level, text)
                current_section = _normalize_section_title(text)
                if _is_reference_title(current_section):
                    ref_index = 0
                if tail:
                    handle_paragraph_text(tail)
                continue

            if b.type == "paragraph":
                handle_paragraph_text(b.text or "")
                continue

            if b.type == "table":
                table_no += 1
                t = b.table or {}
                caption = str(t.get("caption") or "").strip() or f"\u8868{table_no}"
                cols = t.get("columns") if isinstance(t, dict) else None
                rows = t.get("rows") if isinstance(t, dict) else None
                columns = [str(c) for c in cols] if isinstance(cols, list) and cols else ["列1", "列2"]
                body = rows if isinstance(rows, list) and rows else [["—", "—"]]

                cap_p = doc.add_paragraph(f"表{table_no}  {caption}")
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                table = doc.add_table(rows=1, cols=len(columns))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(columns):
                    hdr_cells[i].text = str(col)
                for r in body:
                    rr = r if isinstance(r, list) else [str(r)]
                    row_cells = table.add_row().cells
                    for i in range(len(columns)):
                        row_cells[i].text = str(rr[i] if i < len(rr) else "")
                continue

            if b.type == "figure":
                fig_no += 1
                f = b.figure or {}
                caption = str(f.get("caption") or "").strip() or f"\u56fe{fig_no}"
                kind = str(f.get("type") or "").strip() or "figure"
                cap_p = doc.add_paragraph(f"图{fig_no}  {caption}")
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cairosvg:
                    try:
                        svg, _cap = render_figure_svg(f)
                        png = cairosvg.svg2png(bytestring=svg.encode("utf-8"))
                        pic_p = doc.add_paragraph()
                        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = pic_p.add_run()
                        try:
                            sec = doc.sections[0]
                            width = sec.page_width - sec.left_margin - sec.right_margin
                            run.add_picture(BytesIO(png), width=width)
                        except Exception:
                            run.add_picture(BytesIO(png), width=Cm(15))
                    except Exception:
                        ph = doc.add_paragraph(f"{caption}（图像生成失败）")
                        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    ph = doc.add_paragraph(f"{caption}（图像生成失败）")
                    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
        if _is_reference_title(current_section) and ref_buffer:
            flush_reference_buffer()


