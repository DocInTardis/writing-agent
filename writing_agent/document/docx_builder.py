"""Docx Builder module.

This module belongs to `writing_agent.document` in the writing-agent codebase.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

from writing_agent.agents.citations import CitationAgent
from writing_agent.document.docx_hyperlink import add_hyperlink
from writing_agent.models import Citation, DraftDocument, FormattingRequirements


class DocxBuilder:
    CITE_PATTERN = re.compile(r"\[@(?P<key>[a-zA-Z0-9_-]+)\]")

    def __init__(self) -> None:
        self._citer = CitationAgent()

    def build(self, draft: DraftDocument, citations: dict[str, Citation], formatting: FormattingRequirements) -> bytes:
        doc = Document()
        self._apply_styles(doc, formatting)

        doc.add_heading(draft.title, level=0)

        cite_order: list[str] = []
        cite_numbers: dict[str, int] = {}

        def assign(key: str) -> int:
            if key not in cite_numbers:
                cite_order.append(key)
                cite_numbers[key] = len(cite_order)
            return cite_numbers[key]

        for section in draft.sections:
            doc.add_heading(section.title, level=section.level)
            for para in section.paragraphs:
                text = para.text or ""
                rendered = self.CITE_PATTERN.sub(lambda m: f"[{assign(m.group('key'))}]", text)
                doc.add_paragraph(rendered)

        doc.add_heading("参考文献", level=1)
        for key in cite_order:
            n = cite_numbers[key]
            c = citations.get(key)
            if c is None:
                doc.add_paragraph(f"[{n}] {key}（未找到引用详情）")
                continue

            ref = self._citer.format_reference(c, formatting.citation_style)
            if c.url and ref.endswith(c.url):
                ref_no_url = ref[: -len(c.url)].rstrip()
            else:
                ref_no_url = ref

            p = doc.add_paragraph(f"[{n}] {ref_no_url}")
            if c.url:
                p.add_run(" ")
                add_hyperlink(p, c.url, c.url)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _apply_styles(self, doc: Document, formatting: FormattingRequirements) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = formatting.font_name
        normal.font.size = Pt(formatting.font_size_pt)
        line_spacing = formatting.line_spacing
        if line_spacing >= 4:
            normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            normal.paragraph_format.line_spacing = Pt(line_spacing)
        else:
            normal.paragraph_format.line_spacing = line_spacing
        try:
            r_pr = normal._element.get_or_add_rPr()  # type: ignore[attr-defined]
            r_fonts = r_pr.get_or_add_rFonts()
            r_fonts.set(qn("w:eastAsia"), formatting.font_name_east_asia)
        except Exception:
            pass
