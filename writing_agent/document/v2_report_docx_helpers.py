"""V2 Report Docx Helpers module.

This module belongs to `writing_agent.document` in the writing-agent codebase.
"""

from __future__ import annotations

import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from writing_agent.document import v2_report_docx_content_helpers as content_helpers
from writing_agent.v2.doc_format import DocBlock

FIRST_LINE_INDENT_CM = 0.85


def _docx_aggressive_split_enabled() -> bool:
    raw = str(os.environ.get("WRITING_AGENT_DOCX_AGGRESSIVE_SPLIT", "0")).strip().lower()
    return raw in {"1", "true", "yes", "on"}


def _resolved_field_default_text(instr: str, default_text: str) -> str:
    text = str(default_text or "")
    if text.strip():
        return text
    instr_upper = str(instr or "").upper().strip()
    # PAGE fields with empty display text are commonly flagged as dirty by Word.
    if instr_upper.startswith("PAGE"):
        if "ROMAN" in instr_upper:
            return "I"
        return "1"
    return text


def _add_field_simple(
    paragraph,
    instr: str,
    default_text: str,
    *,
    lock: bool = False,
    dirty: bool = False,
) -> None:
    # Use complex field to preserve switches like \\* ROMAN.
    display_text = _resolved_field_default_text(instr, default_text)

    r_begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    if lock:
        fld_begin.set(qn("w:fldLock"), "true")
    if dirty:
        fld_begin.set(qn("w:dirty"), "true")
    r_begin.append(fld_begin)

    r_instr = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    r_instr.append(instr_text)

    r_sep = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fld_sep)

    r_text = OxmlElement("w:r")
    segments = str(display_text or "").split("\n")
    if not segments:
        segments = [""]
    for idx, seg in enumerate(segments):
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = seg
        r_text.append(t)
        if idx < (len(segments) - 1):
            br = OxmlElement("w:br")
            r_text.append(br)

    r_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end.append(fld_end)

    p = paragraph._p  # type: ignore[attr-defined]
    p.append(r_begin)
    p.append(r_instr)
    p.append(r_sep)
    p.append(r_text)
    p.append(r_end)


def _add_heading(
    doc: Document,
    text: str,
    *,
    level: int,
    style_name: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
):
    p = doc.add_paragraph()
    # Prefer template style, but fall back to built-in Heading styles so
    # Word TOC fields can always pick up heading hierarchy.
    target_style = style_name or f"Heading {level}"
    try:
        p.style = target_style
    except Exception:
        try:
            p.style = f"Heading {max(1, min(9, int(level or 1)))}"
        except Exception:
            pass
    if align is not None:
        p.alignment = align
    _add_inline_runs(p, text)
    return p


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    anchor = str(name or "").strip()
    if not anchor:
        return
    bid = max(1, int(bookmark_id or 1))
    try:
        p = paragraph._p  # type: ignore[attr-defined]
        b_start = OxmlElement("w:bookmarkStart")
        b_start.set(qn("w:id"), str(bid))
        b_start.set(qn("w:name"), anchor)
        b_end = OxmlElement("w:bookmarkEnd")
        b_end.set(qn("w:id"), str(bid))
        # Keep paragraph properties first (<w:pPr>) to avoid XML order issues.
        insert_pos = 0
        if len(p) > 0 and getattr(p[0], "tag", "") == qn("w:pPr"):
            insert_pos = 1
        p.insert(insert_pos, b_start)
        p.append(b_end)
    except Exception:
        pass


def _add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    try:
        p = paragraph._p  # type: ignore[attr-defined]
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("w:anchor"), str(anchor or ""))
        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "Hyperlink")
        r_pr.append(r_style)
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_fonts.set(qn("w:eastAsia"), "宋体")
        r_pr.append(r_fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        r_pr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), "24")
        r_pr.append(sz_cs)
        # Keep TOC appearance close to thesis templates: black + no underline.
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        r_pr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "none")
        r_pr.append(u)
        run.append(r_pr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = str(text or "")
        run.append(t)
        hl.append(run)
        p.append(hl)
    except Exception:
        # Fallback to plain text if hyperlink injection fails.
        paragraph.add_run(str(text or ""))


def _add_inline_runs(paragraph, text: str) -> None:
    for run_text, bold, italic, underline in _split_inline_runs(text):
        if not run_text:
            continue
        _add_runs_with_citations(paragraph, run_text, bold, italic, underline)


def _add_runs_with_citations(paragraph, text: str, bold: bool, italic: bool, underline: bool) -> None:
    # 优化: 先处理 [@citekey] 格式的引用标记
    cite_pattern = re.compile(r'\[@([a-zA-Z0-9_-]+)\]')
    # 分割文本和引用标记
    parts = cite_pattern.split(text)
    
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 0:
            # 普通文本，继续处理 [数字] 格式
            _add_runs_with_number_citations(paragraph, part, bold, italic, underline)
        else:
            # [@citekey] 的key部分，暂时保持原样（需要在生成时统一映射）
            run = paragraph.add_run(f'[@{part}]')
            run.bold = bool(bold)
            run.font.italic = bool(italic)
            run.font.underline = bool(underline)
            run.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色标识


def _add_runs_with_number_citations(paragraph, text: str, bold: bool, italic: bool, underline: bool) -> None:
    """处理 [数字] 格式的引用"""
    parts = re.split(r"(\\[\\d+\\])", text)
    for part in parts:
        if not part:
            continue
        m = re.fullmatch(r"\\[(\\d+)\\]", part)
        if m:
            run = paragraph.add_run(part)
            run.bold = bool(bold)
            run.italic = bool(italic)
            run.underline = bool(underline)
            try:
                run.font.superscript = True
                run.font.size = Pt(8)
            except Exception:
                pass
            continue
        run = paragraph.add_run(part)
        run.bold = bool(bold)
        run.italic = bool(italic)
        run.underline = bool(underline)


def _split_inline_runs(text: str) -> list[tuple[str, bool, bool, bool]]:
    src = text or ""
    bold = False
    italic = False
    underline = False
    buf: list[str] = []
    runs: list[tuple[str, bool, bool, bool]] = []

    def flush() -> None:
        if not buf:
            return
        runs.append(("".join(buf), bold, italic, underline))
        buf.clear()

    def has_closing(idx: int, marker: str) -> bool:
        return src.find(marker, idx + len(marker)) != -1

    i = 0
    while i < len(src):
        if src.startswith("**", i) and (bold or has_closing(i, "**")):
            flush()
            bold = not bold
            i += 2
            continue
        if src.startswith("__", i) and (underline or has_closing(i, "__")):
            flush()
            underline = not underline
            i += 2
            continue
        if src.startswith("*", i) and (italic or has_closing(i, "*")):
            flush()
            italic = not italic
            i += 1
            continue
        buf.append(src[i])
        i += 1
    flush()
    return runs


def _strip_inline_markers(text: str) -> str:
    return re.sub(r"(\*\*|__|\*)", "", text or "")


def _normalize_para(text: str) -> str:
    s = (text or "").replace("\r", "").strip()
    if not s:
        return ""
    s = re.sub(r"\s+\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s


def _split_heading_tail(text: str) -> tuple[str, str]:
    s = _normalize_export_text(text, field="heading").strip()
    if not s:
        return "", ""
    if not _docx_aggressive_split_enabled():
        return s, ""
    head = s
    tail = ""
    common_heads = [
        "\u7eea\u8bba",
        "\u5f15\u8a00",
        "\u9700\u6c42\u5206\u6790",
        "\u6982\u8981\u8bbe\u8ba1",
        "\u8be6\u7ec6\u8bbe\u8ba1",
        "\u7cfb\u7edf\u8bbe\u8ba1",
        "\u7cfb\u7edf\u7684\u5b9e\u73b0",
        "\u7cfb\u7edf\u5b9e\u73b0",
        "\u7cfb\u7edf\u6d4b\u8bd5",
        "\u6d4b\u8bd5\u4e0e\u7ed3\u679c\u5206\u6790",
        "\u603b\u7ed3\u4e0e\u5c55\u671b",
        "\u7ed3\u8bba",
        "\u53c2\u8003\u6587\u732e",
        "\u6c34\u5e73\u4f30\u4f30",
    ]
    for title in common_heads:
        if s.startswith(title) and len(s) > len(title):
            head = title
            tail = s[len(title) :].lstrip("\uff0c\u3002\uFF1A:;\uff1b\u2014- ").strip()
            return head, tail
    # 优化: 修复标题分割正则,避免标题与内容粘连
    # 匹配模式：标题(2-12字) + 连接词 + 至少有动词/实词的句子
    intro = re.match(r"^(.{2,12})(在|本|针对|随着|通过|为了|由于|基于|围绕|结合|面向)(.{4,})", s)
    if intro:
        head_candidate = intro.group(1).strip()
        connector = intro.group(2)
        rest = intro.group(3).strip()
        
        # 只有在后续内容明显是句子内容时才分割
        # 判断标准: 包含动词(完成|实现|进行|分析|设计|建立|提出|探讨|研究|开发|构建)
        if re.search(r"完成|实现|进行|分析|设计|建立|提出|探讨|研究|开发|构建|包括|涵盖|涉及", rest):
            head = head_candidate
            tail = connector + rest
        else:
            # 否则保持完整不分割
            head = s
            tail = ""
    else:
        m = re.search(r"[。！？；]", s)
        if m:
            head = s[: m.start()].strip()
            tail = s[m.start() + 1 :].strip()
    if len(head) > 24:
        m2 = re.search(r"[，、：]", head)
        if m2 and m2.start() >= 6:
            tail = (head[m2.start() + 1 :].strip() + " " + tail).strip()
            head = head[: m2.start()].strip()
        elif len(head) > 24:
            tail = (head[24:].strip() + " " + tail).strip()
            head = head[:24].strip()
    if not head:
        return s, ""
    return head, tail


def _split_paragraph_chunks(text: str) -> list[str]:
    raw = _normalize_para(text)
    if not raw:
        return []
    lines = [s.strip() for s in re.split(r"\n+", raw) if s.strip()]
    if not _docx_aggressive_split_enabled():
        return lines
    out: list[str] = []
    for line in lines:
        if re.match(r"^\s*\[\d+\]\s+", line):
            out.append(line)
            continue
        if re.search(r"\s\d+[.．]\s*", line):
            parts = [p.strip() for p in re.split(r"\s+(?=\d+[.．]\s*)", line) if p.strip()]
            if len(parts) > 1:
                out.extend(parts)
                continue
        if len(re.findall(r"\d+[.．]\s*", line)) >= 2:
            parts = [p.strip() for p in re.split(r"(?=\d+[.．]\s*)", line) if p.strip()]
            out.extend(parts)
            continue
        if len(re.findall(r"\d+、", line)) >= 2:
            parts = [p.strip() for p in re.split(r"(?=\d+、)", line) if p.strip()]
            out.extend(parts)
            continue
        if len(re.findall(r"[（(]\d+[)）]", line)) >= 2:
            parts = [p.strip() for p in re.split(r"(?=[（(]\d+[)）])", line) if p.strip()]
            out.extend(parts)
            continue
        if re.search(r"[：:]\s*[—–-]\s*", line):
            prefix, rest = re.split(r"[：:]\s*[—–-]\s*", line, maxsplit=1)
            prefix = prefix.strip()
            if prefix:
                out.append(prefix + "：")
            items = [s.strip(" -—–") for s in re.split(r"\s*[—–-]\s+", rest) if s.strip()]
            for item in items:
                out.append(f"\u2022 {item}")
            continue
        dash_matches = re.findall(r"[—–-]", line)
        if len(dash_matches) >= 2 and not line.startswith(("-", "—", "–")):
            parts = [p.strip() for p in re.split(r"\s*[—–-]\s*", line) if p.strip()]
            if len(parts) > 1:
                out.append(parts[0])
                for item in parts[1:]:
                    out.append(f"\u2022 {item}")
                continue
        out.append(line)
    return out


_MOJIBAKE_PATTERNS = [
    re.compile(r"(?:Ã.|Â.|â€|â€“|â€”|é«|æ；|ç；)", re.IGNORECASE),
    re.compile(r"�"),
]

_HEADING_NORMALIZATION_TARGETS = (
    "摘要",
    "关键词",
    "参考文献",
    "引言",
    "绪论",
    "结论",
    "标题",
)


def _mojibake_hit_count(text: str) -> int:
    token = str(text or "")
    if not token:
        return 0
    return sum(len(pattern.findall(token)) for pattern in _MOJIBAKE_PATTERNS)



def _normalize_export_text(text: str, *, field: str = "") -> str:
    s = str(text or "")
    if not s:
        return ""
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s)
    s = s.replace("\uFFFE", "").replace("\uFFFF", "")
    compact = s.strip()
    original_hits = _mojibake_hit_count(compact)

    if field in {"heading", "section_title", "title"}:
        for encoding in ("gb18030", "latin-1", "cp1252"):
            try:
                candidate = compact.encode(encoding).decode("utf-8")
            except Exception:
                continue
            candidate = candidate.strip()
            if any(candidate == marker or candidate.startswith(marker) for marker in _HEADING_NORMALIZATION_TARGETS):
                compact = candidate
                break

    if _mojibake_hit_count(compact) > 0:
        best = compact
        best_hits = _mojibake_hit_count(compact)
        for encoding in ("latin-1", "cp1252", "gb18030"):
            try:
                candidate = compact.encode(encoding).decode("utf-8")
            except Exception:
                continue
            candidate = candidate.strip()
            hits = _mojibake_hit_count(candidate)
            if hits < best_hits:
                best = candidate
                best_hits = hits
        compact = best
    if original_hits > 0 and _mojibake_hit_count(compact) > original_hits:
        return s.strip()
    return compact.strip()



def _looks_like_mojibake(text: str) -> bool:
    token = str(text or "").strip()
    if not token:
        return False
    return _mojibake_hit_count(token) > 0



def _collect_export_mojibake_fragments(title: str, blocks: list[DocBlock], *, max_hits: int = 8) -> list[str]:
    hits: list[str] = []
    title_norm = _normalize_export_text(title, field="title")
    if _looks_like_mojibake(title_norm):
        hits.append(title_norm[:120])
    for block in blocks or []:
        block_type = str(getattr(block, "type", "") or "")
        samples: list[str] = []
        if block_type in {"heading", "paragraph"}:
            samples.append(_normalize_export_text(str(getattr(block, "text", "") or ""), field=block_type))
        caption = str(getattr(block, "caption", "") or "").strip()
        if caption:
            samples.append(_normalize_export_text(caption, field="caption"))
        for sample in samples:
            if not sample:
                continue
            if _looks_like_mojibake(sample):
                hits.append(sample[:120])
                if len(hits) >= max_hits:
                    return hits
    return hits


def _sanitize_heading_text(text: str) -> str:
    s = _normalize_export_text(text, field="heading").strip()
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s)
    s = s.replace("\uFFFE", "").replace("\uFFFF", "")
    s = _strip_inline_markers(s)
    s = re.sub(r"^\s*#{1,}\s+", "", s)
    s = s.replace("`", "")
    s = re.sub(r"^\s*[*\-\u2022]\s+", "", s)
    if re.fullmatch(r"\?+", s):
        return ""
    return s.strip()


def _sanitize_paragraph_text(text: str) -> str:
    s = _normalize_export_text(text, field="paragraph").replace("`", "")
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s)
    s = s.replace("\uFFFE", "").replace("\uFFFF", "")
    s = re.sub(r"(?m)^\s*#{2,}\s+", "", s)
    s = re.sub(r"(?m)^\s*[*\-\u2013]\s+", "", s)
    strip_filler = os.environ.get("WRITING_AGENT_STRIP_FILLER", "").strip().lower() in {"1", "true", "yes", "on"}
    if strip_filler:
        filler = [
            "\u4e3a\u4fdd\u8bc1\u8bba\u8ff0\u53ef\u590d\u6838",
            "\u5728\u65b9\u6cd5\u4e0e\u7ed3\u8bba\u4e4b\u95f4",
            "\u672c\u8282\u5f3a\u8c03\u5b9e\u65bd\u5c42\u9762",
            "\u56f4\u7ed5\u6838\u5fc3\u6982\u5ff5\u3001\u7814\u7a76\u8303\u56f4",
        ]
        if any(p in s for p in filler):
            return ""
    return s.strip()


def _strip_chapter_prefix(text: str) -> str:
    s = (text or "").strip()
    s = re.sub(r"^\u7b2c\s*\d+\s*\u7ae0\s*", "", s)
    s = re.sub(r"^\d+(?:\.\d+)*\s*", "", s)
    return s.strip()


def _remove_disallowed_sections(blocks: list[DocBlock]) -> list[DocBlock]:
    # Keep abstract/keywords in export; remove only manual TOC/ack sections.
    disallowed = {
        "\u76ee\u5f55",
        "\u76ee\u6b21",
        "table of contents",
        "contents",
        "\u81f4\u8c22",
        "\u9e23\u8c22",
    }
    out: list[DocBlock] = []
    skip_level: int | None = None
    for b in blocks:
        if b.type == "heading":
            level = int(b.level or 1)
            title = _normalize_section_title(_sanitize_heading_text((b.text or "").strip()))
            if title in disallowed:
                skip_level = level
                continue
            if skip_level is not None and level <= skip_level:
                skip_level = None
        if skip_level is not None:
            continue
        out.append(b)
    return out


def _promote_headings_if_no_h1(blocks: list[DocBlock]) -> list[DocBlock]:
    if any(b.type == "heading" and int(b.level or 0) == 1 for b in blocks):
        return blocks
    out: list[DocBlock] = []
    for b in blocks:
        if b.type == "heading":
            lvl = int(b.level or 1)
            out.append(DocBlock(type="heading", level=max(1, lvl - 1), text=b.text))
        else:
            out.append(b)
    return out


def _normalize_section_title(text: str) -> str:
    s = _normalize_export_text(text, field="section_title").strip()
    s = re.sub(r"^\d+(?:\.\d+)*\s*", "", s)
    s = re.sub(r"^[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u3001\s*", "", s)
    s = re.sub(r"^\u7b2c[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+[\u7ae0\u8282]\s*", "", s)
    return s.strip()


# Stable Unicode-safe predicates for reference title and noise filtering.
def _is_reference_title(title: str) -> bool:
    t = (title or "").strip().lower()
    if not t:
        return False
    return any(
        key in t
        for key in (
            "\u53c2\u8003\u6587\u732e",
            "\u53c2\u8003\u8d44\u6599",
            "references",
            "reference",
            "bibliography",
        )
    )


def _is_reference_noise(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    phrases = [
        "\u793a\u4f8b",
        "\u865a\u6784",
        "\u8bf7\u6ce8\u610f",
        "\u7531\u4e8e\u60a8\u6ca1\u6709\u63d0\u4f9b",
        "\u7531\u4e8e\u672a\u63d0\u4f9b",
        "\u6ca1\u6709\u63d0\u4f9b",
        "\u6211\u5c06\u4e3a\u60a8",
        "\u4ee5\u4e0b\u4e3a",
        "\u4ec5\u4f9b\u53c2\u8003",
        "example",
        "placeholder",
        "for demonstration",
    ]
    return any(p in t for p in phrases)


def _pick_margin(value: float | None, fallback: float) -> float:
    try:
        return float(value) if value is not None else float(fallback)
    except Exception:
        return float(fallback)


def _set_section_page_numbering(sec, *, start_at: int, numbering_format: str | None = None) -> None:
    try:
        sect_pr = sec._sectPr  # type: ignore[attr-defined]
        pg = sect_pr.find(qn("w:pgNumType"))
        if pg is None:
            pg = OxmlElement("w:pgNumType")
            sect_pr.append(pg)
        pg.set(qn("w:start"), str(start_at))
        if numbering_format:
            pg.set(qn("w:fmt"), numbering_format)
    except Exception:
        pass


def _add_bottom_border(paragraph: Paragraph) -> None:
    try:
        p = paragraph._p  # type: ignore[attr-defined]
        p_pr = p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        p_bdr.append(bottom)
    except Exception:
        pass


def _force_paragraph_center(paragraph: Paragraph) -> None:
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    try:
        p = paragraph._p  # type: ignore[attr-defined]
        p_pr = p.get_or_add_pPr()
        jc = p_pr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            p_pr.append(jc)
        jc.set(qn("w:val"), "center")
    except Exception:
        pass


def _strip_non_default_header_footer_refs(sec) -> None:
    try:
        sect_pr = sec._sectPr  # type: ignore[attr-defined]
    except Exception:
        return
    seen_default: set[str] = set()
    for node in list(sect_pr):
        tag = str(getattr(node, "tag", ""))
        if tag.endswith("}headerReference"):
            kind = "header"
        elif tag.endswith("}footerReference"):
            kind = "footer"
        else:
            continue
        ref_type = str(node.get(qn("w:type")) or "default")
        if ref_type != "default":
            sect_pr.remove(node)
            continue
        if kind in seen_default:
            sect_pr.remove(node)
            continue
        seen_default.add(kind)


def _clear_header_footer(sec, *, include_variants: bool = False) -> None:
    containers = [
        getattr(sec, "header", None),
        getattr(sec, "footer", None),
    ]
    if include_variants:
        for attr in ("first_page_header", "first_page_footer", "even_page_header", "even_page_footer"):
            containers.append(getattr(sec, attr, None))
    for container in containers:
        if container is None:
            continue
        try:
            container.is_linked_to_previous = False
        except Exception:
            pass
        for p in list(container.paragraphs):
            p.text = ""
            for run in p.runs:
                run.text = ""
        # Remove extra paragraphs to avoid hidden fields from templates.
        for p in list(container.paragraphs)[1:]:
            try:
                p._element.getparent().remove(p._element)  # type: ignore[attr-defined]
            except Exception:
                pass
    if not include_variants:
        try:
            sec.different_first_page_header_footer = False
        except Exception:
            pass
    _strip_non_default_header_footer_refs(sec)


def _disable_first_page_numbering(sec) -> None:
    try:
        sec.different_first_page_header_footer = True
    except Exception:
        return
    try:
        ftr = sec.first_page_footer
        for p in list(ftr.paragraphs):
            p.text = ""
            for run in p.runs:
                run.text = ""
        for p in list(ftr.paragraphs)[1:]:
            try:
                p._element.getparent().remove(p._element)  # type: ignore[attr-defined]
            except Exception:
                pass
    except Exception:
        pass
    _strip_non_default_header_footer_refs(sec)


def _remove_section_page_numbering(sec) -> None:
    try:
        sect_pr = sec._sectPr  # type: ignore[attr-defined]
        pg = sect_pr.find(qn("w:pgNumType"))
        if pg is not None:
            sect_pr.remove(pg)
    except Exception:
        pass


def _strip_cover_section_numbering(doc: Document) -> None:
    try:
        root = doc.element  # type: ignore[attr-defined]
        for sect in root.findall(".//w:sectPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            pg = sect.find(qn("w:pgNumType"))
            if pg is None:
                continue
            fmt = pg.get(qn("w:fmt"))
            if not fmt:
                sect.remove(pg)
                # Also remove header/footer references so cover stays blank.
                for node in list(sect):
                    if node.tag.endswith("}headerReference") or node.tag.endswith("}footerReference"):
                        sect.remove(node)
    except Exception:
        pass


def _apply_toc_heading_style(paragraph: Paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.bold = True
    try:
        run.font.name = "黑体"
        run.font.size = Pt(16)
        r_pr = run._element.get_or_add_rPr()  # type: ignore[attr-defined]
        r_pr.get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    except Exception:
        pass
    try:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
    except Exception:
        pass


def _enable_update_fields(doc: Document) -> None:
    try:
        settings = doc.settings.element
        upd = settings.find(qn("w:updateFields"))
        if upd is None:
            upd = OxmlElement("w:updateFields")
            settings.append(upd)
        upd.set(qn("w:val"), "true")
    except Exception:
        pass


def _disable_update_fields(doc: Document) -> None:
    try:
        settings = doc.settings.element
        upd = settings.find(qn("w:updateFields"))
        if upd is not None:
            settings.remove(upd)
        # Explicitly hint Word not to auto-update fields when opening.
        no_auto = settings.find(qn("w:doNotAutoUpdateFields"))
        if no_auto is None:
            no_auto = OxmlElement("w:doNotAutoUpdateFields")
            settings.append(no_auto)
    except Exception:
        pass



_clear_doc_body = content_helpers._clear_doc_body
_truncate_template_body = content_helpers._truncate_template_body
_ensure_min_figures = content_helpers._ensure_min_figures
_ensure_min_tables = content_helpers._ensure_min_tables
_ensure_reference_section = content_helpers._ensure_reference_section
_find_template_body_anchor = content_helpers._find_template_body_anchor
_is_heading_para = content_helpers._is_heading_para
_looks_like_heading_text = content_helpers._looks_like_heading_text
_is_toc_style = content_helpers._is_toc_style
_looks_like_toc_entry = content_helpers._looks_like_toc_entry
_detect_template_heading_styles = content_helpers._detect_template_heading_styles
_save_doc = content_helpers._save_doc
_postprocess_toc_footer_numbers = content_helpers._postprocess_toc_footer_numbers

def _ensure_reference_citations(doc: Document) -> None:
    nums: set[int] = set()
    ref_nums: set[int] = set()
    last_body_paragraph = None
    ref_heading_idx = None
    for idx, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if text:
            last_body_paragraph = p
        if "参考文献" in text:
            ref_heading_idx = idx
            break
        for m in re.finditer(r"\\[(\\d+)\\]", text):
            try:
                nums.add(int(m.group(1)))
            except Exception:
                pass

    if ref_heading_idx is None:
        return

    for p in doc.paragraphs[ref_heading_idx + 1 :]:
        t = (p.text or "").strip()
        if not t:
            continue
        m = re.match(r"^\\[(\\d+)\\]", t)
        if m:
            try:
                ref_nums.add(int(m.group(1)))
            except Exception:
                pass
        elif ref_nums:
            # stop if references are done and content resumes
            break

    if not ref_nums:
        return

    max_ref = max(ref_nums)
    missing = [i for i in range(1, max_ref + 1) if i not in nums]
    if not missing:
        return
    if ref_heading_idx > 0:
        last_body_paragraph = doc.paragraphs[ref_heading_idx - 1]
    if last_body_paragraph is None:
        last_body_paragraph = doc.add_paragraph()
    for i in missing:
        run = last_body_paragraph.add_run(f"[{i}]")
        try:
            run.font.superscript = True
            run.font.size = Pt(8)
        except Exception:
            pass


__all__ = [name for name, value in globals().items() if name.startswith("_") and callable(value)]
