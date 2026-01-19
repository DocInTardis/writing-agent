from __future__ import annotations

import io
import re
from html.parser import HTMLParser

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

from writing_agent.models import FormattingRequirements


class HtmlDocxBuilder:
    def build(self, html: str, formatting: FormattingRequirements, resolve_image_path=None) -> bytes:
        doc = Document()
        self._apply_styles(doc, formatting)
        parser = _HtmlToDocxParser(doc, resolve_image_path=resolve_image_path)
        parser.feed(html or "")
        parser.close()
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _apply_styles(self, doc: Document, formatting: FormattingRequirements) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = formatting.font_name
        normal.font.size = Pt(formatting.font_size_pt)
        normal.paragraph_format.line_spacing = formatting.line_spacing
        try:
            r_pr = normal._element.get_or_add_rPr()  # type: ignore[attr-defined]
            r_fonts = r_pr.get_or_add_rFonts()
            r_fonts.set(qn("w:eastAsia"), formatting.font_name_east_asia)
        except Exception:
            pass


class _HtmlToDocxParser(HTMLParser):
    def __init__(self, doc: Document, resolve_image_path=None) -> None:
        super().__init__(convert_charrefs=True)
        self._doc = doc
        self._para = None
        self._bold = 0
        self._italic = 0
        self._underline = 0
        self._font_size_stack: list[Pt | None] = []
        self._font_family_stack: list[str | None] = []
        self._color_stack: list[RGBColor | None] = []
        self._highlight_stack: list[WD_COLOR_INDEX | None] = []
        self._list_stack: list[str] = []  # "ul"|"ol"
        self._heading_level: int | None = None
        self._resolve_image_path = resolve_image_path

        self._in_table = 0
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] | None = None
        self._current_cell: list[str] | None = None

    def handle_starttag(self, tag: str, attrs) -> None:
        t = tag.lower()
        attrs_dict = {str(k).lower(): ("" if v is None else str(v)) for k, v in (attrs or [])}
        if t in {"b", "strong"}:
            self._bold += 1
        elif t in {"i", "em"}:
            self._italic += 1
        elif t == "u":
            self._underline += 1
        elif t == "figure":
            self._para = self._doc.add_paragraph("")
        elif t == "figcaption":
            self._para = self._doc.add_paragraph("")
            self._italic += 1
        elif t == "span":
            fs = _extract_font_size(attrs_dict.get("style", ""))
            self._font_size_stack.append(fs)
            self._font_family_stack.append(_extract_font_family(attrs_dict.get("style", "")))
            self._color_stack.append(_extract_color(attrs_dict.get("style", "")))
            self._highlight_stack.append(_extract_highlight(attrs_dict.get("style", "")))
        elif t == "table":
            self._in_table += 1
            self._table_rows = []
            self._current_row = None
            self._current_cell = None
        elif t == "tr" and self._in_table:
            self._current_row = []
            self._table_rows.append(self._current_row)
        elif t in {"td", "th"} and self._in_table:
            if self._current_row is None:
                self._current_row = []
                self._table_rows.append(self._current_row)
            self._current_cell = []
            self._current_row.append("")
        elif t == "img":
            if self._in_table:
                return
            src = attrs_dict.get("src", "")
            path = None
            if callable(self._resolve_image_path):
                path = self._resolve_image_path(src)
            if path:
                if self._para is None:
                    self._para = self._doc.add_paragraph("")
                width = _extract_img_width(attrs_dict)
                try:
                    if width is not None:
                        self._para.add_run().add_picture(path, width=width)
                    else:
                        self._para.add_run().add_picture(path)
                except Exception:
                    self._para.add_run(f"[图片：{src}]")
        elif t in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(t[1])
            self._heading_level = max(1, min(level, 6))
            self._para = self._doc.add_heading("", level=self._heading_level)
            _apply_paragraph_style(self._para, attrs_dict.get("style", ""))
        elif t == "p":
            self._para = self._doc.add_paragraph("")
            _apply_paragraph_style(self._para, attrs_dict.get("style", ""))
        elif t == "br":
            if self._para is None:
                self._para = self._doc.add_paragraph("")
            self._para.add_run().add_break()
        elif t in {"ul", "ol"}:
            self._list_stack.append(t)
        elif t == "li":
            style = "List Bullet"
            if self._list_stack and self._list_stack[-1] == "ol":
                style = "List Number"
            self._para = self._doc.add_paragraph("", style=style)
            _apply_paragraph_style(self._para, attrs_dict.get("style", ""))

    def handle_endtag(self, tag: str) -> None:
        t = tag.lower()
        if t in {"b", "strong"}:
            self._bold = max(0, self._bold - 1)
        elif t in {"i", "em"}:
            self._italic = max(0, self._italic - 1)
        elif t == "u":
            self._underline = max(0, self._underline - 1)
        elif t == "figcaption":
            self._italic = max(0, self._italic - 1)
        elif t == "span":
            if self._font_size_stack:
                self._font_size_stack.pop()
            if self._font_family_stack:
                self._font_family_stack.pop()
            if self._color_stack:
                self._color_stack.pop()
            if self._highlight_stack:
                self._highlight_stack.pop()
        elif t in {"td", "th"} and self._in_table:
            self._current_cell = None
        elif t == "table" and self._in_table:
            self._in_table = max(0, self._in_table - 1)
            self._emit_table()
            self._table_rows = []
            self._current_row = None
            self._current_cell = None
        elif t in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self._heading_level = None
            self._para = None
        elif t in {"p", "li"}:
            self._para = None
        elif t in {"ul", "ol"}:
            if self._list_stack:
                self._list_stack.pop()

    def handle_data(self, data: str) -> None:
        txt = (data or "").replace("\r", "")
        if not txt.strip():
            return
        if self._in_table and self._current_row is not None:
            # Append to last cell in current row
            if not self._current_row:
                self._current_row.append("")
            idx = len(self._current_row) - 1
            self._current_row[idx] = (self._current_row[idx] or "") + txt
            return
        if self._para is None:
            self._para = self._doc.add_paragraph("")
        run = self._para.add_run(txt)
        run.bold = self._bold > 0
        run.italic = self._italic > 0
        run.underline = self._underline > 0
        fs = self._current_font_size()
        if fs is not None:
            run.font.size = fs
        ff = self._current_font_family()
        if ff:
            run.font.name = ff
        col = self._current_color()
        if col is not None:
            run.font.color.rgb = col
        hl = self._current_highlight()
        if hl is not None:
            run.font.highlight_color = hl

    def _current_font_size(self) -> Pt | None:
        for v in reversed(self._font_size_stack):
            if v is not None:
                return v
        return None

    def _current_font_family(self) -> str | None:
        for v in reversed(self._font_family_stack):
            if v:
                return v
        return None

    def _current_color(self) -> RGBColor | None:
        for v in reversed(self._color_stack):
            if v is not None:
                return v
        return None

    def _current_highlight(self) -> WD_COLOR_INDEX | None:
        for v in reversed(self._highlight_stack):
            if v is not None:
                return v
        return None

    def _emit_table(self) -> None:
        rows = [r for r in self._table_rows if r]
        if not rows:
            return
        cols = max(len(r) for r in rows)
        if cols <= 0:
            return
        table = self._doc.add_table(rows=len(rows), cols=cols)
        for i, r in enumerate(rows):
            for j in range(cols):
                val = (r[j] if j < len(r) else "").strip()
                table.cell(i, j).text = val


_FONT_SIZE_RE = re.compile(r"font-size\\s*:\\s*([0-9]+(?:\\.[0-9]+)?)\\s*(pt|px)\\s*", flags=re.IGNORECASE)
_FONT_FAMILY_RE = re.compile(r"font-family\\s*:\\s*([^;]+)", flags=re.IGNORECASE)
_COLOR_RE = re.compile(r"color\\s*:\\s*(#[0-9a-fA-F]{6})\\s*", flags=re.IGNORECASE)
_BG_RE = re.compile(r"background-color\\s*:\\s*(#[0-9a-fA-F]{6})\\s*", flags=re.IGNORECASE)
_LINE_HEIGHT_RE = re.compile(r"line-height\\s*:\\s*([0-9]+(?:\\.[0-9]+)?)\\s*", flags=re.IGNORECASE)
_TEXT_ALIGN_RE = re.compile(r"text-align\\s*:\\s*(left|center|right|justify)\\s*", flags=re.IGNORECASE)
_TEXT_INDENT_RE = re.compile(r"text-indent\\s*:\\s*([0-9]+(?:\\.[0-9]+)?)\\s*(pt|px|em)\\s*", flags=re.IGNORECASE)
_MARGIN_TOP_RE = re.compile(r"margin-top\\s*:\\s*([0-9]+(?:\\.[0-9]+)?)\\s*(pt|px)\\s*", flags=re.IGNORECASE)
_MARGIN_BOTTOM_RE = re.compile(r"margin-bottom\\s*:\\s*([0-9]+(?:\\.[0-9]+)?)\\s*(pt|px)\\s*", flags=re.IGNORECASE)


def _extract_font_size(style: str) -> Pt | None:
    if not style:
        return None
    m = _FONT_SIZE_RE.search(style)
    if not m:
        return None
    num = float(m.group(1))
    unit = m.group(2).lower()
    if unit == "px":
        num = num * 0.75
    if num < 6 or num > 72:
        return None
    return Pt(num)


def _extract_font_family(style: str) -> str | None:
    if not style:
        return None
    m = _FONT_FAMILY_RE.search(style)
    if not m:
        return None
    fam = m.group(1).strip().strip("\"'")
    if not fam:
        return None
    # Keep first family only
    fam = fam.split(",")[0].strip().strip("\"'")
    return fam[:40] or None


def _extract_color(style: str) -> RGBColor | None:
    if not style:
        return None
    m = _COLOR_RE.search(style)
    if not m:
        return None
    return _hex_to_rgb(m.group(1))


def _extract_highlight(style: str) -> WD_COLOR_INDEX | None:
    if not style:
        return None
    m = _BG_RE.search(style)
    if not m:
        return None
    # Map any background color to a nearest highlight index (limited palette).
    return WD_COLOR_INDEX.YELLOW


def _apply_paragraph_style(paragraph, style: str) -> None:
    if not style:
        return
    m = _TEXT_ALIGN_RE.search(style)
    if m:
        v = m.group(1).lower()
        if v == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif v == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif v == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    m2 = _LINE_HEIGHT_RE.search(style)
    if m2:
        try:
            paragraph.paragraph_format.line_spacing = float(m2.group(1))
        except Exception:
            pass

    m3 = _TEXT_INDENT_RE.search(style)
    if m3:
        try:
            paragraph.paragraph_format.first_line_indent = _css_len_to_pt(m3.group(1), m3.group(2))
        except Exception:
            pass

    m4 = _MARGIN_TOP_RE.search(style)
    if m4:
        try:
            paragraph.paragraph_format.space_before = _css_len_to_pt(m4.group(1), m4.group(2))
        except Exception:
            pass

    m5 = _MARGIN_BOTTOM_RE.search(style)
    if m5:
        try:
            paragraph.paragraph_format.space_after = _css_len_to_pt(m5.group(1), m5.group(2))
        except Exception:
            pass


def _hex_to_rgb(h: str) -> RGBColor:
    s = h.lstrip("#")
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


def _extract_img_width(attrs: dict[str, str]) -> Inches | None:
    w = (attrs.get("width") or "").strip()
    if w.isdigit():
        px = int(w)
        return Inches(min(6.0, max(0.5, px / 96.0)))
    style = (attrs.get("style") or "").lower()
    m = re.search(r"width\\s*:\\s*([0-9]+)px", style)
    if m:
        px = int(m.group(1))
        return Inches(min(6.0, max(0.5, px / 96.0)))
    return None


def _css_len_to_pt(num: str, unit: str) -> Pt:
    v = float(num)
    u = (unit or "").lower()
    if u == "px":
        v = v * 0.75
    elif u == "em":
        v = v * 12.0
    v = max(0.0, min(144.0, v))
    return Pt(v)
