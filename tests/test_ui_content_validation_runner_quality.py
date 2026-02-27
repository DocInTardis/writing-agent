from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scripts import ui_content_validation_runner as runner


def test_is_format_sensitive_case_detects_format_hints() -> None:
    case = {
        "prompt": "请按指定字体字号排版，标题居中，正文行距 1.5 倍。",
        "constraints": {"format_required": False},
        "acceptance": {"required_keywords": ["字体", "字号"], "required_headings": []},
    }
    assert runner.is_format_sensitive_case(case) is True


def test_evaluate_acceptance_enforces_section_richness_for_long_cases() -> None:
    text = "\n".join(
        [
            "# 研究报告",
            "",
            "## 背景",
            "很短。",
            "",
            "## 方法",
            "很短。",
            "",
            "## 结论",
            "很短。",
        ]
    )
    acceptance = {
        "min_chars": 900,
        "required_headings": ["背景", "方法", "结论"],
        "required_keywords": [],
    }
    result = runner.evaluate_acceptance(text, acceptance)
    assert "section_body_too_short" in (result.get("failures") or [])
    assert bool((result.get("section_richness") or {}).get("enforced")) is True


def test_validate_docx_style_conformance_flags_not_centered_title(tmp_path: Path) -> None:
    out = tmp_path / "bad.docx"
    doc = Document()
    p = doc.add_paragraph("文档标题")
    p.style = doc.styles["Heading 1"]
    doc.add_paragraph("正文段落。")
    doc.save(str(out))

    report = runner.validate_docx_style_conformance(out, format_sensitive=True)
    assert report["ok"] is True
    assert report["passed"] is False
    assert "title_or_heading_not_centered" in (report.get("failures") or [])


def test_validate_docx_style_conformance_passes_centered_and_font_bound(tmp_path: Path) -> None:
    out = tmp_path / "good.docx"
    doc = Document()
    p = doc.add_paragraph("文档标题")
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body = doc.add_paragraph("这是一个用于格式校验的正文段落。")
    body.runs[0].font.name = "SimSun"
    doc.save(str(out))

    report = runner.validate_docx_style_conformance(out, format_sensitive=True)
    assert report["ok"] is True
    assert report["passed"] is True


def test_classify_precheck_warnings_keeps_autofix_non_blocking() -> None:
    result = runner.classify_precheck_warnings(["autofix_applied"])
    assert result["blocking"] == []
    assert result["non_blocking"] == ["autofix_applied"]


def test_classify_precheck_warnings_marks_compatibility_risk_blocking() -> None:
    result = runner.classify_precheck_warnings(["compatibility_risk", "corrupt_docx"])
    assert "compatibility_risk" in result["blocking"]
    assert "corrupt_docx" in result["blocking"]


def test_looks_like_connection_error_detects_refused() -> None:
    exc = RuntimeError("Page.goto: net::ERR_CONNECTION_REFUSED at http://127.0.0.1:8000/")
    assert runner._looks_like_connection_error(exc) is True
    assert runner._looks_like_connection_error(RuntimeError("other error")) is False


def test_status_indicates_transport_issue_detects_failed_fetch() -> None:
    assert runner._status_indicates_transport_issue("生成失败: Failed to fetch") is True
    assert runner._status_indicates_transport_issue("完成") is False


def test_ensure_server_available_starts_when_unreachable(monkeypatch) -> None:
    monkeypatch.setattr(runner, "url_ready", lambda *_args, **_kwargs: False)
    sentinel = SimpleNamespace()
    monkeypatch.setattr(runner, "start_local_server", lambda *_args, **_kwargs: sentinel)
    cfg = runner.RunnerConfig(
        base_url="http://127.0.0.1:8000",
        start_server=True,
        disable_ollama=False,
        headless=True,
        timeout_s=60,
        poll_interval_s=0.5,
        group_smoke=False,
        run_all=False,
        max_single=0,
        max_multi=0,
        single_start=1,
        single_end=1,
        multi_start=1,
        multi_end=1,
        export_docx_all=False,
        export_docx_for_format=False,
        checkpoint=None,
        resume=False,
        out_root=Path(".data/out"),
    )
    started = runner.ensure_server_available("http://127.0.0.1:8000", cfg, None)
    assert started is sentinel


def test_should_try_round_acceptance_repair_accepts_keyword_and_length_failures() -> None:
    acceptance = {"failures": ["length_out_of_range:120", "missing_required_keywords"]}
    assert runner.should_try_round_acceptance_repair(acceptance) is True


def test_build_round_acceptance_repair_prompt_includes_missing_terms() -> None:
    prompt = runner.build_round_acceptance_repair_prompt(
        {"min_chars": 800, "max_chars": 1200},
        {
            "char_count": 320,
            "missing_required_keywords": ["Risk Control"],
            "missing_required_headings": ["Conclusion"],
            "short_sections": ["Conclusion"],
            "empty_sections": [],
            "orphan_bilingual_lines": [],
        },
        must_keep=["Background"],
        must_change=["Risk Control"],
    )
    assert "Target range: 800-1200 characters." in prompt
    assert "Risk Control" in prompt
    assert "Conclusion" in prompt


def test_is_heading_line_rejects_numbered_list_item_with_colon() -> None:
    assert runner._is_heading_line("1. Risk: this line describes details.") is False


def test_is_heading_line_accepts_numbered_section_heading() -> None:
    assert runner._is_heading_line("1. Scope and Objectives") is True
