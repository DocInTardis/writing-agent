import io
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from types import SimpleNamespace

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.testclient import TestClient

import writing_agent.web.app_v2 as app_v2


def _create_session(client: TestClient) -> str:
    resp = client.get("/", follow_redirects=False)
    assert resp.status_code == 303
    location = resp.headers.get("location") or ""
    assert location.startswith("/workbench/")
    return location.split("/workbench/")[-1]


def _download_docx(client: TestClient, doc_id: str):
    resp = client.get(f"/download/{doc_id}.docx")
    assert resp.status_code == 200, f"{resp.status_code}: {resp.text[:300]}"
    assert resp.content
    return resp


def _extract_xml(docx_bytes: bytes, name: str) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        with zf.open(name) as handle:
            return handle.read().decode("utf-8", errors="ignore")


def _extract_xml_parts(docx_bytes: bytes, prefix: str) -> list[str]:
    parts: list[str] = []
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        for name in zf.namelist():
            if name.startswith(prefix) and name.endswith(".xml"):
                with zf.open(name) as handle:
                    parts.append(handle.read().decode("utf-8", errors="ignore"))
    return parts


def _build_docx_bytes(parts: dict[str, str]) -> bytes:
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, text in parts.items():
            zf.writestr(name, text.encode("utf-8"))
    return out.getvalue()


def test_export_docx_strict_validation():
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)

    text = "\n".join(
        [
            "# Sample Report",
            "",
            "## Introduction",
            "",
            "This is a test paragraph with consistent formatting.",
            "",
            "## Results",
            "",
            '[[TABLE:{"caption":"Metrics Summary","columns":["Metric","Value"],"rows":[["Accuracy","98%"],["Latency","120ms"]]}]]',
            "",
            '[[FIGURE:{"type":"flow","caption":"Flow Chart","data":{"nodes":["Start","Process","End"]}}]]',
            "",
            "[1] Doe. Sample Study. Journal, 2024.",
            "",
        ]
    )

    payload = {
        "text": text,
        "formatting": {
            "font_name": "Times New Roman",
            "font_name_east_asia": "SimSun",
            "font_size_name": "Small Four",
            "font_size_pt": 12,
            "line_spacing": 1.5,
            "heading1_font_name": "Times New Roman",
            "heading1_font_name_east_asia": "SimHei",
            "heading1_size_pt": 16,
            "heading2_font_name": "Times New Roman",
            "heading2_font_name_east_asia": "SimHei",
            "heading2_size_pt": 14,
            "heading3_font_name": "Times New Roman",
            "heading3_font_name_east_asia": "SimHei",
            "heading3_size_pt": 12,
        },
        "generation_prefs": {
            "include_cover": True,
            "include_toc": True,
            "toc_levels": 3,
            "include_header": True,
            "page_numbers": True,
            "header_text": "Sample Report",
            "footer_text": "",
            "page_margins_cm": 2.5,
            "page_size": "A4",
        },
    }

    env_backup = os.environ.copy()
    os.environ["WRITING_AGENT_EXPORT_MIN_FIGURES"] = "0"
    os.environ["WRITING_AGENT_EXPORT_MIN_TABLES"] = "0"
    try:
        resp = client.post(f"/api/doc/{doc_id}/save", json=payload)
        assert resp.status_code == 200

        resp = _download_docx(client, doc_id)
        docx_bytes = resp.content
    finally:
        os.environ.clear()
        os.environ.update(env_backup)

    assert (resp.headers.get("X-Docx-Export-Backend") or "").strip()
    assert (resp.headers.get("X-Docx-Style-Path") or "").strip()
    assert (resp.headers.get("X-Docx-Validation") or "").strip() in {"ok", "warning"}

    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    assert "Sample Report" in paragraphs
    assert any("Introduction" in p for p in paragraphs)
    assert any("Results" in p for p in paragraphs)
    assert any("Metrics Summary" in p for p in paragraphs)
    assert any("Flow Chart" in p for p in paragraphs)

    tables = doc.tables
    assert len(tables) >= 1
    table_text = " ".join(cell.text.strip() for row in tables[0].rows for cell in row.cells if cell.text.strip())
    assert "Metric" in table_text and "Value" in table_text and "Latency" in table_text

    combined = " ".join(paragraphs + [table_text])
    assert "[[" not in combined and "]]" not in combined
    assert "待补充" not in combined

    normal_style = doc.styles["Normal"]
    assert (normal_style.font.name or "").startswith("Times")
    assert normal_style.font.size is not None and int(normal_style.font.size.pt) == 12
    assert normal_style.paragraph_format.line_spacing in (1.5, 1.5)
    heading2_style = doc.styles["Heading 2"]
    assert heading2_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

    doc_xml = _extract_xml(docx_bytes, "word/document.xml")
    assert "w:hyperlink" in doc_xml
    assert 'w:anchor="toc_' in doc_xml
    assert "w:bookmarkStart" in doc_xml

    header_xmls = _extract_xml_parts(docx_bytes, "word/header")
    assert header_xmls
    assert any("Sample Report" in header_xml for header_xml in header_xmls)


def test_docx_export_keeps_abstract_and_keywords_sections() -> None:
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    text = "\n".join(
        [
            "# Title",
            "",
            "## 摘要",
            "",
            "这是一段摘要内容，用于验证导出链路不会删除摘要。",
            "",
            "## 关键词",
            "",
            "写作代理；DocIR；导出",
            "",
            "## 1 引言",
            "",
            "正文内容。",
            "",
            "## 参考文献",
            "",
            "[1] A. Ref. 2024. https://example.com",
        ]
    )
    payload = {
        "text": text,
        "generation_prefs": {"export_gate_policy": "off", "strict_doc_format": False, "strict_citation_verify": False},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200

    resp = _download_docx(client, doc_id)
    doc = Document(io.BytesIO(resp.content))
    body = "\n".join(p.text for p in doc.paragraphs if p.text)
    assert "摘要" in body
    assert "关键词" in body


def test_docx_export_does_not_set_update_fields_by_default(monkeypatch) -> None:
    monkeypatch.delenv("WRITING_AGENT_DOCX_UPDATE_FIELDS_ON_OPEN", raising=False)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# 测试文档\n\n## 背景\n\n正文。\n\n## 结论\n\n正文。",
        "generation_prefs": {"export_gate_policy": "off", "strict_doc_format": False, "strict_citation_verify": False},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    settings_xml = _extract_xml(resp.content, "word/settings.xml")
    assert "w:updateFields" not in settings_xml


def test_docx_export_sets_update_fields_when_opt_in(monkeypatch) -> None:
    monkeypatch.setenv("WRITING_AGENT_DOCX_UPDATE_FIELDS_ON_OPEN", "1")
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# 测试文档\n\n## 背景\n\n正文。\n\n## 结论\n\n正文。",
        "generation_prefs": {"export_gate_policy": "off", "strict_doc_format": False, "strict_citation_verify": False},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    settings_xml = _extract_xml(resp.content, "word/settings.xml")
    assert "w:updateFields" in settings_xml
    assert 'w:val="true"' in settings_xml


def test_docx_export_disables_toc_footer_postprocess_by_default(monkeypatch) -> None:
    monkeypatch.delenv("WRITING_AGENT_DOCX_TOC_FOOTER_POSTPROCESS", raising=False)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# Title\n\n## Intro\n\nBody",
        "generation_prefs": {"include_cover": True, "include_toc": True, "page_numbers": True, "export_gate_policy": "off"},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    footer_xmls = _extract_xml_parts(resp.content, "word/footer")
    assert footer_xmls
    # Default path should keep native python-docx namespace prefix style.
    assert not any("<ns0:ftr" in x for x in footer_xmls)


def test_docx_export_uses_clickable_toc_links_by_default(monkeypatch) -> None:
    monkeypatch.delenv("WRITING_AGENT_DOCX_DYNAMIC_TOC_FIELD", raising=False)
    monkeypatch.delenv("WRITING_AGENT_DOCX_TOC_HYPERLINK", raising=False)
    monkeypatch.delenv("WRITING_AGENT_DOCX_TOC_FIELD_LOCK", raising=False)
    monkeypatch.delenv("WRITING_AGENT_DOCX_TOC_CLICKABLE_LINKS", raising=False)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# 标题\n\n## 第一章 绪论\n\n正文\n\n### 1.1 背景\n\n正文",
        "generation_prefs": {"include_cover": True, "include_toc": True, "page_numbers": True, "export_gate_policy": "off"},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    document_xml = _extract_xml(resp.content, "word/document.xml")
    assert 'TOC \\\\o' not in document_xml
    assert "w:hyperlink" in document_xml
    assert 'w:anchor="toc_' in document_xml
    assert "w:bookmarkStart" in document_xml
    assert 'w:leader="dot"' in document_xml
    assert "<w:tab/>" in document_xml


def test_docx_export_supports_static_toc_when_dynamic_disabled(monkeypatch) -> None:
    monkeypatch.setenv("WRITING_AGENT_DOCX_DYNAMIC_TOC_FIELD", "0")
    monkeypatch.delenv("WRITING_AGENT_DOCX_TOC_CLICKABLE_LINKS", raising=False)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# 标题\n\n## 第一章 绪论\n\n正文\n\n### 1.1 背景\n\n正文",
        "generation_prefs": {"include_cover": True, "include_toc": True, "page_numbers": True, "export_gate_policy": "off"},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    document_xml = _extract_xml(resp.content, "word/document.xml")
    assert 'TOC \\\\o' not in document_xml
    assert "w:hyperlink" in document_xml
    assert 'w:anchor="toc_' in document_xml


def test_docx_export_supports_dynamic_toc_when_clickable_disabled(monkeypatch) -> None:
    monkeypatch.setenv("WRITING_AGENT_DOCX_TOC_CLICKABLE_LINKS", "0")
    monkeypatch.setenv("WRITING_AGENT_DOCX_DYNAMIC_TOC_FIELD", "1")
    monkeypatch.delenv("WRITING_AGENT_DOCX_TOC_HYPERLINK", raising=False)
    monkeypatch.delenv("WRITING_AGENT_DOCX_TOC_FIELD_LOCK", raising=False)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# 鏍囬\n\n## 绗竴绔?缁\n\n姝ｆ枃\n\n### 1.1 鑳屾櫙\n\n姝ｆ枃",
        "generation_prefs": {"include_cover": True, "include_toc": True, "page_numbers": True, "export_gate_policy": "off"},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    document_xml = _extract_xml(resp.content, "word/document.xml")
    assert 'TOC \\\\o' in document_xml
    assert 'w:fldLock="true"' in document_xml


def test_export_template_policy_default_disables_legacy_auto_pick(monkeypatch) -> None:
    monkeypatch.delenv("WRITING_AGENT_EXPORT_AUTO_TEMPLATE", raising=False)
    monkeypatch.delenv("WRITING_AGENT_EXPORT_DEFAULT_TEMPLATE", raising=False)
    session = SimpleNamespace(template_source_path="")
    resolved = app_v2._resolve_export_template_path(session)
    assert resolved == ""


def test_export_template_policy_allows_explicit_default_template(tmp_path, monkeypatch) -> None:
    tmpl = tmp_path / "clean.docx"
    Document().save(str(tmpl))
    monkeypatch.delenv("WRITING_AGENT_EXPORT_AUTO_TEMPLATE", raising=False)
    monkeypatch.setenv("WRITING_AGENT_EXPORT_DEFAULT_TEMPLATE", str(tmpl))
    session = SimpleNamespace(template_source_path="")
    resolved = app_v2._resolve_export_template_path(session)
    assert resolved == str(tmpl)


def test_dedupe_equivalent_headings_prefers_chinese() -> None:
    raw = "\n".join(
        [
            "# 报告",
            "",
            "## Background",
            "A段。",
            "",
            "## 背景",
            "B段。",
            "",
            "## Conclusion",
            "C段。",
            "",
            "## 结论",
            "D段。",
        ]
    )
    fixed = app_v2._dedupe_equivalent_headings(raw)
    assert "## Background" not in fixed
    assert "## Conclusion" not in fixed
    assert fixed.count("## 背景") == 1
    assert fixed.count("## 结论") == 1
    assert "A段。" in fixed and "B段。" in fixed and "C段。" in fixed and "D段。" in fixed


def test_dedupe_equivalent_headings_compacts_toc_entries() -> None:
    raw = "\n".join(
        [
            "# 项目报告",
            "",
            "## 目录",
            "",
            "1. Background",
            "2. 背景",
            "3. Current State",
            "4. 当前状态",
            "",
            "## Background",
            "内容1。",
            "",
            "## 背景",
            "内容2。",
            "",
            "## Current State",
            "内容3。",
            "",
            "## 当前状态",
            "内容4。",
        ]
    )
    fixed = app_v2._dedupe_equivalent_headings(raw)
    assert "1. 背景" in fixed
    assert "2. 当前状态" in fixed
    assert "3. Current State" not in fixed
    assert "## Background" not in fixed
    assert "## Current State" not in fixed


def test_dedupe_equivalent_headings_handles_synonym_aliases() -> None:
    raw = "\n".join(
        [
            "# \u62a5\u544a",
            "",
            "## \u5f53\u524d\u72b6\u6001\u5f53\u524d\u9636\u6bb5",
            "\u4e2d\u6587\u6bb5\u843d\u3002",
            "",
            "## Current State",
            "English paragraph.",
            "",
            "## \u63a8\u8350\u63aa\u65bd",
            "\u4e2d\u6587\u5efa\u8bae\u3002",
            "",
            "## Recommendations",
            "English recommendations.",
        ]
    )
    fixed = app_v2._dedupe_equivalent_headings(raw)
    assert fixed.count("## \u5f53\u524d\u72b6\u6001\u5f53\u524d\u9636\u6bb5") == 1
    assert "## Current State" not in fixed
    assert fixed.count("## \u63a8\u8350\u63aa\u65bd") == 1
    assert "## Recommendations" not in fixed


def test_fix_section_heading_glue_keeps_plain_numbered_list_items() -> None:
    raw = "\n".join(
        [
            "# T",
            "",
            "1. **Internal Sources**: include technical docs.",
            "2. **External Sources**: include interviews.",
            "",
        ]
    )
    fixed = app_v2._fix_section_heading_glue(raw, ["Internal Sources", "External Sources"])
    assert "## Internal Sources" not in fixed
    assert "## External Sources" not in fixed
    assert "1. **Internal Sources**: include technical docs." in fixed
    assert "2. **External Sources**: include interviews." in fixed


def test_export_autofix_not_persisted_by_default(monkeypatch) -> None:
    monkeypatch.delenv("WRITING_AGENT_EXPORT_PERSIST_AUTOFIX", raising=False)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# Sample\n\n## Background\n\nBody text.\n",
        "generation_prefs": {"strict_doc_format": True, "strict_citation_verify": False, "export_gate_policy": "strict"},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    before = client.get(f"/api/doc/{doc_id}").json().get("text", "")
    assert "## 目录" not in before
    _download_docx(client, doc_id)
    after = client.get(f"/api/doc/{doc_id}").json().get("text", "")
    assert after == before


def test_clean_export_text_preserves_markdown_emphasis_and_list_markers() -> None:
    raw = "\n".join(
        [
            "# Demo",
            "",
            "- first bullet",
            "* second bullet",
            "",
            "This is *emphasis* and **strong** text.",
        ]
    )
    cleaned = app_v2._clean_export_text(raw)
    assert "- first bullet" in cleaned
    assert "* second bullet" in cleaned
    assert "*emphasis*" in cleaned
    assert "**strong**" in cleaned


def test_ensure_reference_section_does_not_inject_placeholder_without_citations() -> None:
    session = SimpleNamespace(citations={}, formatting={}, generation_prefs={})
    text = "# Demo\n\n## Intro\ncontent"
    fixed = app_v2._ensure_reference_section(text, session)
    assert fixed == text
    assert "请补充可核验参考文献" not in fixed


def test_strict_export_defaults_disabled_without_explicit_opt_in(monkeypatch) -> None:
    monkeypatch.delenv("WRITING_AGENT_STRICT_DOC_FORMAT_DEFAULT", raising=False)
    monkeypatch.delenv("WRITING_AGENT_STRICT_CITATION_VERIFY_DEFAULT", raising=False)
    session = SimpleNamespace(generation_prefs={})
    assert app_v2._strict_doc_format_enabled(session) is False
    assert app_v2._strict_citation_verify_enabled(session) is False


def test_normalize_export_text_only_dedupes_headings_in_strict_mode() -> None:
    raw = "\n".join(
        [
            "# Demo",
            "",
            "## Background",
            "English paragraph.",
            "",
            "## 背景",
            "中文段落。",
        ]
    )
    relaxed = app_v2._normalize_export_text(raw, session=SimpleNamespace(generation_prefs={}))
    strict = app_v2._normalize_export_text(raw, session=SimpleNamespace(generation_prefs={"strict_doc_format": True}))
    assert "## Background" in relaxed
    assert "## 背景" in relaxed
    assert strict != relaxed

def test_docx_export_footer_uses_default_refs_and_non_empty_page_fields(monkeypatch) -> None:
    monkeypatch.delenv("WA_USE_RUST_ENGINE", raising=False)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# Sample Report\n\n## Introduction\n\nBody paragraph.\n",
        "generation_prefs": {
            "include_cover": True,
            "include_toc": True,
            "toc_levels": 2,
            "include_header": True,
            "page_numbers": True,
            "header_text": "Sample Report",
            "footer_text": "",
            "export_gate_policy": "off",
            "strict_doc_format": False,
            "strict_citation_verify": False,
        },
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)

    doc_xml = _extract_xml(resp.content, "word/document.xml")
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    root = ET.fromstring(doc_xml)
    footer_refs = root.findall(".//w:sectPr/w:footerReference", ns)
    footer_types = [str(ref.attrib.get(f"{{{ns['w']}}}type") or "default").lower() for ref in footer_refs]
    assert "first" not in footer_types
    assert "even" not in footer_types

    footer_xmls = _extract_xml_parts(resp.content, "word/footer")
    assert footer_xmls
    has_page_field = False
    for footer_xml in footer_xmls:
        if re.search(r"<(?:\w+:)?instrText[^>]*>\s*PAGE[^<]*</(?:\w+:)?instrText>", footer_xml, flags=re.IGNORECASE):
            has_page_field = True
            assert re.search(r"<(?:\w+:)?t[^>]*>\s*(I|1)\s*</(?:\w+:)?t>", footer_xml) is not None
            assert re.search(r"<(?:\w+:)?t[^>]*/>", footer_xml) is None
            assert re.search(r"<(?:\w+:)?t[^>]*>\s*</(?:\w+:)?t>", footer_xml) is None
    assert has_page_field is True


def test_validate_docx_bytes_flags_nondefault_footer_and_empty_page_field() -> None:
    parts = {
        "[Content_Types].xml": """<?xml version=\"1.0\" encoding=\"UTF-8\"?>
<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">
  <Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>
  <Default Extension=\"xml\" ContentType=\"application/xml\"/>
</Types>""",
        "_rels/.rels": """<?xml version=\"1.0\" encoding=\"UTF-8\"?>
<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">
  <Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>
</Relationships>""",
        "word/document.xml": """<?xml version=\"1.0\" encoding=\"UTF-8\"?>
<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"
            xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\">
  <w:body>
    <w:p><w:r><w:t>demo</w:t></w:r></w:p>
    <w:sectPr>
      <w:footerReference w:type=\"first\" r:id=\"rIdFooter1\"/>
    </w:sectPr>
  </w:body>
</w:document>""",
        "word/_rels/document.xml.rels": """<?xml version=\"1.0\" encoding=\"UTF-8\"?>
<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">
  <Relationship Id=\"rIdFooter1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer\" Target=\"footer1.xml\"/>
</Relationships>""",
        "word/footer1.xml": """<?xml version=\"1.0\" encoding=\"UTF-8\"?>
<w:ftr xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">
  <w:p>
    <w:r><w:fldChar w:fldCharType=\"begin\"/></w:r>
    <w:r><w:instrText xml:space=\"preserve\">PAGE</w:instrText></w:r>
    <w:r><w:fldChar w:fldCharType=\"separate\"/></w:r>
    <w:r><w:t xml:space=\"preserve\"></w:t></w:r>
    <w:r><w:fldChar w:fldCharType=\"end\"/></w:r>
  </w:p>
</w:ftr>""",
    }
    bad_docx = _build_docx_bytes(parts)
    issues = app_v2._validate_docx_bytes(bad_docx)
    assert any(x.startswith("footer-ref:nondefault:first") for x in issues)
    assert any(x.startswith("page-field-empty:word/footer1.xml") for x in issues)


def test_docx_export_single_mode_ignores_backend_env_safe(monkeypatch) -> None:
    monkeypatch.setenv("WRITING_AGENT_DOCX_BACKEND_MODE", "safe")
    monkeypatch.delenv("WA_USE_RUST_ENGINE", raising=False)
    monkeypatch.setattr(app_v2, "_doc_ir_has_styles", lambda _doc_ir: True)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# Safe Mode\n\n## Intro\n\nBody text.",
        "generation_prefs": {"export_gate_policy": "off", "strict_doc_format": False, "strict_citation_verify": False},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    assert resp.headers.get("X-Docx-Export-Backend") == "parsed_docx_exporter"
    assert resp.headers.get("X-Docx-Backend-Mode") == "single_parsed"


def test_docx_export_single_mode_ignores_backend_env_auto(monkeypatch) -> None:
    monkeypatch.setenv("WRITING_AGENT_DOCX_BACKEND_MODE", "auto")
    monkeypatch.setattr(app_v2, "_doc_ir_has_styles", lambda _doc_ir: True)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# Auto Mode\n\n## Intro\n\nBody text.",
        "generation_prefs": {"export_gate_policy": "off", "strict_doc_format": False, "strict_citation_verify": False},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    assert resp.headers.get("X-Docx-Export-Backend") == "parsed_docx_exporter"
    assert resp.headers.get("X-Docx-Backend-Mode") == "single_parsed"


def test_docx_export_blocks_when_validation_fails_and_enforced(monkeypatch) -> None:
    monkeypatch.setenv("WRITING_AGENT_DOCX_VALIDATION_ENFORCE", "1")
    monkeypatch.setattr(app_v2, "_validate_docx_bytes", lambda _payload: ["simulated-invalid-docx"])
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# Validate Block\n\n## Intro\n\nBody text.",
        "generation_prefs": {"export_gate_policy": "off", "strict_doc_format": False, "strict_citation_verify": False},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = client.get(f"/download/{doc_id}.docx")
    assert resp.status_code == 500
    assert "DOCX导出失败" in (resp.text or "")


def test_docx_export_uses_fallback_when_first_validation_fails(monkeypatch) -> None:
    monkeypatch.setenv("WRITING_AGENT_DOCX_VALIDATION_ENFORCE", "1")
    calls = {"n": 0}

    def _validate_once_fail(_payload: bytes) -> list[str]:
        calls["n"] += 1
        return ["simulated-invalid-docx"] if calls["n"] == 1 else []

    monkeypatch.setattr(app_v2, "_validate_docx_bytes", _validate_once_fail)
    client = TestClient(app_v2.app)
    doc_id = _create_session(client)
    payload = {
        "text": "# Validate Fallback\n\n## Intro\n\nBody text.",
        "generation_prefs": {"export_gate_policy": "off", "strict_doc_format": False, "strict_citation_verify": False},
    }
    save = client.post(f"/api/doc/{doc_id}/save", json=payload)
    assert save.status_code == 200
    resp = _download_docx(client, doc_id)
    assert resp.headers.get("X-Docx-Repair") == "fallback_no_toc_header_pagenum"
    assert calls["n"] >= 2
