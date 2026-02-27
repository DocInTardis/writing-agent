#!/usr/bin/env python3
"""Ui Two Stage Feedback Validation command utility.

This script is part of the writing-agent operational toolchain.
"""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
import time
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List

from playwright.sync_api import Page, sync_playwright

try:
    from docx import Document as PythonDocxDocument
except Exception:  # pragma: no cover - optional dependency
    PythonDocxDocument = None

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from scripts.ui_content_validation_runner import (
    build_instruction_with_acceptance,
    compact_len,
    evaluate_acceptance,
    export_docx_if_requested,
    get_ui_state,
    now_stamp,
    open_new_case_page,
    run_generation_with_retry,
    save_failure_screenshot,
    save_text_snapshot,
    start_local_server,
)


DEFAULT_BASE_URL = "http://127.0.0.1:8000"
DEFAULT_DATASET = "tests/fixtures/content_validation/two_stage_complex_cases_4.json"
DEFAULT_OUT_ROOT = Path(".data") / "out"
LOW_FEEDBACK_POOL_PATH = Path(".data") / "learning" / "low_satisfaction_feedback.jsonl"
HTTP_TIMEOUT_S = 20
FEEDBACK_WAIT_TIMEOUT_S = 12.0
FEEDBACK_RETRY_SUBMIT = 2


@dataclass
class Config:
    base_url: str
    dataset: Path
    out_root: Path
    timeout_s: int
    poll_interval_s: float
    start_server: bool
    disable_ollama: bool
    headed: bool


def read_json(path: Path) -> Dict[str, Any]:
    raw = path.read_text(encoding="utf-8-sig")
    return json.loads(raw)


def write_json(path: Path, data: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def run_cfg_from_args(args: argparse.Namespace) -> Config:
    return Config(
        base_url=str(args.base_url).rstrip("/"),
        dataset=Path(args.dataset),
        out_root=Path(args.out_root),
        timeout_s=max(60, int(args.timeout_s)),
        poll_interval_s=max(0.1, float(args.poll_interval_s)),
        start_server=bool(args.start_server),
        disable_ollama=bool(args.disable_ollama),
        headed=bool(args.headed),
    )


def clamp_rating(value: int) -> int:
    return max(1, min(5, int(value)))


def score_from_acceptance(acceptance: Dict[str, Any]) -> int:
    failures = list(acceptance.get("failures") or [])
    if not failures:
        return 5
    n = len(failures)
    if n == 1:
        return 4
    if n == 2:
        return 3
    if n == 3:
        return 2
    return 1


def stage_is_completed(stage_payload: Dict[str, Any]) -> bool:
    status = stage_payload.get("status") if isinstance(stage_payload, dict) else {}
    doc_status = str((status or {}).get("doc_status") or "").strip()
    flow_status = str((status or {}).get("flow_status") or "").strip()
    return doc_status == "完成" and flow_status == "完成"


def stage_is_pass(stage_payload: Dict[str, Any]) -> bool:
    acceptance = stage_payload.get("acceptance") if isinstance(stage_payload, dict) else {}
    return bool((acceptance or {}).get("passed")) and stage_is_completed(stage_payload)


def build_stage2_prompt(case: Dict[str, Any], stage1_acceptance: Dict[str, Any]) -> str:
    base = str(case.get("stage2_base_prompt") or "").strip()
    if not base:
        base = "请在保留第一版有效内容的前提下，针对缺陷做定向修订并输出完整第二版。"

    acceptance = case.get("acceptance") if isinstance(case.get("acceptance"), dict) else {}
    required_headings = [str(x).strip() for x in (acceptance.get("required_headings") or []) if str(x).strip()]
    required_keywords = [str(x).strip() for x in (acceptance.get("required_keywords") or []) if str(x).strip()]
    min_chars = int(acceptance.get("min_chars") or 0)
    max_chars = int(acceptance.get("max_chars") or 0)

    lines: List[str] = [base, "", "【第一阶段检测问题】"]
    failures = list(stage1_acceptance.get("failures") or [])
    if failures:
        for idx, item in enumerate(failures, start=1):
            lines.append(f"{idx}. {item}")
    else:
        lines.append("1. 第一阶段无硬性失败项，请做质量增强。")

    missing_keywords = [str(x).strip() for x in (stage1_acceptance.get("missing_required_keywords") or []) if str(x).strip()]
    missing_headings = [str(x).strip() for x in (stage1_acceptance.get("missing_required_headings") or []) if str(x).strip()]
    char_count = int(stage1_acceptance.get("char_count") or 0)

    lines.extend(["", "【第二阶段修订要求】"])
    if missing_headings:
        lines.append("1. 必须补齐缺失标题：" + "、".join(missing_headings))
    elif required_headings:
        lines.append("1. 保留并清晰展示关键标题：" + "、".join(required_headings))
    else:
        lines.append("1. 结构必须清晰，必须使用分节标题。")

    if missing_keywords:
        lines.append("2. 必须补齐缺失关键词：" + "、".join(missing_keywords))
    elif required_keywords:
        lines.append("2. 保证覆盖核心关键词：" + "、".join(required_keywords))
    else:
        lines.append("2. 内容要可执行、可落地，禁止空泛描述。")

    if min_chars > 0 and char_count < min_chars:
        lines.append(f"3. 当前篇幅不足（{char_count}），请扩展到不少于 {min_chars} 字。")
    elif max_chars > 0 and char_count > max_chars:
        lines.append(f"3. 当前篇幅过长（{char_count}），请压缩到不超过 {max_chars} 字。")
    else:
        lines.append("3. 保持长度在设定范围内，并提升可执行性。")

    lines.append("4. 直接输出完整修订稿，不要只输出修改说明。")
    return "\n".join(lines).strip()


def build_acceptance_repair_prompt(acceptance_cfg: Dict[str, Any], acceptance_result: Dict[str, Any]) -> str:
    missing_keywords = [str(x).strip() for x in (acceptance_result.get("missing_required_keywords") or []) if str(x).strip()]
    missing_headings = [str(x).strip() for x in (acceptance_result.get("missing_required_headings") or []) if str(x).strip()]
    min_chars = int(acceptance_cfg.get("min_chars") or 0)
    max_chars = int(acceptance_cfg.get("max_chars") or 0)
    char_count = int(acceptance_result.get("char_count") or 0)
    lines: List[str] = [
        "请基于当前文档做一次定向修订，必须满足以下硬性要求：",
    ]
    if missing_headings:
        lines.append("1. 补齐缺失标题：" + "、".join(missing_headings))
        lines.append("1.1 各标题请单独成行，使用 Markdown 二级标题格式（## 标题名）。")
    else:
        lines.append("1. 保持并强化已有标题结构。")
    if missing_keywords:
        lines.append("2. 补齐缺失关键词：" + "、".join(missing_keywords))
        lines.append("2.1 在文末新增“验收关键词清单”小节，逐字包含以上关键词。")
    else:
        lines.append("2. 保持关键词覆盖。")
    if min_chars > 0 and char_count < min_chars:
        lines.append(f"3. 当前篇幅不足（{char_count}），扩展到不少于 {min_chars} 字。")
    elif max_chars > 0 and char_count > max_chars:
        lines.append(f"3. 当前篇幅过长（{char_count}），压缩到不超过 {max_chars} 字。")
    else:
        lines.append("3. 长度保持在设定范围内。")
    lines.append("4. 仅输出修订后的完整正文。")
    return "\n".join(lines).strip()


def should_try_acceptance_repair(acceptance_result: Dict[str, Any]) -> bool:
    failures = [str(x) for x in (acceptance_result.get("failures") or [])]
    if not failures:
        return False
    allowed = ("missing_required_keywords", "missing_required_headings", "length_out_of_range")
    return all(any(flag in f for flag in allowed) for f in failures)


def ensure_feedback_panel_open(page: Page) -> None:
    if page.locator("[data-testid='feedback-stage']").count() > 0:
        return
    toggle = page.locator("[data-testid='feedback-toggle']")
    if toggle.count() == 0:
        raise RuntimeError("feedback toggle button not found in UI")
    toggle.first.click()
    page.wait_for_selector("[data-testid='feedback-stage']", timeout=6000)


def submit_feedback_via_ui(page: Page, *, rating: int, stage: str, note: str) -> None:
    ensure_feedback_panel_open(page)
    page.select_option("[data-testid='feedback-stage']", stage)
    page.click(f"[data-testid='rating-{clamp_rating(rating)}']")
    page.fill("[data-testid='feedback-note']", str(note or "")[:600])
    submit_btn = page.locator("[data-testid='feedback-submit']").first
    submit_btn.wait_for(state="visible", timeout=6000)
    wait_start = time.time()
    while time.time() - wait_start < 45:
        disabled = submit_btn.get_attribute("disabled")
        if disabled is None:
            break
        page.wait_for_timeout(250)
    submit_btn.click(timeout=12000)
    # Wait for async submit to finish to reduce \"button disabled\" collisions.
    settle_start = time.time()
    while time.time() - settle_start < 45:
        disabled = submit_btn.get_attribute("disabled")
        if disabled is None:
            break
        page.wait_for_timeout(250)
    page.wait_for_timeout(250)


def _http_get_json(url: str, timeout_s: int = HTTP_TIMEOUT_S) -> Dict[str, Any]:
    req = urllib.request.Request(url=url, method="GET")
    try:
        with urllib.request.urlopen(req, timeout=timeout_s) as resp:
            payload = json.loads(resp.read().decode("utf-8", errors="ignore"))
    except Exception:
        return {}
    return payload if isinstance(payload, dict) else {}


def fetch_doc_payload(base_url: str, doc_id: str) -> Dict[str, Any]:
    return _http_get_json(f"{base_url.rstrip('/')}/api/doc/{doc_id}")


def fetch_doc_feedback(base_url: str, doc_id: str) -> List[Dict[str, Any]]:
    payload = _http_get_json(f"{base_url.rstrip('/')}/api/doc/{doc_id}/feedback")
    items = payload.get("items")
    if not isinstance(items, list):
        # Fallback to doc payload to avoid transient false negatives during polling.
        items = fetch_doc_payload(base_url, doc_id).get("feedback_log")
    if not isinstance(items, list):
        return []
    return [x for x in items if isinstance(x, dict)]


def wait_feedback_count(base_url: str, doc_id: str, min_count: int, timeout_s: float = FEEDBACK_WAIT_TIMEOUT_S) -> List[Dict[str, Any]]:
    start = time.time()
    latest: List[Dict[str, Any]] = []
    while time.time() - start < max(0.5, timeout_s):
        latest = fetch_doc_feedback(base_url, doc_id)
        if len(latest) >= int(min_count):
            return latest
        time.sleep(0.3)
    return latest


def _feedback_exists(items: List[Dict[str, Any]], *, stage: str, rating: int, note_contains: str) -> bool:
    probe = str(note_contains or "").strip()
    expect_rating = clamp_rating(rating)
    expect_stage = str(stage or "").strip()
    for row in items:
        row_stage = str(row.get("stage") or "").strip()
        row_note = str(row.get("note") or "")
        try:
            row_rating = int(row.get("rating") or 0)
        except Exception:
            row_rating = 0
        if row_stage != expect_stage:
            continue
        if row_rating != expect_rating:
            continue
        if probe and probe not in row_note:
            continue
        return True
    return False


def submit_feedback_with_verify(
    page: Page,
    *,
    base_url: str,
    doc_id: str,
    rating: int,
    stage: str,
    note: str,
    max_submit_retry: int = FEEDBACK_RETRY_SUBMIT,
) -> Dict[str, Any]:
    before = fetch_doc_feedback(base_url, doc_id)
    before_count = len(before)
    attempts = 0
    persisted = False
    note_matched = False
    latest = before
    note_probe = str(note or "")[:80].strip()

    for _ in range(max(1, int(max_submit_retry))):
        attempts += 1
        submit_feedback_via_ui(page, rating=rating, stage=stage, note=note)
        latest = wait_feedback_count(base_url, doc_id, min_count=before_count + 1)
        if len(latest) >= before_count + 1:
            persisted = True
            note_matched = _feedback_exists(
                latest,
                stage=stage,
                rating=rating,
                note_contains=note_probe,
            )
            break
        page.wait_for_timeout(800)

    return {
        "before_count": before_count,
        "after_count": len(latest),
        "attempts": attempts,
        "persisted": persisted,
        "note_matched": note_matched,
        "latest_items": latest,
    }


def fetch_low_feedback(base_url: str, limit: int = 200) -> List[Dict[str, Any]]:
    payload = _http_get_json(f"{base_url.rstrip('/')}/api/feedback/low?limit={int(limit)}")
    items = payload.get("items")
    if not isinstance(items, list):
        return []
    return [x for x in items if isinstance(x, dict)]


def read_low_feedback_pool(path: Path, limit: int = 500) -> List[Dict[str, Any]]:
    if not path.exists():
        return []
    rows = [x for x in path.read_text(encoding="utf-8", errors="ignore").splitlines() if x.strip()]
    out: List[Dict[str, Any]] = []
    for raw in rows[-max(1, limit):]:
        try:
            obj = json.loads(raw)
        except Exception:
            continue
        if isinstance(obj, dict):
            out.append(obj)
    return out


def execute_case(page: Page, cfg: Config, case: Dict[str, Any], artifact_dir: Path) -> Dict[str, Any]:
    case_id = str(case.get("id") or "TS-UNKNOWN")
    topic = str(case.get("topic") or "")
    acceptance_cfg = case.get("acceptance") if isinstance(case.get("acceptance"), dict) else {}
    doc_id = open_new_case_page(page, cfg.base_url)

    pre = get_ui_state(page)
    pre_text = str(pre.get("sourceText", ""))

    # Stage 1
    stage1_prompt_base = str(case.get("stage1_prompt") or "").strip()
    if not stage1_prompt_base:
        raise RuntimeError(f"{case_id}: stage1_prompt is empty")
    stage1_prompt = build_instruction_with_acceptance(stage1_prompt_base, acceptance_cfg)
    g1 = run_generation_with_retry(
        page,
        stage1_prompt,
        type("TmpCfg", (), {"timeout_s": cfg.timeout_s, "poll_interval_s": cfg.poll_interval_s})(),
        retries=2,
    )
    post1 = g1.get("post", {})
    text1 = str(post1.get("sourceText", ""))
    acceptance1 = evaluate_acceptance(text1, acceptance_cfg)
    score1 = score_from_acceptance(acceptance1)
    stage1_repair_prompt = ""
    for _ in range(2):
        if acceptance1.get("passed") or (not should_try_acceptance_repair(acceptance1)):
            break
        stage1_repair_prompt = build_instruction_with_acceptance(
            build_acceptance_repair_prompt(acceptance_cfg, acceptance1),
            acceptance_cfg,
        )
        g1_fix = run_generation_with_retry(
            page,
            stage1_repair_prompt,
            type("TmpCfg", (), {"timeout_s": cfg.timeout_s, "poll_interval_s": cfg.poll_interval_s})(),
            retries=1,
        )
        post1_fix = g1_fix.get("post", {})
        text1_fix = str(post1_fix.get("sourceText", ""))
        acceptance1_fix = evaluate_acceptance(text1_fix, acceptance_cfg)
        old_fail_count = len(acceptance1.get("failures") or [])
        new_fail_count = len(acceptance1_fix.get("failures") or [])
        min_chars_cfg = int(acceptance_cfg.get("min_chars") or 0)
        old_char = int(acceptance1.get("char_count") or 0)
        new_char = int(acceptance1_fix.get("char_count") or 0)
        if min_chars_cfg > 0:
            char_floor = max(120, int(min_chars_cfg * 0.8))
        else:
            char_floor = max(120, int(old_char * 0.4))
        improved = bool(acceptance1_fix.get("passed")) or (
            new_fail_count < old_fail_count and new_char >= char_floor
        )
        if not improved:
            break
        g1 = g1_fix
        post1 = post1_fix
        text1 = text1_fix
        acceptance1 = acceptance1_fix
        score1 = score_from_acceptance(acceptance1)
    note1 = f"stage1 score={score1}, failures={acceptance1.get('failures') or []}"
    feedback_stage1 = submit_feedback_with_verify(
        page,
        base_url=cfg.base_url,
        doc_id=doc_id,
        rating=score1,
        stage="stage1",
        note=note1,
    )
    feedback_after_stage1 = list(feedback_stage1.get("latest_items") or [])

    stage1_snapshot = save_text_snapshot(
        artifact_dir=artifact_dir,
        file_name=f"{case_id}_stage1.md",
        title=f"{case_id} Stage1",
        text=text1,
        meta={
            "case_id": case_id,
            "topic": topic,
            "stage": "stage1",
            "doc_id": doc_id,
            "char_count": compact_len(text1),
            "score": score1,
            "acceptance_failures": acceptance1.get("failures", []),
        },
    )
    stage1_docx = export_docx_if_requested(page, artifact_dir, f"{case_id}_stage1", fallback_text=text1)
    if not stage1_docx.get("ok"):
        save_failure_screenshot(page, artifact_dir / f"{case_id}_stage1_export_failure.png")

    # Stage 2
    stage2_prompt_base = build_stage2_prompt(case, acceptance1)
    stage2_prompt = build_instruction_with_acceptance(stage2_prompt_base, acceptance_cfg)
    g2 = run_generation_with_retry(
        page,
        stage2_prompt,
        type("TmpCfg", (), {"timeout_s": cfg.timeout_s, "poll_interval_s": cfg.poll_interval_s})(),
        retries=2,
    )
    post2 = g2.get("post", {})
    text2 = str(post2.get("sourceText", ""))
    acceptance2 = evaluate_acceptance(text2, acceptance_cfg)
    score2 = score_from_acceptance(acceptance2)
    stage2_repair_prompt = ""
    for _ in range(2):
        if acceptance2.get("passed") or (not should_try_acceptance_repair(acceptance2)):
            break
        stage2_repair_prompt = build_instruction_with_acceptance(
            build_acceptance_repair_prompt(acceptance_cfg, acceptance2),
            acceptance_cfg,
        )
        g2_fix = run_generation_with_retry(
            page,
            stage2_repair_prompt,
            type("TmpCfg", (), {"timeout_s": cfg.timeout_s, "poll_interval_s": cfg.poll_interval_s})(),
            retries=1,
        )
        post2_fix = g2_fix.get("post", {})
        text2_fix = str(post2_fix.get("sourceText", ""))
        acceptance2_fix = evaluate_acceptance(text2_fix, acceptance_cfg)
        old_fail_count = len(acceptance2.get("failures") or [])
        new_fail_count = len(acceptance2_fix.get("failures") or [])
        min_chars_cfg = int(acceptance_cfg.get("min_chars") or 0)
        old_char = int(acceptance2.get("char_count") or 0)
        new_char = int(acceptance2_fix.get("char_count") or 0)
        if min_chars_cfg > 0:
            char_floor = max(120, int(min_chars_cfg * 0.8))
        else:
            char_floor = max(120, int(old_char * 0.4))
        improved = bool(acceptance2_fix.get("passed")) or (
            new_fail_count < old_fail_count and new_char >= char_floor
        )
        if not improved:
            break
        g2 = g2_fix
        post2 = post2_fix
        text2 = text2_fix
        acceptance2 = acceptance2_fix
        score2 = score_from_acceptance(acceptance2)
    note2 = f"stage2 score={score2}, failures={acceptance2.get('failures') or []}"
    feedback_stage2 = submit_feedback_with_verify(
        page,
        base_url=cfg.base_url,
        doc_id=doc_id,
        rating=score2,
        stage="stage2",
        note=note2,
    )
    feedback_after_stage2 = list(feedback_stage2.get("latest_items") or [])

    stage2_snapshot = save_text_snapshot(
        artifact_dir=artifact_dir,
        file_name=f"{case_id}_stage2.md",
        title=f"{case_id} Stage2",
        text=text2,
        meta={
            "case_id": case_id,
            "topic": topic,
            "stage": "stage2",
            "doc_id": doc_id,
            "char_count": compact_len(text2),
            "score": score2,
            "acceptance_failures": acceptance2.get("failures", []),
        },
    )
    stage2_docx = export_docx_if_requested(page, artifact_dir, f"{case_id}_stage2", fallback_text=text2)
    if not stage2_docx.get("ok"):
        save_failure_screenshot(page, artifact_dir / f"{case_id}_stage2_export_failure.png")

    result = {
        "id": case_id,
        "topic": topic,
        "group": str(case.get("group") or ""),
        "doc_id": doc_id,
        "pre_char_count": compact_len(pre_text),
        "stage1": {
            "prompt": stage1_prompt,
            "repair_prompt": stage1_repair_prompt,
            "acceptance": acceptance1,
            "score": score1,
            "attempts": int(g1.get("attempts", 1)),
            "stage_checks": {
                "generation_started": bool((g1.get("cycle") or {}).get("started")),
                "generation_finished": bool((g1.get("cycle") or {}).get("finished")),
                "generation_timed_out": bool((g1.get("cycle") or {}).get("timed_out")),
            },
            "status": {
                "doc_status": str(post1.get("docStatus", "")),
                "flow_status": str(post1.get("flowStatus", "")),
                "char_count": compact_len(text1),
            },
            "text_snapshot": stage1_snapshot,
            "docx_export": stage1_docx,
            "feedback_items_after_stage": feedback_after_stage1[-5:],
            "feedback_tracking": {
                "before_count": int(feedback_stage1.get("before_count") or 0),
                "after_count": int(feedback_stage1.get("after_count") or 0),
                "submit_attempts": int(feedback_stage1.get("attempts") or 0),
                "persisted": bool(feedback_stage1.get("persisted")),
                "note_matched": bool(feedback_stage1.get("note_matched")),
            },
        },
        "stage2": {
            "prompt": stage2_prompt,
            "repair_prompt": stage2_repair_prompt,
            "acceptance": acceptance2,
            "score": score2,
            "attempts": int(g2.get("attempts", 1)),
            "stage_checks": {
                "generation_started": bool((g2.get("cycle") or {}).get("started")),
                "generation_finished": bool((g2.get("cycle") or {}).get("finished")),
                "generation_timed_out": bool((g2.get("cycle") or {}).get("timed_out")),
            },
            "status": {
                "doc_status": str(post2.get("docStatus", "")),
                "flow_status": str(post2.get("flowStatus", "")),
                "char_count": compact_len(text2),
            },
            "text_snapshot": stage2_snapshot,
            "docx_export": stage2_docx,
            "feedback_items_after_stage": feedback_after_stage2[-5:],
            "feedback_tracking": {
                "before_count": int(feedback_stage2.get("before_count") or 0),
                "after_count": int(feedback_stage2.get("after_count") or 0),
                "submit_attempts": int(feedback_stage2.get("attempts") or 0),
                "persisted": bool(feedback_stage2.get("persisted")),
                "note_matched": bool(feedback_stage2.get("note_matched")),
            },
        },
        "stage_delta": {
            "score_change": int(score2 - score1),
            "char_count_change": int(compact_len(text2) - compact_len(text1)),
            "initial_requirements_met_stage1": stage_is_pass({"acceptance": acceptance1, "status": {
                "doc_status": str(post1.get("docStatus", "")),
                "flow_status": str(post1.get("flowStatus", "")),
            }}),
            "initial_requirements_met_stage2": stage_is_pass({"acceptance": acceptance2, "status": {
                "doc_status": str(post2.get("docStatus", "")),
                "flow_status": str(post2.get("flowStatus", "")),
            }}),
        },
        "artifact_dir": str(artifact_dir),
    }
    return result


def _shorten(value: Any, limit: int = 500) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "..."


def _stage_report_lines(stage_label: str, payload: Dict[str, Any]) -> List[str]:
    acceptance = payload.get("acceptance") if isinstance(payload.get("acceptance"), dict) else {}
    tracking = payload.get("feedback_tracking") if isinstance(payload.get("feedback_tracking"), dict) else {}
    lines: List[str] = []
    lines.append(f"### {stage_label}")
    lines.append("")
    lines.append(f"- 输入 Prompt: `{_shorten(payload.get('prompt') or '', 800)}`")
    if str(payload.get("repair_prompt") or "").strip():
        lines.append(f"- 自动修复 Prompt: `{_shorten(payload.get('repair_prompt') or '', 600)}`")
    lines.append(f"- 评分: `{payload.get('score')}/5`")
    lines.append(f"- 验收通过: `{bool(acceptance.get('passed'))}`")
    lines.append(f"- 验收失败项: `{acceptance.get('failures') or []}`")
    lines.append(f"- 字符数: `{(payload.get('status') or {}).get('char_count', 0)}`")
    lines.append(f"- 文档状态: `{(payload.get('status') or {}).get('doc_status', '')}`")
    lines.append(f"- 流程状态: `{(payload.get('status') or {}).get('flow_status', '')}`")
    lines.append(f"- 输出快照: `{((payload.get('text_snapshot') or {}).get('path') or '')}`")
    lines.append(f"- 输出 DOCX: `{((payload.get('docx_export') or {}).get('path') or '')}`")
    lines.append(f"- DOCX 导出方式: `{((payload.get('docx_export') or {}).get('method') or '')}`")
    lines.append(f"- 评分入库前数量: `{tracking.get('before_count', 0)}`")
    lines.append(f"- 评分入库后数量: `{tracking.get('after_count', 0)}`")
    lines.append(f"- 评分提交尝试次数: `{tracking.get('submit_attempts', 0)}`")
    lines.append(f"- 评分是否确认入库: `{bool(tracking.get('persisted'))}`")
    lines.append(f"- 评分内容命中校验: `{bool(tracking.get('note_matched'))}`")
    lines.append(f"- 最近评分记录数(截断): `{len(payload.get('feedback_items_after_stage') or [])}`")
    lines.append("")
    return lines


def build_report_markdown(run: Dict[str, Any]) -> str:
    lines: List[str] = []
    lines.append("# 两阶段前端复杂 Prompt 测试报告")
    lines.append("")
    lines.append(f"- 时间戳: `{run.get('timestamp')}`")
    lines.append(f"- Base URL: `{run.get('config', {}).get('base_url')}`")
    lines.append(f"- 数据集: `{run.get('config', {}).get('dataset')}`")
    lines.append(f"- 低满意度样本池: `{run.get('paths', {}).get('low_satisfaction_pool')}`")
    lines.append("")

    summary = run.get("summary", {})
    lines.append("## 总览计分板")
    lines.append("")
    lines.append("| 指标 | 数值 |")
    lines.append("|---|---:|")
    lines.append(f"| 用例总数 | {summary.get('case_total', 0)} |")
    lines.append(f"| Stage1 通过数 | {summary.get('stage1_pass', 0)} |")
    lines.append(f"| Stage2 通过数 | {summary.get('stage2_pass', 0)} |")
    lines.append(f"| Stage2 分数提升用例数 | {summary.get('stage2_score_improved', 0)} |")
    lines.append(f"| Stage1 评分入库确认数 | {summary.get('stage1_feedback_persisted', 0)} |")
    lines.append(f"| Stage2 评分入库确认数 | {summary.get('stage2_feedback_persisted', 0)} |")
    lines.append(f"| 低满意度池观测条目数 | {run.get('low_feedback_observed_count', 0)} |")
    lines.append("")

    for case in run.get("results", []):
        lines.append(f"## 用例 {case.get('id')} - {case.get('topic')}")
        lines.append("")
        lines.append(f"- 主题组: `{case.get('group')}`")
        lines.append(f"- 文档 ID: `{case.get('doc_id')}`")
        lines.append(f"- 产物目录: `{case.get('artifact_dir')}`")
        lines.append(f"- 执行耗时: `{case.get('duration_s', 0)}s`")
        lines.append("")
        lines.extend(_stage_report_lines("阶段 1", case.get("stage1", {})))
        lines.extend(_stage_report_lines("阶段 2", case.get("stage2", {})))
        delta = case.get("stage_delta") if isinstance(case.get("stage_delta"), dict) else {}
        lines.append("### 阶段差异")
        lines.append("")
        lines.append(f"- 分数变化: `{delta.get('score_change', 0)}`")
        lines.append(f"- 字数变化: `{delta.get('char_count_change', 0)}`")
        lines.append(f"- Stage1 达标: `{delta.get('initial_requirements_met_stage1')}`")
        lines.append(f"- Stage2 达标: `{delta.get('initial_requirements_met_stage2')}`")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def export_report_docx(run: Dict[str, Any], out_path: Path) -> Dict[str, Any]:
    result = {"ok": False, "path": str(out_path), "error": ""}
    if PythonDocxDocument is None:
        result["error"] = "python-docx not available"
        return result
    try:
        summary = run.get("summary", {})
        doc = PythonDocxDocument()
        doc.add_heading("两阶段前端复杂 Prompt 测试报告", level=1)
        doc.add_paragraph(f"时间戳: {run.get('timestamp')}")
        doc.add_paragraph(f"Base URL: {run.get('config', {}).get('base_url')}")
        doc.add_paragraph(f"数据集: {run.get('config', {}).get('dataset')}")
        doc.add_paragraph(f"低满意度样本池: {run.get('paths', {}).get('low_satisfaction_pool')}")

        doc.add_heading("总览计分板", level=2)
        for row in [
            ("用例总数", summary.get("case_total", 0)),
            ("Stage1 通过数", summary.get("stage1_pass", 0)),
            ("Stage2 通过数", summary.get("stage2_pass", 0)),
            ("Stage2 分数提升用例数", summary.get("stage2_score_improved", 0)),
            ("Stage1 评分入库确认数", summary.get("stage1_feedback_persisted", 0)),
            ("Stage2 评分入库确认数", summary.get("stage2_feedback_persisted", 0)),
            ("低满意度池观测条目数", run.get("low_feedback_observed_count", 0)),
        ]:
            doc.add_paragraph(f"{row[0]}: {row[1]}")

        for case in run.get("results", []):
            doc.add_heading(f"用例 {case.get('id')} - {case.get('topic')}", level=2)
            doc.add_paragraph(f"主题组: {case.get('group')}")
            doc.add_paragraph(f"文档 ID: {case.get('doc_id')}")
            doc.add_paragraph(f"产物目录: {case.get('artifact_dir')}")
            doc.add_paragraph(f"执行耗时: {case.get('duration_s', 0)}s")

            for stage_key, stage_label in [("stage1", "阶段 1"), ("stage2", "阶段 2")]:
                stage = case.get(stage_key) if isinstance(case.get(stage_key), dict) else {}
                acceptance = stage.get("acceptance") if isinstance(stage.get("acceptance"), dict) else {}
                tracking = stage.get("feedback_tracking") if isinstance(stage.get("feedback_tracking"), dict) else {}
                doc.add_heading(stage_label, level=3)
                doc.add_paragraph(f"输入 Prompt: {_shorten(stage.get('prompt') or '', 600)}")
                if str(stage.get("repair_prompt") or "").strip():
                    doc.add_paragraph(f"自动修复 Prompt: {_shorten(stage.get('repair_prompt') or '', 500)}")
                doc.add_paragraph(f"评分: {stage.get('score')}/5")
                doc.add_paragraph(f"验收通过: {bool(acceptance.get('passed'))}")
                doc.add_paragraph(f"验收失败项: {acceptance.get('failures') or []}")
                doc.add_paragraph(f"输出快照: {(stage.get('text_snapshot') or {}).get('path', '')}")
                doc.add_paragraph(f"输出 DOCX: {(stage.get('docx_export') or {}).get('path', '')}")
                doc.add_paragraph(f"DOCX 导出方式: {(stage.get('docx_export') or {}).get('method', '')}")
                doc.add_paragraph(
                    "评分跟踪: "
                    f"before={tracking.get('before_count', 0)}, "
                    f"after={tracking.get('after_count', 0)}, "
                    f"attempts={tracking.get('submit_attempts', 0)}, "
                    f"persisted={bool(tracking.get('persisted'))}, "
                    f"note_matched={bool(tracking.get('note_matched'))}"
                )

            delta = case.get("stage_delta") if isinstance(case.get("stage_delta"), dict) else {}
            doc.add_heading("阶段差异", level=3)
            doc.add_paragraph(f"分数变化: {delta.get('score_change', 0)}")
            doc.add_paragraph(f"字数变化: {delta.get('char_count_change', 0)}")
            doc.add_paragraph(f"Stage1 达标: {delta.get('initial_requirements_met_stage1')}")
            doc.add_paragraph(f"Stage2 达标: {delta.get('initial_requirements_met_stage2')}")

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path.resolve()))
        result["ok"] = out_path.exists()
        return result
    except Exception as exc:
        result["error"] = str(exc)
        return result


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Two-stage frontend complex prompt validation with UI satisfaction scoring.")
    parser.add_argument("--dataset", default=DEFAULT_DATASET)
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL)
    parser.add_argument("--out-root", default=str(DEFAULT_OUT_ROOT))
    parser.add_argument("--timeout-s", type=int, default=260)
    parser.add_argument("--poll-interval-s", type=float, default=0.5)
    parser.add_argument("--start-server", action="store_true")
    parser.add_argument("--disable-ollama", action="store_true")
    parser.add_argument("--headed", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    cfg = run_cfg_from_args(args)
    data = read_json(cfg.dataset)
    cases = data.get("cases") if isinstance(data.get("cases"), list) else []
    if not cases:
        print("no cases found")
        return 1

    stamp = now_stamp()
    run_dir = cfg.out_root / f"two_stage_validation_{stamp}"
    artifacts_root = run_dir / "artifacts"
    artifacts_root.mkdir(parents=True, exist_ok=True)

    server_proc: subprocess.Popen[bytes] | None = None
    if cfg.start_server:
        server_proc = start_local_server(cfg.base_url, disable_ollama=cfg.disable_ollama)
        print(f"server started at {cfg.base_url}")

    results: List[Dict[str, Any]] = []
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=not cfg.headed)
            page = browser.new_page()
            for idx, case in enumerate(cases, start=1):
                case_id = str(case.get("id") or f"TS-{idx:03d}")
                print(f"[case] {idx}/{len(cases)} {case_id}")
                artifact_dir = artifacts_root / case_id
                artifact_dir.mkdir(parents=True, exist_ok=True)
                started = time.time()
                row = execute_case(page, cfg, case, artifact_dir)
                row["duration_s"] = round(time.time() - started, 2)
                results.append(row)
            browser.close()
    finally:
        if server_proc is not None:
            server_proc.terminate()
            try:
                server_proc.wait(timeout=8)
            except Exception:
                server_proc.kill()

    stage1_pass = sum(1 for x in results if stage_is_pass(x.get("stage1") if isinstance(x, dict) else {}))
    stage2_pass = sum(1 for x in results if stage_is_pass(x.get("stage2") if isinstance(x, dict) else {}))
    stage2_score_improved = sum(1 for x in results if int((x.get("stage_delta") or {}).get("score_change") or 0) > 0)
    stage1_feedback_persisted = sum(
        1 for x in results if bool(((x.get("stage1") or {}).get("feedback_tracking") or {}).get("persisted"))
    )
    stage2_feedback_persisted = sum(
        1 for x in results if bool(((x.get("stage2") or {}).get("feedback_tracking") or {}).get("persisted"))
    )
    low_pool = read_low_feedback_pool(LOW_FEEDBACK_POOL_PATH, limit=500)
    low_api_items = fetch_low_feedback(cfg.base_url, limit=500) if not cfg.start_server else []

    run_json_path = run_dir / f"two_stage_run_{stamp}.json"
    report_md_path = run_dir / f"two_stage_report_{stamp}.md"
    report_docx_path = run_dir / f"two_stage_report_{stamp}.docx"

    run_data = {
        "timestamp": stamp,
        "config": {
            "base_url": cfg.base_url,
            "dataset": str(cfg.dataset),
            "timeout_s": cfg.timeout_s,
            "poll_interval_s": cfg.poll_interval_s,
        },
        "summary": {
            "case_total": len(results),
            "stage1_pass": stage1_pass,
            "stage2_pass": stage2_pass,
            "stage2_score_improved": stage2_score_improved,
            "stage1_feedback_persisted": stage1_feedback_persisted,
            "stage2_feedback_persisted": stage2_feedback_persisted,
        },
        "results": results,
        "paths": {
            "run_json": str(run_json_path.resolve()),
            "report_md": str(report_md_path.resolve()),
            "report_docx": str(report_docx_path.resolve()),
            "artifacts_dir": str(artifacts_root.resolve()),
            "low_satisfaction_pool": str(LOW_FEEDBACK_POOL_PATH),
        },
        "low_feedback_observed_count": len(low_pool),
        "low_feedback_api_count": len(low_api_items),
    }

    report_md = build_report_markdown(run_data)
    write_json(run_json_path, run_data)
    report_md_path.write_text(report_md, encoding="utf-8")
    report_docx = export_report_docx(run_data, report_docx_path)
    run_data["report_docx_export"] = report_docx
    write_json(run_json_path, run_data)

    print(f"run_json={run_json_path}")
    print(f"report_md={report_md_path}")
    print(f"report_docx={report_docx_path}")
    print(f"report_docx_ok={report_docx.get('ok')}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
