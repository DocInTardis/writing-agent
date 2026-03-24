#!/usr/bin/env python3
"""
High-quality thesis-style generation and DOCX export batch runner.

Goals:
- Produce real DOCX exports (backend download path).
- Stress generation with multi-case prompts.
- Record quality metrics for each document.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

try:
    from docx import Document as PythonDocxDocument
except Exception:
    PythonDocxDocument = None


def now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def compact_len(text: str) -> int:
    return len(re.sub(r"\s+", "", text or ""))


def normalize_base_url(raw: str) -> str:
    u = (raw or "").strip().rstrip("/")
    if not u.startswith("http://") and not u.startswith("https://"):
        u = "http://" + u
    return u


def url_ready(base_url: str, timeout_s: float = 2.0) -> bool:
    test_url = base_url.rstrip("/") + "/favicon.ico"
    req = urllib.request.Request(test_url, method="GET")
    try:
        with urllib.request.urlopen(req, timeout=timeout_s) as resp:
            return 200 <= int(getattr(resp, "status", 0) or 0) < 500
    except Exception:
        return False


def wait_url_ready(base_url: str, timeout_s: float = 60.0) -> bool:
    start = time.time()
    while time.time() - start < timeout_s:
        if url_ready(base_url, timeout_s=2.0):
            return True
        time.sleep(0.35)
    return False


def start_server(base_url: str, *, model: str, timeout_s: int) -> subprocess.Popen[bytes]:
    parsed = urllib.parse.urlparse(base_url)
    host = parsed.hostname or "127.0.0.1"
    port = parsed.port or 8000

    env = os.environ.copy()
    env["WRITING_AGENT_USE_SVELTE"] = "1"
    env["WRITING_AGENT_USE_OLLAMA"] = "1"
    env["WRITING_AGENT_HOST"] = host
    env["WRITING_AGENT_PORT"] = str(port)
    env["PYTHONPATH"] = str(Path(__file__).resolve().parents[1])
    env["OLLAMA_MODEL"] = model
    env["OLLAMA_TIMEOUT_S"] = str(max(180, int(timeout_s)))
    env["WRITING_AGENT_STREAM_MAX_S"] = str(max(360, int(timeout_s)))
    env["WRITING_AGENT_STREAM_MAX_CAP_S"] = str(max(360, int(timeout_s)))
    env["WRITING_AGENT_STREAM_EVENT_TIMEOUT_S"] = "180"
    env["WRITING_AGENT_STREAM_EVENT_TIMEOUT_CAP_S"] = "240"
    env["WRITING_AGENT_NONSTREAM_MAX_S"] = str(max(360, int(timeout_s)))
    env["WRITING_AGENT_NONSTREAM_EVENT_TIMEOUT_S"] = "180"
    env["WRITING_AGENT_ENFORCE_INSTRUCTION_REQUIREMENTS"] = "1"

    cmd = [
        sys.executable,
        "-m",
        "uvicorn",
        "writing_agent.web.app_v2:app",
        "--host",
        host,
        "--port",
        str(port),
        "--log-level",
        "warning",
    ]
    proc = subprocess.Popen(cmd, env=env, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    if not wait_url_ready(base_url, timeout_s=75):
        proc.terminate()
        try:
            proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            proc.kill()
        raise RuntimeError(f"failed to start server at {base_url}")
    return proc


def post_json(url: str, payload: dict[str, Any], timeout_s: float = 30.0) -> dict[str, Any]:
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        url=url,
        data=data,
        method="POST",
        headers={"Content-Type": "application/json"},
    )
    with urllib.request.urlopen(req, timeout=timeout_s) as resp:
        body = resp.read().decode("utf-8", errors="replace")
    return json.loads(body) if body.strip() else {}


def get_json(url: str, timeout_s: float = 30.0) -> dict[str, Any]:
    req = urllib.request.Request(url=url, method="GET")
    with urllib.request.urlopen(req, timeout=timeout_s) as resp:
        body = resp.read().decode("utf-8", errors="replace")
    return json.loads(body) if body.strip() else {}


def create_doc(base_url: str) -> str:
    req = urllib.request.Request(base_url.rstrip("/") + "/", method="GET")
    with urllib.request.urlopen(req, timeout=30) as resp:
        final_url = str(resp.geturl() or "")
        html = resp.read().decode("utf-8", errors="ignore")

    m = re.search(r"/workbench/([a-f0-9]{32})", final_url)
    if m:
        return m.group(1)

    m = re.search(r'data-doc-id="([a-f0-9]{32})"', html)
    if m:
        return m.group(1)

    raise RuntimeError(f"cannot parse doc_id, final_url={final_url}")


def set_doc_settings(base_url: str, doc_id: str, target_chars: int) -> None:
    payload = {
        "generation_prefs": {
            "purpose": "本科毕业设计论文",
            "target_char_count": int(target_chars),
            "target_length_confirmed": True,
            "expand_outline": False,
        },
        "formatting": {
            "font_name": "宋体",
            "font_name_east_asia": "宋体",
            "font_size_name": "小四",
            "font_size_pt": 12,
            "line_spacing": 28,
            "heading1_font_name": "黑体",
            "heading1_font_name_east_asia": "黑体",
            "heading1_size_pt": 22,
            "heading2_font_name": "黑体",
            "heading2_font_name_east_asia": "黑体",
            "heading2_size_pt": 16,
            "heading3_font_name": "黑体",
            "heading3_font_name_east_asia": "黑体",
            "heading3_size_pt": 16,
        },
    }
    post_json(f"{base_url}/api/doc/{doc_id}/settings", payload, timeout_s=40.0)


def generate_stream(
    base_url: str,
    doc_id: str,
    instruction: str,
    timeout_s: int,
    *,
    text: str = "",
    compose_mode: str = "overwrite",
) -> dict[str, Any]:
    payload = {
        "instruction": instruction,
        "text": text,
        "compose_mode": compose_mode,
    }
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        url=f"{base_url}/api/doc/{doc_id}/generate/stream",
        data=data,
        method="POST",
        headers={"Content-Type": "application/json"},
    )

    start = time.time()
    event_name = ""
    event_count = 0
    final_text = ""
    final_payload: dict[str, Any] = {}
    error = ""
    trace: list[str] = []

    try:
        with urllib.request.urlopen(req, timeout=max(60, timeout_s + 60)) as resp:
            for raw in resp:
                line = raw.decode("utf-8", errors="replace").strip()
                if not line:
                    continue
                if line.startswith("event:"):
                    event_name = line[6:].strip()
                    continue
                if not line.startswith("data:"):
                    continue
                event_count += 1
                body = line[5:].strip()
                try:
                    data_obj = json.loads(body) if body else {}
                except Exception:
                    data_obj = {}
                if event_name == "delta":
                    delta = str(data_obj.get("delta") or "").strip()
                    if delta:
                        trace.append(delta[:180])
                elif event_name == "error":
                    error = str(data_obj.get("message") or data_obj.get("detail") or "stream error").strip()
                elif event_name == "final":
                    final_payload = data_obj if isinstance(data_obj, dict) else {}
                    final_text = str(final_payload.get("text") or "")
    except urllib.error.HTTPError as exc:
        detail = ""
        try:
            detail = exc.read().decode("utf-8", errors="replace")
        except Exception:
            detail = str(exc)
        return {
            "ok": False,
            "error": f"HTTP {exc.code}: {detail[:600]}",
            "duration_s": round(time.time() - start, 2),
            "event_count": event_count,
            "trace": trace[-12:],
            "text": "",
            "payload": {},
        }
    except Exception as exc:
        return {
            "ok": False,
            "error": str(exc),
            "duration_s": round(time.time() - start, 2),
            "event_count": event_count,
            "trace": trace[-12:],
            "text": "",
            "payload": {},
        }

    text = (final_text or "").strip()
    if not text and not error:
        error = "stream ended without final text"
    if error and not text:
        return {
            "ok": False,
            "error": error,
            "duration_s": round(time.time() - start, 2),
            "event_count": event_count,
            "trace": trace[-12:],
            "text": "",
            "payload": final_payload,
        }

    return {
        "ok": bool(text),
        "error": error,
        "duration_s": round(time.time() - start, 2),
        "event_count": event_count,
        "trace": trace[-12:],
        "text": text,
        "payload": final_payload,
    }


def fetch_doc_text(base_url: str, doc_id: str) -> str:
    payload = get_json(f"{base_url}/api/doc/{doc_id}", timeout_s=30.0)
    return str(payload.get("text") or "")


def export_precheck(base_url: str, doc_id: str) -> dict[str, Any]:
    url = f"{base_url}/api/doc/{doc_id}/export/check?format=docx&auto_fix=1"
    try:
        payload = get_json(url, timeout_s=25.0)
        if not isinstance(payload, dict):
            return {"ok": False, "error": "invalid_precheck_response", "can_export": False}
        return {
            "ok": bool(payload.get("ok") == 1 or payload.get("ok") is True),
            "can_export": bool(payload.get("can_export")),
            "issues": payload.get("issues") if isinstance(payload.get("issues"), list) else [],
            "warnings": payload.get("warnings") if isinstance(payload.get("warnings"), list) else [],
            "policy": str(payload.get("policy") or ""),
        }
    except Exception as exc:
        return {"ok": False, "error": str(exc), "can_export": False, "issues": [], "warnings": [], "policy": ""}


def download_docx(base_url: str, doc_id: str, out_path: Path) -> tuple[bool, str]:
    req = urllib.request.Request(f"{base_url}/download/{doc_id}.docx", method="GET")
    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            data = resp.read()
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_bytes(data)
        return out_path.exists() and out_path.stat().st_size > 0, ""
    except urllib.error.HTTPError as exc:
        detail = ""
        try:
            detail = exc.read().decode("utf-8", errors="replace")
        except Exception:
            detail = str(exc)
        return False, f"HTTP {exc.code}: {detail[:300]}"
    except Exception as exc:
        return False, str(exc)


def text_quality_metrics(text: str) -> dict[str, Any]:
    src = text or ""
    title = ""
    m = re.search(r"(?m)^#\s+(.+)$", src)
    if m:
        title = str(m.group(1) or "").strip()
    ref_stats = reference_quality_metrics(src)
    return {
        "title": title,
        "char_count": compact_len(src),
        "h1_count": len(re.findall(r"(?m)^#\s+.+$", src)),
        "h2_count": len(re.findall(r"(?m)^##\s+.+$", src)),
        "h3_count": len(re.findall(r"(?m)^###\s+.+$", src)),
        "has_abstract": bool(
            re.search(r"(?im)(^##?\s*(摘要|abstract)\b|(?:摘要|abstract)\s*[:：])", src)
        ),
        "has_keywords": bool(
            re.search(r"(?im)(^##?\s*(关键词|keywords)\b|(?:关键词|keywords)\s*[:：])", src)
        ),
        "has_references": bool(re.search(r"(?im)^##?\s*(参考文献|references)\b", src)),
        "reference_count": int(ref_stats.get("reference_count") or 0),
        "reference_unique_count": int(ref_stats.get("reference_unique_count") or 0),
        "reference_duplicate_count": int(ref_stats.get("reference_duplicate_count") or 0),
        "reference_split_doi_suspected": bool(ref_stats.get("reference_split_doi_suspected")),
        "title_generic": is_generic_title(title),
    }


def is_generic_title(title: str) -> bool:
    t = str(title or "").strip()
    if not t:
        return True
    if len(t) < 6:
        return True
    low = t.lower()
    if any(low.startswith(x) for x in ("请", "帮我", "给我", "写一篇", "生成一篇", "write ", "generate ")):
        return True
    if t in {"技术报告", "报告", "文档", "文章"}:
        return True
    if "请写" in t or "帮我生成" in t or "给我生成" in t:
        return True
    return False


def extract_reference_items(text: str) -> list[str]:
    src = str(text or "")
    lines = src.splitlines()
    start = -1
    for idx, line in enumerate(lines):
        if re.match(r"^\s*##\s*(参考文献|references?)\s*$", str(line or "").strip(), flags=re.IGNORECASE):
            start = idx + 1
            break
    if start < 0:
        return []

    end = len(lines)
    for idx in range(start, len(lines)):
        if re.match(r"^\s*#{1,3}\s+.+$", str(lines[idx] or "").strip()):
            end = idx
            break
    body = [str(x or "").rstrip() for x in lines[start:end]]
    items: list[str] = []
    cur = ""

    def _flush() -> None:
        nonlocal cur
        if cur.strip():
            items.append(cur.strip())
        cur = ""

    for raw in body:
        row = str(raw or "").strip().replace("↩", "")
        if not row:
            continue
        m_num = re.match(r"^\[(\d+)\]\s*(.+)$", row)
        if m_num:
            content = str(m_num.group(2) or "").strip()
            if re.match(r"^(?:0\.\d{2,9}/|10\.\d{2,9}/)", content) and items:
                prev = items[-1]
                if re.search(r"https?://doi\.org/[0-9.]*$", prev, flags=re.IGNORECASE):
                    items[-1] = (prev + content).strip()
                    continue
            _flush()
            cur = content
            continue
        if cur:
            cur = (cur + " " + row).strip()
        elif items:
            items[-1] = (items[-1] + " " + row).strip()
        else:
            cur = row
    _flush()

    cleaned: list[str] = []
    seen: set[str] = set()
    for item in items:
        norm = re.sub(r"\s+", " ", item).strip().lower()
        if not norm:
            continue
        if norm in seen:
            continue
        seen.add(norm)
        cleaned.append(item.strip())
    return cleaned


def reference_quality_metrics(text: str) -> dict[str, Any]:
    items = extract_reference_items(text)
    raw_count = len(items)
    unique_norm = {
        re.sub(r"\s+", " ", str(x or "").strip().lower())
        for x in items
        if str(x or "").strip()
    }
    split_doi_suspected = False
    for item in items:
        if re.search(r"https?://doi\.org/[0-9.]*\s+(?:0\.\d|10\.)", item, flags=re.IGNORECASE):
            split_doi_suspected = True
            break
    return {
        "reference_count": raw_count,
        "reference_unique_count": len(unique_norm),
        "reference_duplicate_count": max(0, raw_count - len(unique_norm)),
        "reference_split_doi_suspected": split_doi_suspected,
    }


def docx_text_quality_metrics(docx_path: Path) -> dict[str, Any]:
    out: dict[str, Any] = {
        "docx_title": "",
        "docx_title_generic": True,
        "docx_char_count": 0,
        "docx_reference_count": 0,
        "docx_reference_unique_count": 0,
        "docx_reference_duplicate_count": 0,
        "docx_reference_split_doi_suspected": False,
    }
    if PythonDocxDocument is None or not docx_path.exists():
        return out
    try:
        doc = PythonDocxDocument(str(docx_path))
        paras = [str(p.text or "").strip() for p in list(doc.paragraphs or []) if str(p.text or "").strip()]
    except Exception:
        return out
    if not paras:
        return out

    title = ""
    for row in paras:
        if row in {"目录", "目 录"}:
            continue
        title = row
        break
    out["docx_title"] = title
    out["docx_title_generic"] = is_generic_title(title)
    out["docx_char_count"] = compact_len("\n".join(paras))

    ref_start = -1
    for idx, row in enumerate(paras):
        low = row.strip().lower()
        if row.strip() == "参考文献" or low in {"references", "bibliography"}:
            ref_start = idx + 1
            break
    if ref_start < 0:
        return out

    ref_items: list[str] = []
    for row in paras[ref_start:]:
        line = str(row or "").strip()
        if not line:
            continue
        if re.match(r"^\d+(?:\.\d+)*\s+.+$", line):
            break
        if line in {"附录", "致谢"} or line.lower().startswith("acknowledg"):
            break
        m_num = re.match(r"^\[(\d+)\]\s*(.+)$", line)
        if m_num:
            ref_items.append(str(m_num.group(2) or "").strip())
        elif ref_items:
            ref_items[-1] = (ref_items[-1] + " " + line).strip()
        else:
            ref_items.append(line)

    cleaned: list[str] = []
    seen: set[str] = set()
    split_doi = False
    for item in ref_items:
        value = re.sub(r"\s+", " ", str(item or "").strip())
        value = re.sub(r"https?://doi\.org/\s+", "https://doi.org/", value, flags=re.IGNORECASE)
        if re.search(r"https?://doi\.org/[0-9.]*\s+(?:0\.\d|10\.)", value, flags=re.IGNORECASE):
            split_doi = True
        norm = value.lower()
        if not norm or norm in seen:
            continue
        seen.add(norm)
        cleaned.append(value)

    out["docx_reference_count"] = len(cleaned)
    out["docx_reference_unique_count"] = len(cleaned)
    out["docx_reference_duplicate_count"] = max(0, len(ref_items) - len(cleaned))
    out["docx_reference_split_doi_suspected"] = split_doi
    return out


def canonicalize_docx_file(docx_path: Path) -> tuple[bool, str]:
    if PythonDocxDocument is None:
        return False, "python-docx unavailable"
    if not docx_path.exists():
        return False, "docx missing"
    try:
        doc = PythonDocxDocument(str(docx_path))
        tmp = docx_path.with_suffix(".normalized.docx")
        doc.save(str(tmp))
        if not tmp.exists() or tmp.stat().st_size <= 0:
            return False, "normalized output missing"
        tmp.replace(docx_path)
        return True, ""
    except Exception as exc:
        return False, str(exc)


def docx_metrics(docx_path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {
        "docx_exists": docx_path.exists(),
        "docx_bytes": int(docx_path.stat().st_size if docx_path.exists() else 0),
        "xml_ok": False,
        "paragraphs": 0,
        "headings": 0,
        "errors": [],
    }
    if not docx_path.exists():
        result["errors"] = ["docx_missing"]
        return result

    try:
        with ZipFile(docx_path, "r") as zin:
            names = set(zin.namelist())
            required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
            missing = [x for x in required if x not in names]
            if missing:
                result["errors"].append("missing_openxml_parts")
            else:
                result["xml_ok"] = True
    except BadZipFile:
        result["errors"].append("bad_zip")
        return result
    except Exception as exc:
        result["errors"].append(f"zip_error:{exc}")
        return result

    if PythonDocxDocument is None:
        return result
    try:
        doc = PythonDocxDocument(str(docx_path))
        paragraphs = list(doc.paragraphs or [])
        result["paragraphs"] = len(paragraphs)
        heading_count = 0
        for p in paragraphs:
            style_name = str(getattr(getattr(p, "style", None), "name", "") or "").lower()
            if "heading" in style_name:
                heading_count += 1
        result["headings"] = heading_count
    except Exception as exc:
        result["errors"].append(f"python_docx_error:{exc}")
    return result


def quality_score(text_m: dict[str, Any], docx_m: dict[str, Any], precheck: dict[str, Any]) -> int:
    score = 0
    chars = int(text_m.get("char_count") or 0)
    docx_chars = int(docx_m.get("docx_char_count") or 0)
    effective_chars = max(chars, docx_chars)
    if effective_chars >= 3800:
        score += 30
    elif effective_chars >= 3200:
        score += 24
    elif effective_chars >= 2600:
        score += 18
    elif effective_chars >= 2000:
        score += 10

    if int(text_m.get("h1_count") or 0) >= 1:
        score += 8
    if int(text_m.get("h2_count") or 0) >= 6:
        score += 16
    elif int(text_m.get("h2_count") or 0) >= 4:
        score += 10

    if bool(text_m.get("has_abstract")):
        score += 12
    if bool(text_m.get("has_keywords")):
        score += 8
    if bool(text_m.get("has_references")):
        score += 14
    ref_unique = max(
        int(text_m.get("reference_unique_count") or 0),
        int(docx_m.get("docx_reference_unique_count") or 0),
    )
    if ref_unique >= 8:
        score += 6
    elif ref_unique >= 5:
        score += 3
    if bool(text_m.get("title_generic")):
        score -= 12
    if bool(text_m.get("reference_split_doi_suspected")) or bool(docx_m.get("docx_reference_split_doi_suspected")):
        score -= 10
    if bool(docx_m.get("docx_title_generic")):
        score -= 8

    if bool(precheck.get("ok")) and bool(precheck.get("can_export")):
        score += 6

    if bool(docx_m.get("xml_ok")):
        score += 4
    if int(docx_m.get("docx_bytes") or 0) >= 80_000:
        score += 2
    return min(100, score)


def needs_quality_repair(text_m: dict[str, Any], case: "Case") -> bool:
    min_chars = max(2600, int(int(case.target_chars or 2600) * 0.82))
    return (
        int(text_m.get("char_count") or 0) < min_chars
        or int(text_m.get("h2_count") or 0) < 6
        or not bool(text_m.get("has_abstract"))
        or not bool(text_m.get("has_keywords"))
        or not bool(text_m.get("has_references"))
        or int(text_m.get("reference_unique_count") or 0) < 6
        or bool(text_m.get("reference_split_doi_suspected"))
        or bool(text_m.get("title_generic"))
    )


def build_quality_repair_instruction(case: Case, text_m: dict[str, Any]) -> str:
    min_chars = max(2600, int(int(case.target_chars or 2600) * 0.82))
    missing: list[str] = []
    if bool(text_m.get("title_generic")):
        missing.append(f"将一级标题改为“{case.title}”，禁止使用指令句或口语化标题")
    if not bool(text_m.get("has_abstract")):
        missing.append("补全“摘要”章节")
    if not bool(text_m.get("has_keywords")):
        missing.append("补全“关键词/Keywords”章节")
    if not bool(text_m.get("has_references")):
        missing.append("补全“参考文献”章节")
    if int(text_m.get("h2_count") or 0) < 6:
        missing.append("增加必要的二级标题，保证结构完整")
    if int(text_m.get("char_count") or 0) < min_chars:
        missing.append(f"在保留现有结构前提下扩写技术细节、实验细节与结果分析，正文不少于{min_chars}字")
    if int(text_m.get("reference_unique_count") or 0) < 6:
        missing.append("参考文献至少8条且内容唯一，不得重复")
    if bool(text_m.get("reference_split_doi_suspected")):
        missing.append("修复被拆断的 DOI/URL，每条参考文献必须一行完整给出链接")
    missing_text = "；".join(missing) if missing else "增强论证深度与篇幅"
    return (
        f"请在当前文稿基础上进行学术化补强，不要重写标题。主题仍为《{case.title}》。"
        f"要求：{missing_text}。"
        "保留已有章节并继续扩写，每个核心章节至少2段，补充可执行细节与可量化指标。"
        "参考文献格式要求：[n] 作者. 标题. 年份. DOI或URL。禁止将一条文献拆成多条。"
        "输出完整Markdown全文，不要解释过程。"
    )


@dataclass
class Case:
    case_id: str
    title: str
    instruction: str
    target_chars: int


def default_cases() -> list[Case]:
    return [
        Case(
            case_id="HQ-101",
            title="RAG学术写作辅助系统设计与实现",
            target_chars=2900,
            instruction=(
                "请写一篇本科毕业设计论文风格的完整Markdown文稿，题目《基于RAG的学术写作辅助系统设计与实现》。"
                "必须包含并按顺序输出以下章节："
                "## 摘要、## 关键词、## Abstract、## Keywords、## 1 引言、## 2 相关工作、"
                "## 3 需求分析、## 4 系统架构设计、## 5 关键算法与实现、## 6 实验与结果分析、"
                "## 7 结论与展望、## 参考文献。"
                "要求：语言学术化、论证完整、有技术细节与指标说明，避免空话。"
            ),
        ),
        Case(
            case_id="HQ-102",
            title="双引擎生成与路由兜底机制研究",
            target_chars=2900,
            instruction=(
                "撰写题为《面向长文档生成的双引擎路由与兜底机制研究》的毕业论文风格文稿。"
                "必须包含：摘要、关键词、引言、问题定义、方法设计、路由策略、实验设计、结果分析、威胁与局限、结论、参考文献。"
                "在方法设计中明确说明resume/replay差异、节点状态持久化、失败恢复流程。"
            ),
        ),
        Case(
            case_id="HQ-103",
            title="结构化输出约束在文本编辑系统中的应用",
            target_chars=2800,
            instruction=(
                "请生成毕业论文风格文稿《结构化输出约束在智能文本编辑系统中的应用研究》。"
                "要求包含：摘要、关键词、引言、理论基础、系统实现、实验评测、工程落地建议、结论、参考文献。"
                "正文需阐明schema约束、strict模式、PATCH原子性、JSON Patch test前置断言的工程价值。"
            ),
        ),
        Case(
            case_id="HQ-104",
            title="高校教学平台微服务系统设计",
            target_chars=3000,
            instruction=(
                "撰写《高校教学平台的微服务系统设计与性能优化》本科毕业设计论文文稿。"
                "章节包含：摘要、关键词、需求分析、总体设计、数据库设计、核心服务实现、性能测试、安全与运维、结论、参考文献。"
                "需要给出可量化的测试指标与结果解释。"
            ),
        ),
        Case(
            case_id="HQ-105",
            title="可追溯引用管理与导出质量控制",
            target_chars=2800,
            instruction=(
                "请生成《可追溯引用管理与文档导出质量控制研究》论文风格文稿。"
                "必须覆盖：摘要、关键词、研究背景、方法、实现、质量控制策略、案例分析、结论、参考文献。"
                "明确说明引用核验、导出前检查、样式一致性检测的流程。"
            ),
        ),
        Case(
            case_id="HQ-106",
            title="人机协同写作工作台的交互与工程实现",
            target_chars=2800,
            instruction=(
                "写一篇《人机协同写作工作台的交互设计与工程实现》毕业论文风格文稿。"
                "要求章节包括：摘要、关键词、引言、需求与场景、系统设计、交互策略、实验评估、结论、参考文献。"
                "内容要体现编辑模式、资料模式、协作模式的设计取舍与可用性评估。"
            ),
        ),
        Case(
            case_id="HQ-107",
            title="中文学术文档自动排版与规范化输出",
            target_chars=2600,
            instruction=(
                "请写《中文学术文档自动排版与规范化输出研究》文稿，采用本科论文风格。"
                "需要有：摘要、关键词、引言、相关规范、排版策略、实现方案、实验与讨论、结论、参考文献。"
                "强调字体字号、标题层级、段落格式与导出一致性。"
            ),
        ),
        Case(
            case_id="HQ-108",
            title="面向毕业设计的写作代理质量评测体系",
            target_chars=2600,
            instruction=(
                "撰写《面向毕业设计的写作代理质量评测体系构建》论文风格文稿。"
                "章节需包含：摘要、关键词、引言、评测维度设计、数据集与实验、结果分析、改进策略、结论、参考文献。"
                "评测维度要包括可解释性、鲁棒性、结构完整性与导出质量。"
            ),
        ),
    ]


def rows_to_markdown(rows: list[dict[str, Any]], summary: dict[str, Any]) -> str:
    out: list[str] = []
    out.append("# High Quality DOCX Batch Report")
    out.append("")
    out.append(f"- Timestamp: `{summary['timestamp']}`")
    out.append(f"- Total cases: `{summary['cases_total']}`")
    out.append(f"- Passed: `{summary['cases_passed']}`")
    out.append(f"- Pass rate: `{summary['pass_rate']}%`")
    out.append(f"- Avg char count: `{summary['avg_char_count']}`")
    out.append(f"- Avg quality score: `{summary['avg_quality_score']}`")
    out.append("")
    out.append("| ID | Generate | Export | Chars | H2 | Refs(Uniq) | TitleOK | DOCX Refs | Canon | Score | Ready |")
    out.append("|---|---|---|---:|---:|---:|---|---:|---|---:|---|")
    for row in rows:
        out.append(
            "| {id} | {gen} | {exp} | {chars} | {h2} | {refs_u} | {title_ok} | {docx_refs} | {canon} | {score} | {ready} |".format(
                id=row.get("id", ""),
                gen=row.get("ok_generate", False),
                exp=row.get("ok_docx_download", False),
                chars=row.get("char_count", 0),
                h2=row.get("h2_count", 0),
                refs_u=row.get("reference_unique_count", 0),
                title_ok=(not bool(row.get("title_generic"))),
                docx_refs=row.get("docx_reference_unique_count", 0),
                canon=row.get("docx_canonicalized", False),
                score=row.get("quality_score", 0),
                ready=row.get("release_ready", False),
            )
        )
    out.append("")
    out.append("## DOCX Paths")
    out.append("")
    for row in rows:
        path = str(row.get("docx_path") or "").strip()
        if path:
            out.append(f"- `{row.get('id')}`: `{path}`")
    out.append("")
    out.append("## Errors")
    out.append("")
    errs = [r for r in rows if str(r.get("error") or "").strip()]
    if not errs:
        out.append("- None")
    else:
        for row in errs:
            out.append(f"- `{row.get('id')}`: {row.get('error')}")
    out.append("")
    return "\n".join(out)


def run_case(
    *,
    base_url: str,
    case: Case,
    output_docx_dir: Path,
    timeout_s: int,
    retries: int,
    quality_repair: bool,
) -> dict[str, Any]:
    started = time.time()
    row: dict[str, Any] = {
        "id": case.case_id,
        "title": case.title,
        "doc_id": "",
        "ok_generate": False,
        "ok_export_check": False,
        "ok_docx_download": False,
        "char_count": 0,
        "h1_count": 0,
        "h2_count": 0,
        "h3_count": 0,
        "has_abstract": False,
        "has_keywords": False,
        "has_references": False,
        "reference_count": 0,
        "reference_unique_count": 0,
        "reference_duplicate_count": 0,
        "reference_split_doi_suspected": False,
        "title_generic": False,
        "quality_score": 0,
        "docx_path": "",
        "docx_bytes": 0,
        "docx_canonicalized": False,
        "docx_title": "",
        "docx_title_generic": True,
        "docx_char_count": 0,
        "docx_reference_count": 0,
        "docx_reference_unique_count": 0,
        "docx_reference_duplicate_count": 0,
        "docx_reference_split_doi_suspected": False,
        "release_ready": False,
        "duration_s": 0.0,
        "error": "",
        "attempts": 0,
    }

    try:
        doc_id = create_doc(base_url)
        row["doc_id"] = doc_id
    except Exception as exc:
        row["error"] = f"create_doc_failed: {exc}"
        row["duration_s"] = round(time.time() - started, 2)
        return row

    attempt_errors: list[str] = []
    text = ""
    for attempt in range(1, retries + 2):
        row["attempts"] = attempt
        target_chars = max(1800, case.target_chars - (attempt - 1) * 500)
        try:
            set_doc_settings(base_url, doc_id, target_chars)
        except Exception as exc:
            attempt_errors.append(f"set_settings_attempt_{attempt}:{exc}")

        result = generate_stream(
            base_url,
            doc_id,
            case.instruction,
            timeout_s=timeout_s,
            text="",
            compose_mode="overwrite",
        )
        if result.get("ok"):
            row["ok_generate"] = True
            text = str(result.get("text") or "").strip()
            break
        attempt_errors.append(f"generate_attempt_{attempt}:{result.get('error')}")
        time.sleep(1.2)

    if not row["ok_generate"]:
        row["error"] = " | ".join(attempt_errors)[:2000]
        row["duration_s"] = round(time.time() - started, 2)
        return row

    if not text:
        text = fetch_doc_text(base_url, doc_id)

    text_m = text_quality_metrics(text)
    if quality_repair:
        for repair_round in range(1, 3):
            if not needs_quality_repair(text_m, case):
                break
            repair_instruction = build_quality_repair_instruction(case, text_m)
            repair = generate_stream(
                base_url,
                doc_id,
                repair_instruction,
                timeout_s=max(240, timeout_s),
                text=text,
                compose_mode="continue",
            )
            if repair.get("ok"):
                text = str(repair.get("text") or text)
                text_m = text_quality_metrics(text)
            else:
                attempt_errors.append(f"quality_repair_round_{repair_round}_failed:{repair.get('error')}")
                break

    row.update(text_m)

    precheck = export_precheck(base_url, doc_id)
    row["ok_export_check"] = bool(precheck.get("ok")) and bool(precheck.get("can_export"))

    docx_name = f"{case.case_id}_{doc_id}.docx"
    docx_path = output_docx_dir / docx_name
    ok_docx, docx_err = download_docx(base_url, doc_id, docx_path)
    row["ok_docx_download"] = ok_docx
    if ok_docx:
        canon_ok, canon_err = canonicalize_docx_file(docx_path)
        row["docx_canonicalized"] = bool(canon_ok)
        if not canon_ok:
            attempt_errors.append(f"docx_canonicalize_failed:{canon_err}")
        row["docx_path"] = str(docx_path)
    else:
        attempt_errors.append(f"download_docx:{docx_err}")

    d_m = docx_metrics(docx_path)
    d_text_m = docx_text_quality_metrics(docx_path)
    d_all = dict(d_m)
    d_all.update(d_text_m)
    row.update(d_text_m)
    row["docx_bytes"] = int(d_m.get("docx_bytes") or 0)
    row["quality_score"] = quality_score(text_m, d_all, precheck)
    row["release_ready"] = release_gate_passed(row, case)

    if attempt_errors:
        row["error"] = " | ".join(attempt_errors)[:2000]
    row["duration_s"] = round(time.time() - started, 2)
    return row


def release_gate_passed(row: dict[str, Any], case: Case) -> bool:
    min_chars = max(2600, int(int(case.target_chars or 2600) * 0.82))
    ref_unique = max(
        int(row.get("reference_unique_count") or 0),
        int(row.get("docx_reference_unique_count") or 0),
    )
    title_ok = not bool(row.get("title_generic")) and not bool(row.get("docx_title_generic"))
    split_doi = bool(row.get("reference_split_doi_suspected")) or bool(row.get("docx_reference_split_doi_suspected"))
    quality_ok = int(row.get("quality_score") or 0) >= 80
    char_ok = max(int(row.get("char_count") or 0), int(row.get("docx_char_count") or 0)) >= min_chars
    return bool(
        row.get("ok_generate")
        and row.get("ok_docx_download")
        and row.get("docx_canonicalized")
        and quality_ok
        and char_ok
        and title_ok
        and ref_unique >= 6
        and not split_doi
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="High-quality thesis-style docx export batch runner")
    parser.add_argument("--base-url", default=os.environ.get("WA_BASE_URL", "http://127.0.0.1:8000"))
    parser.add_argument("--start-server", action="store_true", help="Start local uvicorn if base-url is unavailable")
    parser.add_argument("--model", default=os.environ.get("OLLAMA_MODEL", "qwen2.5:7b"))
    parser.add_argument("--timeout-s", type=int, default=720, help="Per-case generation timeout (seconds)")
    parser.add_argument("--retries", type=int, default=1, help="Generation retries per case")
    parser.add_argument("--quality-repair", action="store_true", help="Enable secondary quality repair pass")
    parser.add_argument("--max-cases", type=int, default=8, help="Run first N cases")
    parser.add_argument(
        "--case-ids",
        default="",
        help="Comma-separated case IDs to run, for example HQ-101,HQ-104",
    )
    parser.add_argument("--out-root", default=".data/out")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    base_url = normalize_base_url(args.base_url)
    run_ts = now_stamp()
    out_root = Path(args.out_root).resolve()
    run_dir = out_root / f"high_quality_exports_{run_ts}"
    docx_dir = run_dir / "docx"
    run_dir.mkdir(parents=True, exist_ok=True)
    docx_dir.mkdir(parents=True, exist_ok=True)

    server_proc: subprocess.Popen[bytes] | None = None
    try:
        if not url_ready(base_url, timeout_s=2.0):
            if not args.start_server:
                print(f"base url unavailable: {base_url}", file=sys.stderr)
                return 2
            server_proc = start_server(base_url, model=args.model, timeout_s=args.timeout_s)
            print(f"server started at {base_url}")

        all_cases = default_cases()
        raw_ids = str(args.case_ids or "").strip()
        if raw_ids:
            wanted = {x.strip().upper() for x in raw_ids.split(",") if x.strip()}
            selected = [c for c in all_cases if c.case_id.upper() in wanted]
            if not selected:
                print(f"no matching case ids: {raw_ids}", file=sys.stderr)
                return 3
            cases = selected
        else:
            cases = all_cases[: max(1, int(args.max_cases))]
        rows: list[dict[str, Any]] = []
        for idx, case in enumerate(cases, start=1):
            print(f"[{idx}/{len(cases)}] {case.case_id} {case.title}")
            row = run_case(
                base_url=base_url,
                case=case,
                output_docx_dir=docx_dir,
                timeout_s=max(240, int(args.timeout_s)),
                retries=max(0, int(args.retries)),
                quality_repair=bool(args.quality_repair),
            )
            rows.append(row)
            print(
                "  generate={g} export={e} chars={c} score={s} docx={d} ready={r}".format(
                    g=row.get("ok_generate"),
                    e=row.get("ok_docx_download"),
                    c=row.get("char_count"),
                    s=row.get("quality_score"),
                    d="yes" if row.get("docx_path") else "no",
                    r=row.get("release_ready"),
                )
            )

        passed = [r for r in rows if r.get("release_ready")]
        avg_chars = int(sum(int(r.get("char_count") or 0) for r in rows) / max(1, len(rows)))
        avg_score = round(sum(float(r.get("quality_score") or 0) for r in rows) / max(1, len(rows)), 2)
        summary = {
            "timestamp": run_ts,
            "base_url": base_url,
            "cases_total": len(rows),
            "cases_passed": len(passed),
            "pass_rate": round((len(passed) / max(1, len(rows))) * 100, 2),
            "avg_char_count": avg_chars,
            "avg_quality_score": avg_score,
        }
        report = {"summary": summary, "rows": rows}

        json_path = run_dir / "high_quality_export_report.json"
        md_path = run_dir / "high_quality_export_report.md"
        json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        md_path.write_text(rows_to_markdown(rows, summary), encoding="utf-8")

        print("")
        print("done")
        print(f"json={json_path}")
        print(f"md={md_path}")
        return 0 if len(passed) == len(rows) else 1
    finally:
        if server_proc is not None:
            server_proc.terminate()
            try:
                server_proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                server_proc.kill()


if __name__ == "__main__":
    raise SystemExit(main())
