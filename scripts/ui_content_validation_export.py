"""Ui Content Validation Export command utility.

This script is part of the writing-agent operational toolchain.
"""

from __future__ import annotations

import io
import json
import re
import urllib.error
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any, Dict, Iterable, List
from zipfile import BadZipFile, ZipFile

from playwright.sync_api import Page

from scripts.ui_content_validation_text_eval import compact_len

try:
    from docx import Document as PythonDocxDocument
except Exception:  # pragma: no cover - best effort fallback import
    PythonDocxDocument = None


def should_export_docx(case: Dict[str, Any], cfg) -> bool:
    if bool(getattr(cfg, "export_docx_all", False)):
        return True
    if not bool(getattr(cfg, "export_docx_for_format", False)):
        return False
    return bool((case.get("constraints") or {}).get("format_required"))


def extract_issue_codes(items: object) -> List[str]:
    if not isinstance(items, list):
        return []
    out: List[str] = []
    for item in items:
        if isinstance(item, dict):
            code = str(item.get("code") or item.get("message") or "").strip()
            if code:
                out.append(code)
        elif item is not None:
            txt = str(item).strip()
            if txt:
                out.append(txt)
    return out


def classify_precheck_warnings(warnings: Iterable[str]) -> Dict[str, List[str]]:
    blocking: List[str] = []
    non_blocking: List[str] = []
    for raw in warnings:
        code = str(raw or "").strip()
        if not code:
            continue
        lower = code.lower()
        if lower in {"autofix_applied"}:
            non_blocking.append(code)
            continue
        # Keep strict for obviously risky warnings.
        if any(tok in lower for tok in ("compat", "corrupt", "invalid", "broken", "mismatch")):
            blocking.append(code)
            continue
        non_blocking.append(code)
    return {"blocking": blocking, "non_blocking": non_blocking}


def _http_get_json(url: str, timeout_s: float = 20.0) -> Dict[str, Any]:
    req = urllib.request.Request(url=url, method="GET")
    try:
        with urllib.request.urlopen(req, timeout=timeout_s) as resp:
            raw = resp.read().decode("utf-8", errors="ignore")
    except Exception:
        return {}
    try:
        payload = json.loads(raw)
    except Exception:
        return {}
    return payload if isinstance(payload, dict) else {}


def fetch_export_precheck(base_url: str, doc_id: str) -> Dict[str, Any]:
    result = {
        "attempted": True,
        "ok": False,
        "can_export": None,
        "policy": "",
        "issues": [],
        "warnings": [],
        "error": "",
    }
    url = f"{base_url.rstrip('/')}/api/doc/{doc_id}/export/check?format=docx&auto_fix=1"
    payload = _http_get_json(url, timeout_s=20.0)
    if not payload:
        result["error"] = "empty_or_invalid_response"
        return result
    result["ok"] = bool(payload.get("ok") == 1 or payload.get("ok") is True)
    result["can_export"] = bool(payload.get("can_export")) if payload.get("can_export") is not None else None
    result["policy"] = str(payload.get("policy") or "").strip()
    result["issues"] = extract_issue_codes(payload.get("issues"))
    result["warnings"] = extract_issue_codes(payload.get("warnings"))
    if not result["ok"]:
        result["error"] = "precheck_not_ok"
    return result


def probe_docx_download_headers(base_url: str, doc_id: str) -> Dict[str, Any]:
    result = {
        "attempted": True,
        "ok": False,
        "error": "",
        "status_code": 0,
        "export_backend": "",
        "style_path": "",
        "validation": "",
        "warn": "",
        "policy": "",
    }
    url = f"{base_url.rstrip('/')}/download/{urllib.parse.quote(str(doc_id))}.docx"
    req = urllib.request.Request(url=url, method="GET")
    try:
        with urllib.request.urlopen(req, timeout=35) as resp:
            result["status_code"] = int(getattr(resp, "status", 200) or 200)
            # Read a small prefix to force response headers/data path but avoid extra memory.
            resp.read(256)
            headers = resp.headers
            result["export_backend"] = str(headers.get("X-Docx-Export-Backend") or "").strip()
            result["style_path"] = str(headers.get("X-Docx-Style-Path") or "").strip()
            result["validation"] = str(headers.get("X-Docx-Validation") or "").strip().lower()
            result["warn"] = str(headers.get("X-Docx-Warn") or "").strip()
            result["policy"] = str(headers.get("X-Docx-Export-Policy") or "").strip()
            result["ok"] = True
            return result
    except urllib.error.HTTPError as exc:
        result["status_code"] = int(getattr(exc, "code", 0) or 0)
        try:
            detail = exc.read().decode("utf-8", errors="ignore")
        except Exception:
            detail = str(exc)
        result["error"] = f"http:{result['status_code']}:{detail[:240]}"
        return result
    except Exception as exc:
        result["error"] = str(exc)
        return result


def validate_docx_style_conformance(docx_path: Path, *, format_sensitive: bool) -> Dict[str, Any]:
    result = {
        "attempted": True,
        "ok": False,
        "passed": False,
        "failures": [],
        "metrics": {},
        "error": "",
    }
    if not docx_path.exists():
        result["error"] = "docx_not_found"
        result["failures"] = ["docx_not_found"]
        return result
    failures: List[str] = []
    metrics: Dict[str, Any] = {}
    try:
        raw = docx_path.read_bytes()
    except Exception as exc:
        result["error"] = str(exc)
        result["failures"] = ["docx_read_failed"]
        return result

    xml_issues: List[str] = []
    document_xml = ""
    styles_xml = ""
    try:
        with ZipFile(io.BytesIO(raw), "r") as zin:
            names = set(zin.namelist())
            required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
            missing = [x for x in required if x not in names]
            if missing:
                xml_issues.append("missing_openxml_parts")
            for name in names:
                if not name.lower().endswith(".xml"):
                    continue
                try:
                    ET.fromstring(zin.read(name))
                except Exception:
                    xml_issues.append(f"xml_parse_failed:{name}")
            if "word/document.xml" in names:
                document_xml = zin.read("word/document.xml").decode("utf-8", errors="ignore")
            if "word/styles.xml" in names:
                styles_xml = zin.read("word/styles.xml").decode("utf-8", errors="ignore")
    except BadZipFile:
        xml_issues.append("bad_zip")
    except Exception:
        xml_issues.append("zip_open_failed")

    if xml_issues:
        failures.extend(xml_issues)

    rfonts_doc = len(re.findall(r"w:rFonts", document_xml))
    rfonts_styles = len(re.findall(r"w:rFonts", styles_xml))
    center_doc = len(re.findall(r"w:jc[^>]+w:val=['\"]center['\"]", document_xml, flags=re.IGNORECASE))
    metrics["rfonts_document"] = rfonts_doc
    metrics["rfonts_styles"] = rfonts_styles
    metrics["center_alignment_tags"] = center_doc

    paragraph_count = 0
    non_empty_count = 0
    centered_preview = 0
    run_font_count = 0
    run_with_text = 0
    unique_fonts: List[str] = []
    if PythonDocxDocument is not None:
        try:
            doc = PythonDocxDocument(str(docx_path))
            paragraphs = list(doc.paragraphs or [])
            paragraph_count = len(paragraphs)
            non_empty = [p for p in paragraphs if str(p.text or "").strip()]
            non_empty_count = len(non_empty)
            for p in non_empty[:12]:
                align = p.paragraph_format.alignment
                if align is not None and int(align) == 1:
                    centered_preview += 1
            for p in non_empty[:140]:
                for run in p.runs:
                    if not str(run.text or "").strip():
                        continue
                    run_with_text += 1
                    fname = str(run.font.name or "").strip()
                    if fname:
                        run_font_count += 1
                        unique_fonts.append(fname.lower())
        except Exception as exc:
            failures.append("python_docx_open_failed")
            result["error"] = str(exc)

    metrics["paragraph_total"] = paragraph_count
    metrics["paragraph_non_empty"] = non_empty_count
    metrics["centered_preview_count"] = centered_preview
    metrics["run_with_text"] = run_with_text
    metrics["run_with_explicit_font"] = run_font_count
    unique_font_set = sorted(set(unique_fonts))
    metrics["unique_fonts"] = unique_font_set[:16]
    metrics["unique_fonts_count"] = len(unique_font_set)

    if format_sensitive:
        if centered_preview <= 0 and center_doc <= 0:
            failures.append("title_or_heading_not_centered")
        if (rfonts_doc + rfonts_styles) <= 0:
            failures.append("font_metadata_missing")
        if run_with_text > 0 and run_font_count == 0 and (rfonts_doc + rfonts_styles) <= 2:
            failures.append("explicit_font_binding_missing")
        if len(unique_font_set) >= 9:
            failures.append("font_family_too_fragmented")

    result["ok"] = True
    result["failures"] = failures
    result["metrics"] = metrics
    result["passed"] = len(failures) == 0
    return result


def export_docx_from_text_fallback(text: str, out_path: Path) -> Dict[str, Any]:
    result = {"attempted": True, "ok": False, "path": "", "error": ""}
    if PythonDocxDocument is None:
        result["error"] = "python-docx not available"
        return result
    try:
        doc = PythonDocxDocument()
        normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not normalized:
            doc.add_paragraph("")
        else:
            for block in re.split(r"\n{2,}", normalized):
                b = block.strip()
                if not b:
                    continue
                if b.startswith("# "):
                    doc.add_heading(b[2:].strip() or "Untitled", level=1)
                    continue
                if b.startswith("## "):
                    doc.add_heading(b[3:].strip() or "Section", level=2)
                    continue
                if b.startswith("### "):
                    doc.add_heading(b[4:].strip() or "Subsection", level=3)
                    continue
                for line in b.split("\n"):
                    clean = line.strip()
                    if clean:
                        doc.add_paragraph(clean)
        doc.save(str(out_path.resolve()))
        result["ok"] = out_path.exists()
        result["path"] = str(out_path)
        return result
    except Exception as exc:
        result["error"] = str(exc)
        return result


def export_docx_if_requested(
    page: Page,
    artifact_dir: Path,
    case_id: str,
    fallback_text: str = "",
) -> Dict[str, Any]:
    result = {"attempted": True, "ok": False, "path": "", "error": "", "method": ""}
    out_path = artifact_dir / f"{case_id}.docx"
    errors: List[str] = []

    # First try the standard top-bar export button in the UI.
    for _ in range(2):
        try:
            button = page.locator(".top-actions button", has_text="Word")
            if button.count() == 0:
                button = page.locator("button", has_text="Word")
            if button.count() == 0:
                errors.append("export_button_not_found")
                break
            page.wait_for_timeout(900)
            with page.expect_download(timeout=45000) as dl_info:
                button.first.click()
            dl = dl_info.value
            dl.save_as(str(out_path.resolve()))
            result["ok"] = out_path.exists()
            result["path"] = str(out_path)
            result["method"] = "ui_button"
            if result["ok"]:
                return result
        except Exception as exc:
            errors.append(str(exc))
            page.wait_for_timeout(1200)

    # Fallback to direct browser download route for this doc id.
    try:
        doc_id = str(page.evaluate("window.__waGetStore ? window.__waGetStore('docId') : ''") or "").strip()
        if not doc_id:
            errors.append("missing_doc_id_for_direct_download")
            result["error"] = " | ".join([e for e in errors if e])
            return result
        with page.expect_download(timeout=45000) as dl_info:
            page.evaluate(
                """
                (docId) => {
                  const a = document.createElement('a');
                  a.href = `/download/${docId}.docx`;
                  a.download = '';
                  a.style.display = 'none';
                  document.body.appendChild(a);
                  a.click();
                  a.remove();
                }
                """,
                doc_id,
            )
        dl = dl_info.value
        dl.save_as(str(out_path.resolve()))
        result["ok"] = out_path.exists()
        result["path"] = str(out_path)
        result["method"] = "direct_download_fallback"
        if result["ok"]:
            return result
        errors.append("download_saved_but_file_missing")
    except Exception as exc:
        errors.append(str(exc))

    local_fallback = export_docx_from_text_fallback(fallback_text, out_path)
    if local_fallback.get("ok"):
        result["ok"] = True
        result["path"] = str(out_path)
        result["method"] = "local_text_fallback"
        return result
    if local_fallback.get("error"):
        errors.append(f"local_fallback:{local_fallback.get('error')}")

    result["error"] = " | ".join([e for e in errors if e]) or "unknown_export_failure"
    return result
