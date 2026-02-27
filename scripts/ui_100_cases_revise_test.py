# coding: utf-8
"""Ui 100 Cases Revise Test command utility.

This script is part of the writing-agent operational toolchain.
"""

from __future__ import annotations

from pathlib import Path
from playwright.sync_api import sync_playwright
from urllib.request import Request, urlopen
from docx import Document
import os
import random
import re
import time

BASE = os.environ.get("WA_BASE", "http://127.0.0.1:8000/")
OUT = Path(".data") / "out"
OUT.mkdir(parents=True, exist_ok=True)
LOG = OUT / "ui_100_cases_revise_test.log"

BASE_SEED = 7

TOPICS = [
    "教室管理系统",
    "工资管理系统",
    "宿舍管理系统",
    "图书管理系统",
    "选课系统",
    "设备维修管理系统",
    "库存管理系统",
    "考试管理系统",
    "项目进度管理系统",
    "车辆管理系统",
    "固定资产管理系统",
    "客户关系管理系统",
    "门禁管理系统",
    "实验室预约系统",
    "会议室管理系统",
]
PURPOSES = [
    "课程设计",
    "课程报告",
    "毕业设计",
    "项目总结",
    "调研报告",
    "参考作用",
]
LENGTHS = [
    "1500字",
    "1800字",
    "2000字",
    "2200字",
    "约2000字",
    "大约1800字",
]
FORMATS = [
    "默认格式",
    "用默认格式",
    "格式默认就行",
    "没有格式要求",
]
SCOPE_HINTS = [
    "主要写系统的设计与实现",
    "重点放在需求分析和总体设计",
    "写清楚业务流程和数据库设计",
    "包括系统架构、功能模块和测试",
]
TONE = [
    "语气正式一点",
    "偏学术",
    "更工程化一点",
    "语言简洁",
]


def log(msg: str) -> None:
    ts = time.strftime("%H:%M:%S")
    prev = LOG.read_text(encoding="utf-8") if LOG.exists() else ""
    LOG.write_text(prev + f"[{ts}] {msg}\n", encoding="utf-8")


def last_system_message(page) -> str:
    return page.evaluate(
        """
() => {
  const rows = Array.from(document.querySelectorAll('#chatHistory .chat-msg.system'));
  const last = rows[rows.length - 1];
  return last ? last.innerText : '';
}
"""
    )


def answer_for(question: str, length_hint: str, topic: str) -> str:
    q = (question or "").strip()
    if not q:
        return "默认"
    if re.search(r"模板|格式", q):
        return "默认"
    if re.search(r"目标长度|字数|多少字|页", q):
        return length_hint
    if re.search(r"目的|用途", q):
        return "课程设计"
    if re.search(r"范围|包含哪些部分|内容", q):
        return "系统的设计与实现，包含需求分析、总体设计、数据库设计和测试"
    if re.search(r"标题|题目", q):
        return topic
    return "默认"


def wait_ready(page) -> None:
    page.wait_for_function("() => window.__wa_ready === true", timeout=60000)


def close_modal(page) -> None:
    for _ in range(5):
        try:
            if not page.is_visible("#modalRoot:not(.hidden)"):
                return
        except Exception:
            return
        try:
            btns = page.locator("#modalFoot button")
            if btns.count() > 0:
                btns.first.click(force=True)
            else:
                page.click("#modalRoot [data-close='1']", force=True)
        except Exception:
            pass
        page.wait_for_timeout(200)


def wait_generate_done(page, timeout_ms=90000) -> None:
    page.wait_for_function(
        """
() => {
  const el = document.querySelector('#docStatus');
  if (!el) return false;
  const t = el.innerText || '';
  return t.includes('完成');
}
""",
        timeout=timeout_ms,
    )


def slugify(name: str) -> str:
    s = re.sub(r"[\\\\/:*?\"<>|]+", "_", name).strip()
    return re.sub(r"\\s+", "_", s)[:40] or "doc"


def export_docx(page, idx: int, topic: str) -> Path:
    doc_id = page.evaluate("() => document.querySelector('.app')?.dataset?.docId || ''")
    if not doc_id:
        raise RuntimeError("missing doc id")
    url = f"{BASE}download/{doc_id}.docx"
    filename = f"ui_case_{idx:03d}_{slugify(topic)}.docx"
    path = OUT / filename
    req = Request(url, headers={"User-Agent": "ui-test"})
    with urlopen(req, timeout=60) as resp:
        data = resp.read()
    path.write_bytes(data)
    return path


def validate_docx(path: Path) -> dict:
    doc = Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    toc_found = any(t in p for p in paras for t in ("目录", "目 录"))
    ref_found = any("参考文献" in p for p in paras)
    bad_mark = any("####" in p or "乱码占位" in p or "？？？" in p for p in paras)
    heading_styles = 0
    for p in doc.paragraphs:
        style_name = ""
        try:
            style_name = (p.style.name or "").lower()
        except Exception:
            style_name = ""
        if style_name and ("heading" in style_name or "标题" in style_name):
            if p.text and p.text.strip():
                heading_styles += 1
    headings = [
        p for p in paras if re.match(r"^第\\d+章\\s*\\S+", p) or re.match(r"^\\d+\\.\\d+\\s+\\S+", p)
    ]
    return {
        "toc": toc_found,
        "ref": ref_found,
        "bad": bad_mark,
        "headings": max(len(headings), heading_styles),
        "paras": len(paras),
    }


def do_revision(page, instruction: str) -> bool:
    page.fill("#instruction", instruction)
    try:
        page.wait_for_function(
            "() => !document.querySelector('#btnRevise')?.disabled",
            timeout=8000,
        )
    except Exception:
        return False
    page.click("#btnRevise")
    try:
        page.wait_for_selector("#modalRoot:not(.hidden)", timeout=20000)
    except Exception:
        return False
    # apply revision
    btn = page.locator("#modalFoot button", has_text="应用修订")
    if btn.count() == 0:
        btn = page.locator("#modalFoot button", has_text="应用修改")
    if btn.count() == 0:
        btn = page.locator("#modalFoot button").last
    btn.first.click(force=True)
    page.wait_for_timeout(400)
    return True


def build_prompt(topic: str, rng: random.Random) -> str:
    return "，".join(
        [
            rng.choice(FORMATS) + "给我写一份" + topic + "的报告",
            "没有模板",
            "用于" + rng.choice(PURPOSES),
            rng.choice(LENGTHS),
            rng.choice(SCOPE_HINTS),
            rng.choice(TONE),
        ]
    )


def run_case(page, idx: int) -> None:
    rng = random.Random(BASE_SEED + idx)
    topic = rng.choice(TOPICS)
    length_hint = rng.choice(LENGTHS)
    prompt = build_prompt(topic, rng)

    page.goto(BASE, wait_until="domcontentloaded")
    wait_ready(page)
    page.wait_for_selector("#instruction")
    close_modal(page)

    page.fill("#instruction", prompt)
    page.click("#btnGenerate")

    # answer pending questions
    for _ in range(12):
        try:
            page.wait_for_function(
                """
() => {
  const btn = document.querySelector('#btnGenerate');
  return btn && btn.dataset.mode === 'reply';
}
""",
                timeout=12000,
            )
        except Exception:
            break
        q = last_system_message(page)
        a = answer_for(q, length_hint, topic)
        page.fill("#instruction", a)
        page.click("#btnGenerate")
        page.wait_for_timeout(200)

    try:
        wait_generate_done(page)
    except Exception:
        log(f"case {idx:03d} generation timeout")
        return

    # revisions
    ok1 = do_revision(page, "把语言调整为更正式的学术表达")
    ok2 = do_revision(page, "把本章小节衔接更顺畅，避免重复")

    # export + validate
    docx_path = export_docx(page, idx, topic)
    check = validate_docx(docx_path)

    # log summary
    text_len = page.evaluate("() => (document.querySelector('#source')?.value || '').length")
    log(
        f"case {idx:03d} topic={topic} len={text_len} revise_ok={ok1 and ok2} "
        f"toc={check['toc']} ref={check['ref']} bad={check['bad']} "
        f"headings={check['headings']} paras={check['paras']} file={docx_path.name}"
    )


if __name__ == "__main__":
    case_count = int(os.environ.get("UI_CASE_COUNT", "100"))
    start = int(os.environ.get("UI_CASE_START", "1"))
    end = int(os.environ.get("UI_CASE_END", str(case_count)))
    headless = os.environ.get("UI_HEADLESS", "1") != "0"
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=headless)
        page = browser.new_page()
        page.set_default_timeout(30000)
        page.set_default_navigation_timeout(60000)
        for i in range(start, min(end, case_count) + 1):
            try:
                run_case(page, i)
            except Exception as e:
                log(f"case {i:03d} failed: {e}")
                try:
                    close_modal(page)
                except Exception:
                    pass
                # attempt continue
        browser.close()
